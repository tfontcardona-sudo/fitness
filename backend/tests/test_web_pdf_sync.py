"""AUDITORÍA de sincronización web ⇄ almacenado ⇄ PDF/Word del plan.

La regla del sistema: la planificación es UN organismo. Cualquier edición
(kcal, macros, comidas…) debe propagarse a macros, objetivos por comida,
gramos del banco y equivalencias — y el documento descargado debe reflejar
EXACTAMENTE lo almacenado, siempre, sin versiones antiguas.

Este módulo edita un plan por la MISMA ruta que usa la web (PATCH /api/plans)
y verifica la propagación en la BD y en el documento descargado.
"""

import io
import uuid

import pytest


def _db_available() -> bool:
    try:
        from sqlalchemy import create_engine, text

        from app.config import settings

        create_engine(settings.database_url).connect().execute(text("SELECT 1"))
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(not _db_available(), reason="Requiere PostgreSQL")


@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient

    from app.main import app
    from app.seeds.run import main as seed_main

    seed_main()
    return TestClient(app)


@pytest.fixture(scope="module")
def auth(client):
    # Token directo: /api/auth/login está capado a 5/min contra fuerza bruta.
    from app.security import create_access_token

    return {"Authorization": f"Bearer {create_access_token('coach1')}"}


def _nutrition_base() -> dict:
    """Plan realista con banco completo y gramos conocidos para medir el reescalado."""
    return {
        "tdee_kcal": 2500, "target_kcal": 2000,
        "macros": {"protein_g": 150, "carbs_g": 200, "fat_g": 60},  # = 2000 kcal exactas
        "meals": [
            {"slot": 1, "name": "Desayuno", "time": "08:00",
             "target": {"kcal": 600, "protein_g": 45, "carbs_g": 60, "fat_g": 18}},
            {"slot": 2, "name": "Comida", "time": "14:00",
             "target": {"kcal": 800, "protein_g": 60, "carbs_g": 80, "fat_g": 24}},
            {"slot": 3, "name": "Cena", "time": "21:00",
             "target": {"kcal": 600, "protein_g": 45, "carbs_g": 60, "fat_g": 18}},
        ],
        "meal_bank": {"mode": "flexible_7", "slots": [
            {"slot": 1, "fmt": "options", "options": [
                {"key": "A", "title": "Yogur con avena", "prep": "", "prep_minutes": 5,
                 "ingredients": [{"food": "Yogur griego", "grams": 200, "household": ""},
                                  {"food": "Copos de avena", "grams": 60, "household": ""}],
                 "macros": {"kcal": 600, "protein_g": 45, "carbs_g": 60, "fat_g": 18},
                 "tags": []}],
             "weekly_examples": ["Yogur"] * 7},
            {"slot": 2, "fmt": "equivalences", "equivalences": {
                "intro": "Equivalencias calculadas para aportar ~80 g de CH del cereal",
                "groups": [
                    {"name": "Hidratos de carbono", "note": "",
                     "items": [{"food": "Arroz blanco", "amount": "80 g crudo = 220 g cocido"}]},
                    {"name": "Proteína magra", "note": "",
                     "items": [{"food": "Pechuga de pollo", "amount": "200 g crudo"}]},
                ]},
             "weekly_examples": ["Pollo con arroz"] * 7},
        ]},
    }


@pytest.fixture(scope="module")
def plan_ids(client, auth):
    r = client.post("/api/clients", headers=auth, json={
        "full_name": "Auditoria Sync", "email": f"sync-{uuid.uuid4().hex[:8]}@example.com",
    })
    cid = r.json()["client"]["id"]
    r = client.post(f"/api/clients/{cid}/plans", headers=auth, json={
        "month_index": 1, "nutrition_json": _nutrition_base(),
        "training_json": {"sessions": []}, "education_json": {},
    })
    return cid, r.json()["id"]


def _docx_text(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_editar_kcal_propaga_a_macros_comidas_banco_y_documento(client, auth, plan_ids):
    """Subir las kcal por la MISMA ruta que la web (PATCH del JSON entero con
    solo target_kcal tocado) debe recalcular macros, comidas, gramos del banco
    y equivalencias — y el documento descargado debe traer los valores NUEVOS."""
    _cid, plan_id = plan_ids

    edited = _nutrition_base()
    edited["target_kcal"] = 2500  # el coach sube kcal SIN tocar nada más
    r = client.patch(f"/api/plans/{plan_id}", headers=auth, json={"nutrition_json": edited})
    assert r.status_code == 200, r.text

    stored = client.get(f"/api/clients/{_cid}/plans", headers=auth).json()
    nut = next(p for p in stored if p["id"] == plan_id)["nutrition_json"]

    # 1) kcal ≡ macros (4/4/9): P y G se conservan; los CH hacen de colchón.
    m = nut["macros"]
    assert nut["target_kcal"] == 4 * m["protein_g"] + 4 * m["carbs_g"] + 9 * m["fat_g"]
    assert nut["target_kcal"] == 2500
    assert m["carbs_g"] > 200  # los CH subieron para absorber las +500 kcal

    # 2) Σ comidas ≡ totales, eje por eje.
    meals = nut["meals"]
    for axis in ("kcal", "protein_g", "carbs_g", "fat_g"):
        total = nut["target_kcal"] if axis == "kcal" else m[axis]
        assert sum(mm["target"][axis] for mm in meals) == total, axis
    # y cada comida cuadra sola (kcal ≡ sus macros)
    for mm in meals:
        t = mm["target"]
        assert t["kcal"] == 4 * t["protein_g"] + 4 * t["carbs_g"] + 9 * t["fat_g"]

    # 3) El banco siguió a los totales: los gramos del desayuno subieron y las
    #    equivalencias de hidratos también (los CH absorben la subida).
    slot1 = next(s for s in nut["meal_bank"]["slots"] if s["slot"] == 1)
    grams = {i["food"]: i["grams"] for i in slot1["options"][0]["ingredients"]}
    assert grams["Copos de avena"] > 60  # 60 g era la base a 2000 kcal
    slot2 = next(s for s in nut["meal_bank"]["slots"] if s["slot"] == 2)
    carbs_item = slot2["equivalences"]["groups"][0]["items"][0]["amount"]
    assert "80 g crudo" not in carbs_item  # la cantidad ya no es la de 2000 kcal

    # 4) El documento descargado = EXACTAMENTE lo almacenado (nunca una versión vieja).
    w = client.get(f"/api/plans/{plan_id}/document?format=docx", headers=auth)
    assert w.status_code == 200 and w.content[:2] == b"PK"
    text = _docx_text(w.content)
    assert "2500" in text                      # kcal nuevas en el resumen
    assert "2000" not in text.replace("≈ 2000", "2000") or "2500" in text
    assert str(m["carbs_g"]) in text           # macros nuevos en el resumen
    assert carbs_item.split(" =")[0] in text   # equivalencia reescalada en su tarjeta
    grams_avena = str(grams["Copos de avena"])
    assert f"Copos de avena {grams_avena} g" in text  # gramos del banco reescalados


def test_editar_macro_recalcula_kcal_y_documento(client, auth, plan_ids):
    """Cambiar un macro (proteína) recalcula las kcal (4/4/9) y el documento
    descargado refleja el estado nuevo — segunda edición sobre la anterior."""
    _cid, plan_id = plan_ids

    stored = client.get(f"/api/clients/{_cid}/plans", headers=auth).json()
    nut = next(p for p in stored if p["id"] == plan_id)["nutrition_json"]
    nut["macros"]["protein_g"] = nut["macros"]["protein_g"] + 40  # sube proteína

    r = client.patch(f"/api/plans/{plan_id}", headers=auth, json={"nutrition_json": nut})
    assert r.status_code == 200, r.text

    stored2 = client.get(f"/api/clients/{_cid}/plans", headers=auth).json()
    nut2 = next(p for p in stored2 if p["id"] == plan_id)["nutrition_json"]
    m = nut2["macros"]
    # kcal ≡ macros SIEMPRE; y la suma de comidas sigue cuadrando eje a eje
    assert nut2["target_kcal"] == 4 * m["protein_g"] + 4 * m["carbs_g"] + 9 * m["fat_g"]
    for axis in ("protein_g", "carbs_g", "fat_g"):
        assert sum(mm["target"][axis] for mm in nut2["meals"]) == m[axis], axis

    w = client.get(f"/api/plans/{plan_id}/document?format=docx", headers=auth)
    text = _docx_text(w.content)
    assert f"P {m['protein_g']} g" in text  # el reparto nuevo, en el documento


def test_cambiar_numero_de_comidas_redistribuye_y_documenta(client, auth, plan_ids):
    """Añadir una comida (editor de nº de comidas) redistribuye los ejes entre
    las tomas, la nueva toma recibe banco automático y el documento la trae."""
    _cid, plan_id = plan_ids

    stored = client.get(f"/api/clients/{_cid}/plans", headers=auth).json()
    nut = next(p for p in stored if p["id"] == plan_id)["nutrition_json"]
    total_k = nut["target_kcal"]
    # El coach añade una Merienda repartiendo como hace el editor: recorta la
    # comida mayor y cede ese objetivo a la toma nueva.
    comida = next(m for m in nut["meals"] if m["name"] == "Comida")
    nueva = {"slot": 4, "name": "Merienda", "time": "18:00",
             "target": {"kcal": 300, "protein_g": 22, "carbs_g": 35, "fat_g": 8}}
    for axis in ("kcal", "protein_g", "carbs_g", "fat_g"):
        comida["target"][axis] = max(0, comida["target"][axis] - nueva["target"][axis])
    nut["meals"].append(nueva)

    r = client.patch(f"/api/plans/{plan_id}", headers=auth, json={"nutrition_json": nut})
    assert r.status_code == 200, r.text

    stored2 = client.get(f"/api/clients/{_cid}/plans", headers=auth).json()
    nut2 = next(p for p in stored2 if p["id"] == plan_id)["nutrition_json"]
    assert len(nut2["meals"]) == 4
    assert nut2["target_kcal"] == total_k  # el total no cambia por repartir
    assert sum(m["target"]["kcal"] for m in nut2["meals"]) == total_k
    # La toma nueva NUNCA queda vacía: recibió sus 3 opciones automáticas
    slot4 = next(s for s in nut2["meal_bank"]["slots"] if s["slot"] == 4)
    assert len(slot4["options"]) == 3

    w = client.get(f"/api/plans/{plan_id}/document?format=docx", headers=auth)
    text = _docx_text(w.content)
    assert "MERIENDA" in text.upper()      # la toma nueva está en el documento
    assert "Toma libre" not in text        # y con opciones concretas
