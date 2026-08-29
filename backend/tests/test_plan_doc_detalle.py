"""El PDF del cliente explica, no solo lista.

Estas regresiones cubren lo que la auditoría encontró perdido: el POR QUÉ del
enfoque y el margen de maniobra se pautaban en el plan y nunca llegaban al
documento, y las cifras salían sin contexto ni formato español.
"""
import io

from docx import Document

from app.services.docs.plan_doc import _ajuste_text, _macro_lines, generate_plan_doc
from app.services.docs.word_base import DocBrand

BRAND = DocBrand(name="DQR", color_primary="#8B1A2B", color_secondary="#4A7BA8",
                 font_family="Calibri", contact_email="hola@dqr.es")

NUTRICION = {
    "tdee_kcal": 2650, "target_kcal": 2200,
    "macros": {"carbs_g": 210, "protein_g": 165, "fat_g": 68},
    "rationale": "Déficit moderado y proteína alta para no perder fuerza.",
    "flexibility_rules": ["Una comida libre a la semana."],
    "refeed_or_break": "Recarga de carbohidratos el día de pierna.",
    "meals": [{"slot": 1, "name": "Desayuno", "time": "08:00",
               "target": {"kcal": 700, "protein_g": 50, "carbs_g": 70, "fat_g": 20}}],
    "meal_bank": {"mode": "flexible", "slots": []},
}


def _texto(**kwargs) -> str:
    data = generate_plan_doc(
        brand=BRAND, client_name="Mario", month_index=3, goal_type="fat_loss",
        diet_mode=kwargs.pop("diet_mode", "flexible"),
        nutrition=kwargs.pop("nutrition", NUTRICION),
        training={}, education=kwargs.pop("education", {}),
        include_training=False, include_nutrition=True,
        **kwargs,
    )
    doc = Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            partes += [c.text for c in row.cells]
    return "\n".join(partes)


def test_el_documento_explica_el_porque_y_el_margen():
    texto = _texto()
    # El criterio del coach viaja al cliente (antes se quedaba en el sistema).
    assert "POR QUÉ ESTE ENFOQUE" in texto
    assert "Déficit moderado y proteína alta" in texto
    assert "Una comida libre a la semana." in texto
    assert "Recarga de carbohidratos el día de pierna." in texto
    # Mapa del documento y origen de la cifra.
    assert "EN ESTE DOCUMENTO" in texto
    assert "De dónde sale tu cifra" in texto
    assert "2.650" in texto and "2.200" in texto   # formato es-ES


def test_sin_argumentario_no_se_pintan_secciones_vacias():
    nut = {k: v for k, v in NUTRICION.items()
           if k not in ("rationale", "flexibility_rules", "refeed_or_break")}
    texto = _texto(nutrition=nut)
    assert "POR QUÉ ESTE ENFOQUE" not in texto
    assert "TU MARGEN DE MANIOBRA" not in texto


def test_los_porcentajes_del_reparto_suman_cien():
    lineas = _macro_lines({"carbs_g": 210, "protein_g": 165, "fat_g": 68}, 2200)
    pct = [int(x.rstrip("%")) for x in lineas[1].split(" de ")[0].split(" · ")]
    assert sum(pct) == 100


def test_el_ajuste_no_repite_el_signo_con_la_palabra():
    # "Déficit -450 kcal" decía dos veces lo mismo y con un guion feo.
    assert _ajuste_text({"tdee_kcal": 2650, "target_kcal": 2200}, "fat_loss") \
        == "Déficit de 450 kcal (17%)"
    assert _ajuste_text({"tdee_kcal": 2000, "target_kcal": 2000}, "maintenance") \
        == "Mantenimiento · sin ajuste"


def test_el_menu_cerrado_trae_su_lista_de_la_compra():
    """Con menú cerrado el cliente tenía que sumar a mano los ingredientes de
    28 platos para ir al supermercado: la pieza que lo resuelve existía entera
    (con tests) y no la llamaba nadie."""
    def _plato(titulo, *ings):
        return {"title": titulo, "prep": "Cocinar",
                "ingredients": [{"food": f, "grams": g} for f, g in ings],
                "macros": {"kcal": 700, "protein_g": 50, "carbs_g": 70, "fat_g": 20}}

    dias = [{"day": d, "meals": [{"slot": 1, "dish": _plato(
        "Pollo con arroz", ("Pollo", 200), ("Arroz", 80))}]}
        for d in ("lunes", "martes", "miércoles", "jueves", "viernes",
                  "sábado", "domingo")]
    nut = {**NUTRICION, "meal_bank": {"mode": "strict", "days": dias}}
    texto = _texto(nutrition=nut, diet_mode="strict")
    assert "lista de la compra" in texto.lower()
    assert "1400 g" in texto or "1.400" in texto     # 7 × 200 g de pollo


def test_el_educativo_tambien_respeta_las_alergias():
    """Era la ÚNICA sección del documento sin filtro de alérgenos, y encima su
    texto se cachea por split: la píldora escrita para otro cliente acababa
    impresa en el PDF de un alérgico a la leche."""
    educacion = {
        "pills": [
            {"topic": "Proteína", "for_client": "Un yogur griego después de entrenar."},
            {"topic": "Sueño", "for_client": "Duerme 7-8 h para recuperar mejor."},
        ],
        "faq": [{"q": "¿Puedo tomar leche?", "a": "Sí, un vaso al día."},
                {"q": "¿Y el cardio?", "a": "Camina 8.000 pasos."}],
        "biomech_by_pattern": [],
    }
    texto = _texto(education=educacion, food_allergies=["lactosa"])
    assert "Duerme 7-8 h" in texto and "8.000 pasos" in texto
    assert "yogur" not in texto.lower()
    assert "leche" not in texto.lower()


def test_la_fecha_del_plan_es_la_de_generacion():
    """El PDF se construye en cada descarga: sin este dato, el mismo plan salía
    con una fecha distinta cada vez (y el del mes 3 podía parecer posterior al
    del mes 4)."""
    from datetime import date

    data = generate_plan_doc(
        brand=BRAND, client_name="Mario", month_index=3, goal_type="fat_loss",
        diet_mode="flexible", nutrition=NUTRICION, training={}, education={},
        include_training=False, include_nutrition=True,
        generated_on=date(2026, 3, 14),
    )
    doc = Document(io.BytesIO(data))
    cabeceras = " ".join(
        p.text for s in doc.sections
        for hdr in (s.header, s.first_page_header, s.even_page_header)
        for p in hdr.paragraphs
    )
    tablas = " ".join(
        c.text for s in doc.sections
        for hdr in (s.header, s.first_page_header, s.even_page_header)
        for t in hdr.tables for row in t.rows for c in row.cells
    )
    assert "14/03/2026" in (cabeceras + tablas)
