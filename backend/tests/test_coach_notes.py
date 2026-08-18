"""Indicaciones personalizadas por ejercicio (coach_notes): el coach las escribe
en el editor y el cliente las ve en su portal y en el PDF."""
from __future__ import annotations

import uuid

import pytest

from tests.test_portal import _db_available, _generate_plan_content  # infra compartida

pytestmark = pytest.mark.skipif(not _db_available(), reason="Requiere PostgreSQL")


@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient

    from app.main import app
    from app.seeds.run import main as seed_main

    seed_main()
    return TestClient(app)


@pytest.fixture(scope="module")
def auth():
    from app.security import create_access_token

    return {"Authorization": f"Bearer {create_access_token('coach1')}"}


def _client_with_published_plan(client, auth) -> tuple[int, str, list[int]]:
    """Crea cliente + anamnesis + plan publicado. Devuelve (id, token, exercise_ids)."""
    body = client.post(
        "/api/clients", headers=auth,
        json={"full_name": "Notas Tester", "email": f"notas-{uuid.uuid4().hex[:8]}@example.com"},
    ).json()
    cid = body["client"]["id"]
    token = body["links"]["portal_token"]

    anam = {
        "sex": "male", "birth_date": "1994-03-10", "height_cm": 180, "start_weight_kg": 82,
        "goal_type": "fat_loss", "goal_weight_kg": 76, "level": "intermediate",
        "training_days": 4, "session_max_min": 75, "training_place": "gym",
        "equipment": ["barra"], "meals_per_day": 4,
        "meal_schedule": [{"slot": i, "name": n, "time": t} for i, n, t in
                          [(1, "Desayuno", "08:00"), (2, "Comida", "14:00"),
                           (3, "Merienda", "18:00"), (4, "Cena", "21:30")]],
        "food_allergies": [], "food_dislikes": [], "food_likes": ["pollo"],
        "diet_mode": "flexible_7", "consent_accepted": True,
    }
    assert client.post(f"/api/p/{token}/anamnesis", json=anam).status_code == 200

    nutrition, training, education, flags = _generate_plan_content()
    plan = client.post(f"/api/clients/{cid}/plans", headers=auth, json={
        "month_index": 1, "nutrition_json": nutrition, "training_json": training,
        "education_json": education, "guardrail_flags": flags, "generated_by": "test",
    }).json()
    assert client.post(f"/api/plans/{plan['id']}/publish", headers=auth).json()["status"] == "published"

    ex_ids: list[int] = []
    for s in training.get("sessions", []):
        for e in s.get("exercises", []):
            if e["exercise_id"] not in ex_ids:
                ex_ids.append(e["exercise_id"])
    assert ex_ids, "el plan de prueba debe referenciar al menos un ejercicio"
    return cid, token, ex_ids


def test_coach_notes_llegan_al_portal(client, auth):
    cid, token, _ = _client_with_published_plan(client, auth)
    plan = client.get(f"/api/clients/{cid}/plans", headers=auth).json()[0]

    tr = plan["training_json"]
    tr["sessions"][0]["exercises"][0]["coach_notes"] = (
        "Por tu hombro: baja solo hasta 90º y sube el respaldo un punto."
    )
    assert client.patch(f"/api/plans/{plan['id']}", headers=auth,
                        json={"training_json": tr}).status_code == 200

    sesiones = client.get(f"/api/p/{token}/training").json()["sessions"]
    primera = sesiones[0]["exercises"][0]
    assert "hombro" in (primera.get("coach_notes") or "")


def test_coach_notes_en_el_doc_completo(client, auth):
    """El doc con entrenamiento (include_training) se genera sin romperse y la
    indicación aparece etiquetada en la celda de clave técnica. (El PDF que se
    entrega es solo dieta: la rutina —con las indicaciones— vive en el portal.)"""
    import io

    from docx import Document

    from app.services.docs.plan_doc import generate_plan_doc
    from app.services.docs.word_base import DocBrand

    cid, _, ex_ids = _client_with_published_plan(client, auth)
    plan = client.get(f"/api/clients/{cid}/plans", headers=auth).json()[0]
    tr = plan["training_json"]
    tr["sessions"][0]["exercises"][0]["coach_notes"] = "Agarre neutro por la muñeca."

    raw = generate_plan_doc(
        brand=DocBrand(name="Professional", color_primary="#8B1A2B",
                       color_secondary="#4A7BA8", font_family="Inter"),
        client_name="Test", month_index=1, goal_type="fat_loss",
        diet_mode="flexible_7", nutrition=plan["nutrition_json"], training=tr,
        education=plan["education_json"] or {},
        exercise_names={i: f"Ejercicio {i}" for i in ex_ids},
        include_training=True,
    )
    doc = Document(io.BytesIO(raw))
    texto = "\n".join(p.text for p in doc.paragraphs) + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    assert "Indicación para ti: Agarre neutro por la muñeca." in texto
