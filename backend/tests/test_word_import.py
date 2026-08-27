"""IDA Y VUELTA del Word editable: el coach descarga el .docx, lo edita en Word
y lo sube — el sistema detecta los cambios y los aplica por el PATCH de siempre.

Se prueba el ciclo REAL: generar el documento por el endpoint, mutar celdas con
python-docx (como haría Word), subirlo a /import-word y aplicar. Requiere
PostgreSQL (se salta sin él)."""
import io
import uuid

import pytest


def _db_available() -> bool:
    try:
        from sqlalchemy import create_engine, text

        from app.config import settings

        create_engine(settings.database_url).connect().execute(text("SELECT 1"))
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(not _db_available(), reason="Requiere PostgreSQL")


@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient

    from app.main import app
    from app.seeds.run import main as seed_main

    seed_main()
    return TestClient(app)


@pytest.fixture(scope="module")
def auth(client):
    from app.security import create_access_token

    return {"Authorization": f"Bearer {create_access_token('coach1')}"}


def _nutrition() -> dict:
    return {
        "tdee_kcal": 2500, "target_kcal": 2000,
        "macros": {"protein_g": 150, "carbs_g": 200, "fat_g": 60},
        "meals": [
            {"slot": 1, "name": "Desayuno", "time": "08:00",
             "target": {"kcal": 1000, "protein_g": 75, "carbs_g": 100, "fat_g": 30}},
            {"slot": 2, "name": "Cena", "time": "21:00",
             "target": {"kcal": 1000, "protein_g": 75, "carbs_g": 100, "fat_g": 30}},
        ],
        "supplements": [{"name": "Creatina", "dose": "5 g", "timing": "post-entreno",
                         "evidence_note": "sólida"}],
        "rationale": "Déficit moderado para empezar.",
        "flexibility_rules": ["Una comida social a la semana"],
        "refeed_or_break": None,
        "meal_bank": {"mode": "flexible_7", "slots": []},
    }


def _training(ex1: int, ex2: int) -> dict:
    def ex(eid, sets=3):
        return {"exercise_id": eid, "sets": sets, "rep_range": "8-10", "rir": "2",
                "tempo": None, "rest_sec": 90, "start_weight_hint_kg": None,
                "progression_rule": "Sube 2,5 kg al completar el rango",
                "technique_cue": "Controla la bajada", "biomech_cue": "",
                "coach_notes": None}

    return {
        "split_name": "Torso/Pierna", "split_rationale": "2 días",
        "weekly_progression": [
            {"week": 1, "intent": "Base", "load_pct": 70, "rir_target": "3",
             "volume_note": "arranque"},
            {"week": 2, "intent": "Progresión", "load_pct": 75, "rir_target": "2",
             "volume_note": ""},
            {"week": 3, "intent": "Pico", "load_pct": 80, "rir_target": "1-2",
             "volume_note": ""},
            {"week": 4, "intent": "Deload", "load_pct": 60, "rir_target": "4",
             "volume_note": "descarga"},
        ],
        "sessions": [
            {"day": "Lunes", "name": "Torso A", "warmup": "5 min bici",
             "exercises": [ex(ex1), ex(ex2, sets=4)], "cooldown": "estiramientos"},
        ],
        "cardio": {"daily_steps": 8000,
                   "sessions": [{"type": "liss", "minutes": 30,
                                 "times_per_week": 2, "notes": None}]},
        "deload_instructions": "Semana 4: volumen a la mitad.",
    }


@pytest.fixture(scope="module")
def setup(client, auth):
    exs = client.get("/api/exercises", headers=auth).json()
    assert len(exs) >= 3
    ex1, ex2, ex3 = exs[0], exs[1], exs[2]
    r = client.post("/api/clients", headers=auth, json={
        "full_name": "Word Import", "email": f"word-{uuid.uuid4().hex[:8]}@example.com",
    })
    cid = r.json()["client"]["id"]
    r = client.post(f"/api/clients/{cid}/plans", headers=auth, json={
        "month_index": 1, "nutrition_json": _nutrition(),
        "training_json": _training(ex1["id"], ex2["id"]), "education_json": {},
    })
    assert r.status_code == 201, r.text
    return {"cid": cid, "plan_id": r.json()["id"], "ex": (ex1, ex2, ex3)}


def _download_docx(client, auth, plan_id: int) -> bytes:
    r = client.get(f"/api/plans/{plan_id}/document?format=docx", headers=auth)
    assert r.status_code == 200
    return r.content


def _edit_docx(content: bytes, edits) -> bytes:
    """Aplica `edits(doc)` sobre el documento y lo devuelve como bytes."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    edits(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _find_table(doc, first_header: str):
    for t in doc.tables:
        try:
            if t.rows[0].cells[0].text.strip().lower().startswith(first_header):
                return t
        except Exception:
            continue
    return None


def test_import_word_detecta_y_aplica_cambios(client, auth, setup):
    plan_id = setup["plan_id"]
    ex1, ex2, ex3 = setup["ex"]
    original = _download_docx(client, auth, plan_id)

    def edits(doc):
        # kcal y macros del resumen energético
        t = _find_table(doc, "calorías")
        assert t is not None
        t.rows[1].cells[0].text = "≈ 2200 kcal"
        t.rows[1].cells[1].text = "CH 220 g · P 160 g · G 62 g"
        # hora de la cena
        tt = _find_table(doc, "hora")
        assert tt is not None
        tt.rows[2].cells[0].text = "20:30"
        # sesión: series del primer ejercicio, descanso del segundo y CAMBIO de
        # ejercicio (el segundo pasa a ser ex3, otro de la biblioteca)
        ts = _find_table(doc, "ejercicio")
        assert ts is not None
        ts.rows[1].cells[1].text = "5×6-8"
        ts.rows[2].cells[0].text = ex3["canonical_name"]
        ts.rows[2].cells[3].text = "120s"
        # progresión: carga de la semana 2
        tp = _find_table(doc, "semana")
        assert tp is not None
        tp.rows[2].cells[2].text = "78%"

    editado = _edit_docx(original, edits)
    r = client.post(
        f"/api/plans/{plan_id}/import-word", headers=auth,
        files={"file": ("plan.docx", editado,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["has_changes"]
    frases = " | ".join(data["changes"])
    assert "2000" in frases and "2200" in frases          # kcal detectadas
    assert "20:30" in frases                              # hora de la cena
    assert "5×6-8" in frases or "5x6-8" in frases         # series
    assert ex3["canonical_name"] in frases                # cambio de ejercicio
    assert "78" in frases                                 # carga semana 2

    # Aplicar por el MISMO camino que el editor (PATCH con base_rev)
    r2 = client.patch(f"/api/plans/{plan_id}", headers=auth, json={
        "nutrition_json": data["nutrition_json"],
        "training_json": data["training_json"],
        "base_rev": data["base_rev"],
    })
    assert r2.status_code == 200, r2.text
    guardado = r2.json()
    # El PATCH reconcilia (kcal ≡ 4/4/9 de los macros redondeados): tolerancia
    # de unas pocas kcal alrededor de lo escrito en el Word.
    assert abs(guardado["nutrition_json"]["target_kcal"] - 2200) <= 5
    assert guardado["nutrition_json"]["macros"]["protein_g"] == 160
    ses = guardado["training_json"]["sessions"][0]
    assert ses["exercises"][0]["sets"] == 5
    assert ses["exercises"][0]["rep_range"] == "6-8"
    assert ses["exercises"][1]["exercise_id"] == ex3["id"]
    assert ses["exercises"][1]["rest_sec"] == 120
    assert guardado["training_json"]["weekly_progression"][1]["load_pct"] == 78
    assert guardado["nutrition_json"]["meals"][1]["time"] == "20:30"


def test_import_word_rechaza_archivos_invalidos(client, auth, setup):
    plan_id = setup["plan_id"]
    # No es un zip/docx
    r = client.post(f"/api/plans/{plan_id}/import-word", headers=auth,
                    files={"file": ("plan.docx", b"%PDF-1.4 no soy word", "application/pdf")})
    assert r.status_code == 422
    # Es un docx pero NO nuestro plan (sin ninguna tabla reconocible)
    from docx import Document as _D

    buf = io.BytesIO()
    d = _D()
    d.add_paragraph("Documento cualquiera")
    d.save(buf)
    r2 = client.post(f"/api/plans/{plan_id}/import-word", headers=auth,
                     files={"file": ("otro.docx", buf.getvalue(),
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r2.status_code == 422
    assert "No reconozco" in r2.json()["detail"]


def test_import_word_sin_cambios_no_inventa_nada(client, auth, setup):
    plan_id = setup["plan_id"]
    original = _download_docx(client, auth, plan_id)
    r = client.post(
        f"/api/plans/{plan_id}/import-word", headers=auth,
        files={"file": ("plan.docx", original,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    assert r.json()["changes"] == []


def _find_box(doc, prefix: str):
    """Caja 1×1 cuyo texto empieza por `prefix` (normalizado)."""
    from app.services.word_import import _norm

    for t in doc.tables:
        if len(t.rows) == 1 and len(t.rows[0].cells) == 1:
            if _norm(t.rows[0].cells[0].text).startswith(_norm(prefix)):
                return t.rows[0].cells[0]
    return None


def _find_bar(doc, text: str):
    from app.services.word_import import _es_barra, _iter_blocks, _norm

    for k, b in _iter_blocks(doc):
        if k == "p" and _es_barra(b) and _norm(b.text) == _norm(text):
            return b
    return None


def _upload(client, auth, plan_id: int, contenido: bytes) -> dict:
    r = client.post(
        f"/api/plans/{plan_id}/import-word", headers=auth,
        files={"file": ("plan.docx", contenido,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    return r.json()


def test_reparto_reordenado_estrategia_descanso_min_y_rename(client, auth, setup):
    """Extensiones de la ronda de ingesta: macros en OTRO orden y con etiquetas
    naturales, estrategia de la toma, «2 min» de descanso, «4X12» con X
    mayúscula, enfoque en minúscula y rename de la sesión desde su barra."""
    plan_id = setup["plan_id"]
    original = _download_docx(client, auth, plan_id)

    def edits(doc):
        t = _find_table(doc, "calorías")
        t.rows[1].cells[1].text = "Prot 165 g · G 58 g · CH 210 g"  # reordenado
        tt = _find_table(doc, "hora")
        tt.rows[1].cells[2].text = "Prioriza la proteína al levantarte."
        ts = _find_table(doc, "ejercicio")
        ts.rows[1].cells[3].text = "2 min"
        ts.rows[1].cells[1].text = "4X12"
        tp = _find_table(doc, "semana")
        tp.rows[2].cells[1].text = "pico"      # minúscula → mapea a "Pico"
        bar = _find_bar(doc, "Lunes · Torso A")
        assert bar is not None
        bar.text = "Lunes · Empuje fuerte"

    data = _upload(client, auth, plan_id, _edit_docx(original, edits))
    nut, tr = data["nutrition_json"], data["training_json"]
    assert nut["macros"] == {"protein_g": 165, "carbs_g": 210, "fat_g": 58}
    assert nut["meals"][0]["strategy"] == "Prioriza la proteína al levantarte."
    ses = tr["sessions"][0]
    assert ses["name"] == "Empuje fuerte"
    assert ses["exercises"][0]["rest_sec"] == 120
    assert ses["exercises"][0]["sets"] == 4
    assert ses["exercises"][0]["rep_range"] == "12"
    assert tr["weekly_progression"][1]["intent"] == "Pico"


def test_margen_rationale_cardio_calentamiento_y_deload(client, auth, setup):
    """Las cajas que antes se perdían EN SILENCIO ahora se importan: margen de
    maniobra, «Por qué este enfoque», sesiones de cardio (y pasos con la
    etiqueta reescrita), calentamiento y vaciado del deload."""
    plan_id = setup["plan_id"]
    original = _download_docx(client, auth, plan_id)

    def edits(doc):
        b = _find_box(doc, "Déficit moderado")
        assert b is not None
        b.text = "Enfoque revisado por el coach en Word."
        m = _find_box(doc, "• Una comida social")
        assert m is not None
        m.text = "• Una comida social a la semana\n• Fruta libre entre horas"
        c = _find_box(doc, "Pasos diarios objetivo")
        assert c is not None
        c.text = "9.500 pasos cada día\nLISS: 40 min × 3/sem — en ayunas suave"
        w = _find_box(doc, "Calentamiento")
        assert w is not None
        w.text = "Calentamiento: 8 min de movilidad de hombro y cadera."
        d = _find_box(doc, "Semana 4: volumen a la mitad")
        assert d is not None
        d.text = ""

    data = _upload(client, auth, plan_id, _edit_docx(original, edits))
    nut, tr = data["nutrition_json"], data["training_json"]
    assert nut["rationale"] == "Enfoque revisado por el coach en Word."
    assert nut["flexibility_rules"] == ["Una comida social a la semana",
                                       "Fruta libre entre horas"]
    assert tr["cardio"]["daily_steps"] == 9500
    cs = tr["cardio"]["sessions"][0]
    assert (cs["minutes"], cs["times_per_week"]) == (40, 3)
    assert cs["notes"] == "en ayunas suave"
    assert "movilidad de hombro" in tr["sessions"][0]["warmup"]
    assert tr["deload_instructions"] == ""     # vaciar la caja QUITA el deload


def test_cabecera_modificada_avisa_no_calla(client, auth, setup):
    """Renombrar una columna de una tabla nuestra ya no la hace invisible: se
    avisa de que esa tabla no se pudo importar."""
    plan_id = setup["plan_id"]
    original = _download_docx(client, auth, plan_id)

    def edits(doc):
        ts = _find_table(doc, "ejercicio")
        ts.rows[0].cells[2].text = "RIR objetivo"
        ts.rows[1].cells[1].text = "7×7"       # cambio que se PERDERÍA

    data = _upload(client, auth, plan_id, _edit_docx(original, edits))
    assert any("cabecera modificada" in w for w in data["warnings"]), data["warnings"]
    # y el cambio de esa tabla efectivamente NO se aplicó (ni en silencio)
    assert data["training_json"]["sessions"][0]["exercises"][0]["sets"] != 7


def test_kcal_ilegibles_no_destruyen_el_plan(client, auth, setup):
    """Una cifra absurda (coma anglosajona mal leída) NO reescala el plan a la
    nada: aviso y el resumen energético se queda como estaba. Y «2,150» se
    entiende como 2150 (separador de miles), no como 2,15 kcal."""
    from app.services.word_import import _num

    assert _num("2,150") == 2150
    assert _num("120,50") == 120.5
    plan_id = setup["plan_id"]
    original = _download_docx(client, auth, plan_id)

    def edits(doc):
        t = _find_table(doc, "calorías")
        t.rows[1].cells[0].text = "180 kcal"

    planes = client.get(f"/api/clients/{setup['cid']}/plans", headers=auth).json()
    actual = next(p for p in planes if p["id"] == plan_id)["nutrition_json"]["target_kcal"]
    data = _upload(client, auth, plan_id, _edit_docx(original, edits))
    assert any("no son creíbles" in w for w in data["warnings"]), data["warnings"]
    assert data["nutrition_json"]["target_kcal"] == actual  # intacto


def test_parsers_unitarios():
    from app.services.word_import import _parse_reparto, _parse_rest

    assert _parse_rest("2 min") == 120
    assert _parse_rest("1,5 min") == 90
    assert _parse_rest("1 min 30") == 90
    assert _parse_rest("90s") == 90
    assert _parse_rest("120") == 120
    r = _parse_reparto("P 170 g · G 60 g · CH 230 g\n55% · 25% · 20% de tus calorías")
    assert r == {"protein_g": 170, "fat_g": 60, "carbs_g": 230}
    assert _parse_reparto("hidratos 200 g y proteína 150 g") == {
        "carbs_g": 200, "protein_g": 150}


@pytest.fixture(scope="module")
def scaffold_flexible(client, auth):
    """Cliente completo + BASE A MANO (banco determinista real): para probar la
    ida y vuelta de las RECETAS con macros recalculados desde `foods`."""
    r = client.post("/api/clients", headers=auth, json={
        "full_name": "Word Recetas", "email": f"word-rec-{uuid.uuid4().hex[:8]}@example.com",
    })
    cid = r.json()["client"]["id"]
    client.patch(f"/api/clients/{cid}", headers=auth, json={
        "sex": "male", "birth_date": "1990-01-01", "height_cm": 178,
        "start_weight_kg": 82, "goal_type": "fat_loss", "level": "intermediate",
        "training_days": 3, "session_max_min": 60, "training_place": "gym",
        "diet_mode": "flexible_7"})
    r = client.post(f"/api/clients/{cid}/scaffold-plan", headers=auth)
    assert r.status_code in (200, 201), r.text
    return {"cid": cid, "plan": r.json()}


def test_scaffold_usa_dias_reales_de_la_semana(scaffold_flexible):
    """Regresión (auditoría 27-08): la base a mano decía "Día 1" y el portal
    nunca detectaba la sesión de HOY (compara con lunes…domingo)."""
    dias = [s["day"] for s in scaffold_flexible["plan"]["training"]["sessions"]]
    assert dias[0] == "Lunes", dias
    validos = {"Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"}
    assert all(d in validos for d in dias), dias


def test_recetas_banco_flexible_ida_y_vuelta(client, auth, scaffold_flexible):
    """Editar una receta del banco en el Word («Opción 1…») la aplica al plan
    con los MACROS RECALCULADOS por el backend (Atwater 4/4/9 de `foods`)."""
    import re as _re

    plan = scaffold_flexible["plan"]
    plan_id = plan["id"]
    meal1 = plan["nutrition"]["meals"][0]
    original = _download_docx(client, auth, plan_id)

    gramos_nuevos = {}

    def edits(doc):
        from app.services.word_import import _norm
        b = _find_box(doc, "Opción 1")
        assert b is not None
        p0 = b.paragraphs[0]
        m = _re.match(r"^Opción 1\.\s+(.+?) — (.+?)\.?$", p0.text)
        assert m, p0.text
        ings = m.group(2)
        mg = _re.search(r"(\d+)\s*g", ings)
        gramos_nuevos["g"] = int(mg.group(1)) + 25
        ings2 = ings[:mg.start()] + f"{gramos_nuevos['g']} g" + ings[mg.end():]
        for run in p0.runs:
            run.text = ""
        p0.add_run(f"Opción 1. Receta editada en Word — {ings2}.")

    data = _upload(client, auth, plan_id, _edit_docx(original, edits))
    assert any("opción 1" in c.lower() for c in data["changes"]), data["changes"]
    slots = {s["slot"]: s for s in data["nutrition_json"]["meal_bank"]["slots"]}
    op = slots[meal1["slot"]]["options"][0]
    assert op["title"] == "Receta editada en Word"
    assert round(op["ingredients"][0]["grams"]) == gramos_nuevos["g"]
    mac = op["macros"]
    assert round(mac["kcal"]) == round(mac["protein_g"] * 4 + mac["carbs_g"] * 4
                                       + mac["fat_g"] * 9)


def test_menu_strict_ida_y_vuelta(client, auth):
    """El MENÚ CERRADO por días también va y vuelve, incluida la pauta de la
    comida libre semanal (que el cliente pidió en su anamnesis)."""
    import re as _re

    r = client.post("/api/clients", headers=auth, json={
        "full_name": "Word Strict", "email": f"word-str-{uuid.uuid4().hex[:8]}@example.com",
    })
    cid = r.json()["client"]["id"]
    client.patch(f"/api/clients/{cid}", headers=auth, json={
        "sex": "female", "birth_date": "1992-05-05", "height_cm": 165,
        "start_weight_kg": 60, "goal_type": "recomp", "level": "beginner",
        "training_days": 3, "session_max_min": 60, "training_place": "gym",
        "diet_mode": "strict", "strict_free_meal_enabled": True})
    plan = client.post(f"/api/clients/{cid}/scaffold-plan", headers=auth).json()
    bank = plan["nutrition"]["meal_bank"]
    assert bank and bank["mode"] == "strict"
    assert bank.get("free_meal_guidelines")  # la pidió → pauta determinista
    original = _download_docx(client, auth, plan["id"])

    def edits(doc):
        b = _find_box(doc, "Tienes UNA comida libre")
        assert b is not None
        b.text = "Comida libre: el sábado, y de vuelta al plan en la cena."
        # primer plato del MARTES: título nuevo
        from app.services.word_import import _es_barra, _iter_blocks, _norm
        barra = ""
        for k, blk in _iter_blocks(doc):
            if k == "p" and _es_barra(blk):
                barra = _norm(blk.text)
            elif barra == "martes" and len(blk.rows) == 1 and len(blk.rows[0].cells) == 1:
                p0 = blk.rows[0].cells[0].paragraphs[0]
                m = _re.match(r"^(.+?)\.\s+(.+?) — (.+?)\.?$", p0.text)
                assert m, p0.text
                for run in p0.runs:
                    run.text = ""
                p0.add_run(f"{m.group(1)}. Plato editado en Word — {m.group(3)}.")
                return

    data = _upload(client, auth, plan["id"], _edit_docx(original, edits))
    nut = data["nutrition_json"]
    assert "el sábado" in nut["meal_bank"]["free_meal_guidelines"]
    martes = next(d for d in nut["meal_bank"]["days"] if d["day"] == "martes")
    assert martes["meals"][0]["dish"]["title"] == "Plato editado en Word"


def test_la_linea_guia_no_suplanta_a_la_barra_de_seccion():
    """Regresión: al imprimir una línea guía bajo cada barra, las cajas de
    Cardio y de la semana de descarga dejaban de importarse EN SILENCIO — la
    frase de la guía pasaba a ser la 'última barra' y ya no casaba."""
    from docx import Document

    from app.services.docs.word_base import DocBrand
    from app.services.docs.plan_doc import generate_plan_doc
    from app.services.word_import import _es_barra, _iter_blocks

    brand = DocBrand(name="DQR", color_primary="#8B1A2B", color_secondary="#4A7BA8",
                     font_family="Calibri")
    data = generate_plan_doc(
        brand=brand, client_name="Mario", month_index=1, goal_type="fat_loss",
        diet_mode="flexible",
        nutrition={"tdee_kcal": 2500, "target_kcal": 2200,
                   "macros": {"carbs_g": 200, "protein_g": 170, "fat_g": 70},
                   "meals": [{"slot": 1, "name": "Desayuno", "time": "08:00"}],
                   "meal_bank": {"mode": "flexible", "slots": []}},
        training={"split_name": "Full", "sessions": [],
                  "cardio": {"daily_steps": 9000, "sessions": []},
                  "deload_instructions": "Semana 4: mitad de series."},
        education={}, include_training=True, include_nutrition=True,
    )
    doc = Document(io.BytesIO(data))
    barras = [b.text.strip() for k, b in _iter_blocks(doc) if k == "p" and _es_barra(b)]
    assert "CARDIO Y NEAT" in barras, barras
    assert "SEMANA DE DESCARGA (DELOAD)" in barras, barras
    # Y la línea guía NO se cuela como barra.
    assert not any("pasos diarios pesan" in b.lower() for b in barras)
