"""Lector UNIVERSAL de documentos: cualquier fichero que se suba a DQR
(PDF, Word, fotos, Excel, texto) se lee ENTERO, con la estructura que tenga, y
alimenta lo que toque: la ficha (anamnesis), los adjuntos (analítica, informes)
y la planificación (plan ajeno → borrador con cifras del backend).

Todo con IA FALSA (respuestas guionizadas): aquí se prueba la fontanería —
normalización, ingesta, fusión en la ficha, contratos— no el modelo.
"""
import io
import json
import os
import uuid
from datetime import date

import pytest


def _db_available() -> bool:
    try:
        from sqlalchemy import create_engine, text

        from app.config import settings

        create_engine(settings.database_url).connect().execute(text("SELECT 1"))
        return True
    except Exception:
        return False


db_required = pytest.mark.skipif(not _db_available(), reason="Requiere PostgreSQL")

_DIAS_SEMANA = ("lunes", "martes", "miercoles", "jueves", "viernes", "sabado", "domingo")


# ------------------------------------------------------------ utilidades ----

def _pdf(paginas: int = 2, texto: str = "Peso 82 kg") -> bytes:
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    for i in range(paginas):
        c.drawString(72, 720, f"{texto} · página {i + 1}")
        c.showPage()
    c.save()
    return buf.getvalue()


def _png(w: int = 800, h: int = 1000, modo: str = "RGB") -> bytes:
    from PIL import Image

    img = Image.new(modo, (w, h), (240, 240, 240) if modo == "RGB" else (240, 240, 240, 0))
    b = io.BytesIO()
    img.save(b, "PNG")
    return b.getvalue()


def _docx(parrafos: list[str], tabla: list[list[str]] | None = None) -> bytes:
    from docx import Document

    d = Document()
    for p in parrafos:
        d.add_paragraph(p)
    if tabla:
        t = d.add_table(rows=len(tabla), cols=len(tabla[0]))
        for i, fila in enumerate(tabla):
            for j, celda in enumerate(fila):
                t.cell(i, j).text = celda
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _xlsx(filas: list[list]) -> bytes:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dieta"
    for f in filas:
        ws.append(f)
    b = io.BytesIO()
    wb.save(b)
    return b.getvalue()


class ScriptedDocs:
    """AIClient falso para el lector: `_raw_call_with_blocks` devuelve las
    respuestas guionizadas en orden y guarda lo que se le pidió."""

    def __init__(self, respuestas: list):
        self.respuestas = [r if isinstance(r, str) else json.dumps(r, ensure_ascii=False)
                           for r in respuestas]
        self.calls: list[dict] = []

    def instalar(self, monkeypatch):
        from app.services.ai.client import AIClient

        yo = self

        def _falso(self_, *, model, system, user, bloques, temperature=None, max_tokens=None):
            yo.calls.append({"model": model, "system": system, "user": user,
                             "bloques": bloques, "temperature": temperature,
                             "max_tokens": max_tokens})
            if not yo.respuestas:
                raise RuntimeError("guion agotado")
            return yo.respuestas.pop(0)

        monkeypatch.setattr(AIClient, "_raw_call_with_blocks", _falso)
        # Los endpoints construyen `AIClient()` sin clave explícita: sin una
        # en settings el constructor corta con «Falta ANTHROPIC_API_KEY».
        from app.config import settings

        monkeypatch.setattr(settings, "anthropic_api_key", "test-key")
        return self


# ========================================================== normalización ====

def test_pdf_se_lee_nativo_y_se_cuentan_paginas():
    from app.services.document_reader import normalizar

    d = normalizar(_pdf(3), "Cuestionario Otro Coach.PDF")
    assert d.origen == "pdf" and d.paginas == 3 and d.extension == "pdf"
    assert [b["type"] for b in d.bloques] == ["document"]
    assert d.bloques[0]["source"]["media_type"] == "application/pdf"
    assert "3 páginas" in d.descripcion
    assert d.nombre.endswith(".pdf") and d.contenido.startswith(b"%PDF-")


def test_pdf_largo_se_trocea_en_bloques_de_100_paginas():
    from app.services.document_reader import normalizar

    d = normalizar(_pdf(205, "x"), "historia.pdf")
    assert d.paginas == 205
    assert len(d.bloques) == 3                       # 100 + 100 + 5
    assert any("3 partes" in a for a in d.avisos)


def test_una_foto_se_endereza_reescala_y_recomprime_a_jpeg():
    from PIL import Image

    from app.services.document_reader import IMAGE_MAX_SIDE, normalizar

    d = normalizar(_png(3000, 4000, "RGBA"), "IMG_0001.PNG")
    assert d.origen == "imagen" and d.extension == "jpg" and d.paginas == 1
    assert d.bloques[0]["type"] == "image"
    assert d.bloques[0]["source"]["media_type"] == "image/jpeg"
    img = Image.open(io.BytesIO(d.contenido))
    assert max(img.size) == IMAGE_MAX_SIDE and img.mode == "RGB"


def test_varias_fotos_son_un_documento_guardado_como_pdf_de_n_paginas():
    from app.services.document_reader import _paginas_pdf, normalizar_varios

    d = normalizar_varios([(_png(600, 800), "p1.png"), (_png(600, 800), "p2.png"),
                           (_png(600, 800), "p3.png")], "anamnesis")
    assert d.origen == "imagen" and d.paginas == 3 and d.extension == "pdf"
    assert len(d.bloques) == 3 and all(b["type"] == "image" for b in d.bloques)
    assert _paginas_pdf(d.contenido) == 3          # lo que se GUARDA
    assert "3 fotos" in d.descripcion


def test_pdf_mas_fotos_se_unen_en_un_solo_pdf():
    from app.services.document_reader import _paginas_pdf, normalizar_varios

    d = normalizar_varios([(_pdf(2), "a.pdf"), (_png(), "b.png")], "mixto")
    assert d.origen == "pdf" and _paginas_pdf(d.contenido) == 3
    assert any("unidos" in a for a in d.avisos)


def test_word_sin_libreoffice_se_lee_por_texto_con_tablas(monkeypatch):
    from app.services.docs import pdf_convert
    from app.services.document_reader import normalizar

    monkeypatch.setattr(pdf_convert, "office_bytes_to_pdf",
                        lambda raw, ext, timeout=120: (_ for _ in ()).throw(RuntimeError("sin soffice")))
    d = normalizar(_docx(["ANAMNESIS de otro centro"], [["Peso", "82 kg"], ["Alergias", "lactosa"]]),
                   "anamnesis otro centro.docx")
    assert d.origen == "office" and d.conversion == "python-docx" and d.extension == "docx"
    assert d.bloques[0]["type"] == "text"
    assert "Peso | 82 kg" in d.texto and "Alergias | lactosa" in d.texto
    assert any("sin maquetación" in a for a in d.avisos)


def test_word_con_libreoffice_llega_como_pdf_nativo(monkeypatch):
    from app.services.docs import pdf_convert
    from app.services.document_reader import normalizar

    monkeypatch.setattr(pdf_convert, "office_bytes_to_pdf", lambda raw, ext, timeout=120: _pdf(4))
    d = normalizar(_docx(["hola"]), "plan.docx")
    assert d.origen == "pdf" and d.conversion == "libreoffice" and d.paginas == 4
    assert "convertido desde Word" in d.descripcion


def test_excel_se_vuelca_como_tabla_de_texto():
    from app.services.document_reader import normalizar

    d = normalizar(_xlsx([["Comida", "Alimento", "Gramos"], ["Desayuno", "Avena", 60],
                          [None, None, None], ["Comida", "Arroz", 120]]), "dieta.xlsx")
    assert d.origen == "hoja" and d.extension == "xlsx"
    assert "Desayuno\tAvena\t60" in d.texto and "Comida\tArroz\t120" in d.texto
    assert "\n\n" not in d.texto.split("---")[1].strip()   # la fila vacía no está


def test_texto_csv_y_markdown_se_leen_tal_cual():
    from app.services.document_reader import normalizar

    d = normalizar("Peso: 80 kg\nAlergias: frutos secos\n".encode("utf-8"), "notas.txt")
    assert d.origen == "texto" and d.extension == "txt"
    assert "Peso: 80 kg" in d.bloques[0]["text"]
    d2 = normalizar("a;b\n1;2\n".encode("latin-1"), "x.csv")
    assert d2.extension == "csv" and "tabla CSV" in d2.bloques[0]["text"]


@pytest.mark.parametrize("raw, nombre, trozo", [
    (b"\x00\x01\x02" * 40, "raro.bin", "No reconozco el formato"),
    (b"\x00\x00\x00\x18ftypheic" + b"\x00" * 40, "IMG_1.HEIC", "HEIC"),
    (b"", "vacio.pdf", "vacío"),
    (b"PK\x03\x04" + b"\x00" * 60, "raro.zip", "No reconozco"),
])
def test_lo_ilegible_se_rechaza_con_mensaje_para_humanos(raw, nombre, trozo):
    from app.services.document_reader import DocumentoIlegible, normalizar

    with pytest.raises(DocumentoIlegible) as exc:
        normalizar(raw, nombre)
    assert trozo in str(exc.value)


def test_la_magia_manda_sobre_el_nombre():
    """Un «.pdf» que en realidad es una foto se lee como foto (y se guarda
    como JPG), no se manda como PDF roto a la IA."""
    from app.services.document_reader import normalizar

    d = normalizar(_png(), "escaneo.pdf")
    assert d.origen == "imagen" and d.extension == "jpg"


def test_bloques_con_cache_marca_solo_el_ultimo():
    from app.services.document_reader import bloques_con_cache, normalizar_varios

    d = normalizar_varios([(_png(), "a.png"), (_png(), "b.png")])
    bs = bloques_con_cache(d)
    assert "cache_control" not in bs[0] and bs[1]["cache_control"] == {"type": "ephemeral"}
    assert "cache_control" not in d.bloques[1]    # no muta el documento


# ============================================== AIClient.read_document_json ===

class _Esquema(__import__("pydantic").BaseModel):
    peso: float
    nombre: str


def test_read_document_json_reintenta_con_el_error_inyectado(monkeypatch):
    from app.services.ai.client import AIClient
    from app.services.document_reader import normalizar

    guion = ScriptedDocs(['{"peso": "no sé"}', '{"peso": 82, "nombre": "Ana"}']).instalar(monkeypatch)
    out = AIClient(api_key="test").read_document_json(
        model="m", system="s", user="u", documento=normalizar(_pdf(1), "a.pdf"), schema=_Esquema)
    assert out.peso == 82 and out.nombre == "Ana"
    assert len(guion.calls) == 2
    assert "CORRECCIÓN REQUERIDA" in guion.calls[1]["user"]
    # el documento viaja como bloque nativo, cacheado, delante del texto
    assert guion.calls[0]["bloques"][0]["type"] == "document"
    assert guion.calls[0]["bloques"][-1]["cache_control"] == {"type": "ephemeral"}


def test_read_document_json_falla_claro_tras_dos_intentos(monkeypatch):
    from app.services.ai.client import AIClient, AIGenerationError
    from app.services.document_reader import normalizar

    ScriptedDocs(["basura", "más basura"]).instalar(monkeypatch)
    with pytest.raises(AIGenerationError):
        AIClient(api_key="test").read_document_json(
            model="m", system="s", user="u", documento=normalizar(b"hola", "n.txt"), schema=_Esquema)


# ============================================= anamnesis: cualquier forma ====

_EXTRACCION_WHATSAPP = {
    "sex": "mujer", "birth_date": "1992-06-03", "height_cm": 1.65 * 100,
    "start_weight_kg": 68.4, "goal_type": "definición", "level": "principiante",
    "training_days": 3, "training_place": "casa", "equipment": ["gomas", "mancuernas"],
    "diet_mode": "flexible", "diet_pattern": "vegetariana", "meals_per_day": 4,
    "meal_schedule": [{"name": "Desayuno", "time": "08:00"}, {"name": "Comida"}],
    "food_allergies": ["frutos secos"], "food_dislikes": ["brócoli"],
    "medical_notes": "- Clínica: hipotiroidismo\n- Analítica (2026-05-12): TSH 4.9 mUI/L (ALTO; ref 0.4-4.0)",
    "medication_notes": "- Eutirox — 50 mcg — diario",
    "injuries_notes": None, "lifestyle_notes": "- Motivo: verse mejor\n- Trabajo: oficina",
    "deep_analysis": "- Empezar suave\n- Vigilar TSH",
    "document_kind": "notas",
    "source_inventory": ["- Mensajes de WhatsApp con datos personales", "- Foto de analítica"],
    "unmapped_info": ["Se casa en octubre", "Tiene perro y pasea 40 min al día"],
    "confidence": {"sex": 1, "start_weight_kg": 0.7, "birth_date": 1},
}
_VERIFICACION_OK = {"sex": "female", "birth_date": "1992-06-03", "height_cm": 165,
                    "start_weight_kg": 68.4, "goal_type": "fat_loss",
                    "food_allergies": ["frutos secos"], "medication_notes": "- Eutirox 50",
                    "medical_notes": "- hipotiroidismo", "injuries_notes": None, "omissions": []}


def test_extraccion_desde_notas_de_whatsapp_mapea_por_significado(monkeypatch):
    """Ni casillas ni secciones del PDF oficial: un .txt con mensajes. Los
    enums se normalizan («mujer», «definición», «casa», «vegetariana»), las
    tomas se autocompletan y lo sin casilla no se pierde."""
    from app.services.ai.client import AIClient
    from app.services.ai.extraction import extract_anamnesis_from_document
    from app.services.document_reader import normalizar

    guion = ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    doc = normalizar("Hola, peso 68,4, mido 1,65…".encode(), "whatsapp.txt")
    lectura = extract_anamnesis_from_document(doc, AIClient(api_key="test"))
    e = lectura.extraction
    assert e.sex == "female" and e.goal_type == "fat_loss" and e.level == "beginner"
    assert e.training_place == "home" and e.diet_pattern == "vegetariano" and e.diet_mode == "flexible_7"
    assert [m.slot for m in e.meal_schedule] == [1, 2] and e.meal_schedule[1].name == "Comida"
    assert e.unmapped_info == ["Se casa en octubre", "Tiene perro y pasea 40 min al día"]
    assert e.confidence["start_weight_kg"] == 0.7
    # segundo pase: mismo documento (cacheado), otro system, campos coinciden
    assert len(guion.calls) == 2
    assert guion.calls[1]["bloques"][0]["text"] == guion.calls[0]["bloques"][0]["text"]
    assert "REVISOR" in guion.calls[1]["system"]
    assert lectura.verification["needs_review"] is False
    assert lectura.verification["discrepancies"] == []
    # el prompt de usuario describe el documento REAL, no «el PDF oficial»
    assert "whatsapp.txt" in guion.calls[0]["user"] and "texto" in guion.calls[0]["user"]
    assert "PDF oficial" not in guion.calls[0]["user"]


def test_el_segundo_pase_marca_discrepancias_criticas_sin_resolverlas_solo(monkeypatch):
    from app.services.ai.client import AIClient
    from app.services.ai.extraction import extract_anamnesis_from_document
    from app.services.document_reader import normalizar

    ver = {**_VERIFICACION_OK, "start_weight_kg": 86.4, "food_allergies": ["frutos secos", "marisco"],
           "injuries_notes": "- Hombro derecho: tendinitis", "omissions": ["Toma omeprazol"]}
    ScriptedDocs([_EXTRACCION_WHATSAPP, ver]).instalar(monkeypatch)
    lectura = extract_anamnesis_from_document(normalizar(b"x", "n.txt"), AIClient(api_key="test"))
    v = lectura.verification
    assert v["needs_review"] is True
    txt = " ".join(v["discrepancies"])
    assert "peso actual" in txt and "86.4" in txt
    assert "alergias" in txt and "marisco" in txt
    assert "lesiones" in txt
    assert v["omissions"] == ["Toma omeprazol"]
    # La extracción NO se toca: el coach decide.
    assert lectura.extraction.start_weight_kg == 68.4
    assert lectura.extraction.food_allergies == ["frutos secos"]


def test_una_duda_por_confianza_baja_dice_en_que_campo(monkeypatch):
    """Sin desajustes pero con la relectura sin encontrar un dato crítico, la
    duda existe (§5) y tiene que decir EN QUÉ campo: antes salía «la relectura
    no coincide en 0 datos» y la ficha no pintaba nada."""
    from app.services.ai.extraction import comparar_pases, resumen_de_dudas

    r = comparar_pases({"sex": "male", "height_cm": 175.0, "start_weight_kg": 80.0,
                        "food_allergies": []},
                       {"sex": "male", "height_cm": 175.0, "start_weight_kg": None,
                        "food_allergies": [], "omissions": []})
    assert r["needs_review"] is True and r["discrepancies"] == []
    assert r["low_confidence"] == ["start_weight_kg"]
    assert r["low_confidence_labels"] == ["peso actual"]
    assert resumen_de_dudas(r) == "confianza baja en peso actual"
    # con todo coincidente, ninguna duda (y la confianza sube)
    ok = comparar_pases({"sex": "male", "start_weight_kg": 80.0}, {"sex": "male", "start_weight_kg": 80.0})
    assert ok["needs_review"] is False and resumen_de_dudas(ok) is None
    assert ok["confidence"]["start_weight_kg"] >= 0.9
    # varios motivos, cada uno con su nombre
    mix = comparar_pases({"sex": "male", "start_weight_kg": 80.0},
                         {"sex": "female", "start_weight_kg": None, "omissions": ["toma omeprazol"]})
    txt = resumen_de_dudas(mix)
    assert "no coincide en 1 dato" in txt and "confianza baja en peso actual" in txt
    assert "1 dato que la relectura echa en falta" in txt


def test_si_el_segundo_pase_falla_la_lectura_sigue_valiendo(monkeypatch):
    from app.services.ai.client import AIClient
    from app.services.ai.extraction import extract_anamnesis_from_document
    from app.services.document_reader import normalizar

    ScriptedDocs([_EXTRACCION_WHATSAPP]).instalar(monkeypatch)   # sin 2ª respuesta
    lectura = extract_anamnesis_from_document(normalizar(b"x", "n.txt"), AIClient(api_key="test"))
    assert lectura.extraction.sex == "female"
    assert "skipped" in lectura.verification and lectura.verification["needs_review"] is False


def test_extract_anamnesis_from_pdf_sigue_funcionando_sin_segundo_pase(monkeypatch):
    from app.services.ai.client import AIClient
    from app.services.ai.extraction import extract_anamnesis_from_pdf

    guion = ScriptedDocs([_EXTRACCION_WHATSAPP]).instalar(monkeypatch)
    e = extract_anamnesis_from_pdf(_pdf(1), AIClient(api_key="test"))
    assert e.sex == "female" and len(guion.calls) == 1


# ===================================================== adjuntos (unidad) ====

def _analitica():
    from app.services.attachments import AttachmentExtraction

    return AttachmentExtraction.model_validate({
        "document_kind": "analitica", "document_date": "2026-05-12", "title": "Bioquímica",
        "summary": "- Glucosa alta\n- Resto normal",
        "lab_values": [
            {"marker": "Glucosa", "value": "101", "unit": "mg/dL", "reference": "70-99", "flag": "H"},
            {"name": "Hemoglobina", "value": 14.1, "unit": "g/dL", "flag": "normal"},
            {"marker": "TSH", "value": "4.9", "unit": "mUI/L", "reference": "0.4-4.0", "flag": "↑"},
            {"marker": "Ferritina", "value": "30"}],
        "clinical": ["Prediabetes según el informe"],
        "medications": [{"name": "Metformina", "dose": "850 mg", "timing": "cena"}],
        "alerts": ["Glucosa en ayunas 101 mg/dL"], "other": None,
    })


def test_adjunto_se_fusiona_en_la_ficha_sin_pisar_y_de_forma_idempotente():
    from types import SimpleNamespace

    from app.services import attachments as at

    c = SimpleNamespace(medical_notes="- Clínica: hipotiroidismo", injuries_notes=None,
                        medication_notes="- Eutirox 50", current_supplements=None,
                        sport_history=None, lifestyle_notes=None)
    ext = _analitica()
    assert at.merge_into_client(c, ext, "adjunto_analitica_ab12.pdf") == ["medical_notes", "medication_notes"]
    assert c.medical_notes.startswith("- Clínica: hipotiroidismo")          # lo del coach, intacto
    assert "[Adjunto: analitica_ab12]" in c.medical_notes
    assert "Glucosa 101 mg/dL (ALTO; ref 70-99)" in c.medical_notes         # fuera de rango, línea propia
    assert "resto normal: Hemoglobina" in c.medical_notes                    # normales, agrupados
    assert "Metformina — 850 mg (cena)" in c.medication_notes
    at.merge_into_client(c, ext, "adjunto_analitica_ab12.pdf")             # releer
    assert c.medical_notes.count("[Adjunto:") == 1
    at.remove_from_client(c, "adjunto_analitica_ab12.pdf")
    assert c.medical_notes == "- Clínica: hipotiroidismo" and c.medication_notes == "- Eutirox 50"


def test_el_bloque_del_adjunto_no_se_come_lo_que_el_coach_escribe_debajo():
    """El bloque se AÑADE al final de la columna, así que todo lo que el coach
    teclee después caía dentro de él: releer o borrar el adjunto le borraba sus
    propias notas. Ahora el bloque va delimitado (cabecera + cierre) y se retira
    EXACTAMENTE lo que se escribió."""
    from types import SimpleNamespace

    from app.services import attachments as at

    def _cli(**kw):
        base = dict(medical_notes=None, injuries_notes=None, medication_notes=None,
                    current_supplements=None, sport_history=None, lifestyle_notes=None)
        base.update(kw)
        return SimpleNamespace(**base)

    ext = _analitica()
    c = _cli(medical_notes="- Clínica: hipotiroidismo (coach)")
    at.merge_into_client(c, ext, "adjunto_analitica_ab12.pdf")
    c.medical_notes += "\n- Clínica: alergia al ibuprofeno (coach, tras llamada)"
    at.merge_into_client(c, ext, "adjunto_analitica_ab12.pdf")     # releer
    assert "ibuprofeno" in c.medical_notes and c.medical_notes.count("[Adjunto:") == 1
    at.remove_from_client(c, "adjunto_analitica_ab12.pdf")          # borrar
    assert "hipotiroidismo" in c.medical_notes and "ibuprofeno" in c.medical_notes
    assert "[Adjunto:" not in c.medical_notes and "Glucosa" not in c.medical_notes

    # Nota del coach ENTRE dos adjuntos: borrar el primero no toca ni la nota
    # ni el bloque del segundo.
    c2 = _cli()
    at.merge_into_client(c2, ext, "adjunto_a_1.pdf")
    c2.medical_notes += "\n- Clínica: nota del coach entre adjuntos"
    at.merge_into_client(c2, ext, "adjunto_b_2.pdf")
    at.remove_from_client(c2, "adjunto_a_1.pdf")
    assert "entre adjuntos" in c2.medical_notes
    assert "[Adjunto: b_2]" in c2.medical_notes and "[Adjunto: a_1]" not in c2.medical_notes

    # Un bloque sin cierre (escrito a mano) también se sabe retirar.
    c3 = _cli(medical_notes="- Clínica: previa\n- [Adjunto: viejo_9] analitica\n- Analítica: algo")
    at.remove_from_client(c3, "adjunto_viejo_9.pdf")
    assert c3.medical_notes == "- Clínica: previa"


def test_dos_adjuntos_conviven_y_el_contexto_para_la_ia_resume_los_leidos(tmp_path, monkeypatch):
    from types import SimpleNamespace

    from app.services import attachments as at

    monkeypatch.setenv("STORAGE_PATH", str(tmp_path))
    from app.config import settings

    monkeypatch.setattr(settings, "storage_path", str(tmp_path), raising=False)
    c = SimpleNamespace(medical_notes=None, injuries_notes=None, medication_notes=None,
                        current_supplements=None, sport_history=None, lifestyle_notes=None)
    at.merge_into_client(c, _analitica(), "adjunto_analitica_1.pdf")
    at.merge_into_client(c, at.AttachmentExtraction(document_kind="informe_fisio",
                                                    injuries=["Rodilla izq: evitar sentadilla profunda"]),
                         "adjunto_fisio_2.jpg")
    assert "[Adjunto: analitica_1]" in c.medical_notes and "[Adjunto: fisio_2]" in c.injuries_notes
    at.save_sidecar(99001, "adjunto_analitica_1.pdf", _analitica(), ["aviso"])
    ctx = at.attachment_context(99001)
    assert ctx and "ADJUNTOS DEL CLIENTE" in ctx and "Glucosa 101" in ctx and "⚠" in ctx
    assert at.resumen_para_ui(at.load_sidecars(99001)[0])["out_of_range"] == ["Glucosa: 101 mg/dL (alto)"]
    at.delete_sidecar(99001, "adjunto_analitica_1.pdf")
    assert at.attachment_context(99001) is None


# ================================================== endpoints (con BD) =======

@pytest.fixture()
def db():
    from app.db import SessionLocal

    s = SessionLocal()
    yield s
    s.close()


@pytest.fixture()
def http():
    from fastapi.testclient import TestClient

    from app.main import app

    with TestClient(app) as c:
        yield c


def _auth():
    from app.security import create_access_token

    return {"Authorization": f"Bearer {create_access_token(os.environ.get('ADMIN_1_USER', 'coach1'))}"}


def _cliente(db, **kw):
    from app.models import Client
    from app.security import new_portal_token

    base = dict(full_name="Lector Universal", email=f"lu-{uuid.uuid4().hex[:8]}@example.com",
                status="onboarding", portal_token="tmp", package_tier="full")
    base.update(kw)
    c = Client(**base)
    db.add(c)
    db.flush()
    c.portal_token = new_portal_token(c.id)
    db.commit()
    return c


@db_required
def test_la_anamnesis_entra_en_word_y_rellena_la_ficha(http, db, monkeypatch):
    from app.services.docs import pdf_convert
    from app.services.storage import anamnesis_documents

    monkeypatch.setattr(pdf_convert, "office_bytes_to_pdf",
                        lambda raw, ext, timeout=120: (_ for _ in ()).throw(RuntimeError("sin soffice")))
    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    docx = _docx(["Cuestionario de otro centro"], [["Peso", "68,4"], ["Alergias", "frutos secos"]])
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                  files={"file": ("cuestionario otro centro.docx", docx,
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["read_ok"] is True and body["format"] == "docx"
    assert body["verification"]["needs_review"] is False
    docs = anamnesis_documents(c.id)
    assert len(docs) == 1 and docs[0]["format"] == "docx" and docs[0]["kind"] == "anamnesis"
    db.expire_all()
    from app.models import Client

    cli = db.get(Client, c.id)
    assert cli.sex == "female" and cli.start_weight_kg == 68.4 and cli.diet_pattern == "vegetariano"
    assert cli.food_allergies == ["frutos secos"]
    assert "Otros datos del documento: Se casa en octubre" in cli.lifestyle_notes
    # constancia de la lectura, disponible sin volver a pagar
    a = http.get(f"/api/clients/{c.id}/anamnesis-analysis", headers=_auth()).json()
    assert a["document_kind"] == "notas" and len(a["source_inventory"]) == 2
    assert a["unmapped_info"][0] == "Se casa en octubre"
    assert a["document"]["name"].endswith(".docx")
    # y se sirve con su MIME real
    v = http.get(f"/api/clients/{c.id}/documents/{docs[0]['name']}", headers=_auth())
    assert v.status_code == 200 and "wordprocessingml" in v.headers["content-type"]


@db_required
def test_diez_fotos_del_cuestionario_son_una_anamnesis_en_pdf(http, db, monkeypatch):
    from app.services.document_reader import _paginas_pdf
    from app.services.storage import abs_path, anamnesis_documents

    guion = ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    fotos = [("files", (f"p{i}.png", _png(500, 700), "image/png")) for i in range(3)]
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(), files=fotos)
    assert r.status_code == 200, r.text
    assert r.json()["format"] == "pdf" and "3 fotos" in r.json()["document"]
    docs = anamnesis_documents(c.id)
    assert len(docs) == 1 and _paginas_pdf(abs_path(docs[0]["rel_path"]).read_bytes()) == 3
    # la IA recibió las 3 imágenes (mejor que un PDF rasterizado)
    assert [b["type"] for b in guion.calls[0]["bloques"]] == ["image", "image", "image"]
    # «Leer con IA» después vuelve a leer ese PDF guardado
    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    r2 = http.post(f"/api/clients/{c.id}/read-anamnesis", headers=_auth())
    assert r2.status_code == 200 and r2.json()["document"]["description"].startswith("PDF de 3")


@db_required
def test_una_analitica_subida_como_anamnesis_se_desvia_a_adjunto_sin_tocar_la_anamnesis(http, db, monkeypatch):
    """El botón equivocado no puede costar el cuestionario: la subida LEE antes
    de guardar y, si la IA ve una analítica/informe/plan, lo guarda como
    ADJUNTO (fusión sin pisar), sin 2º pase y sin tocar la anamnesis."""
    from app.models import Client
    from app.services.storage import anamnesis_documents, list_documents

    # 1) una anamnesis de verdad
    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    r0 = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                   files={"file": ("cuestionario.pdf", _pdf(1), "application/pdf")})
    assert r0.status_code == 200 and r0.json()["redirected_to"] is None
    nombre_anamnesis = anamnesis_documents(c.id)[0]["name"]
    db.expire_all()
    notas_antes = db.get(Client, c.id).medical_notes
    # 2) la analítica por el botón equivocado
    ext = {**_EXTRACCION_WHATSAPP, "document_kind": "analitica", "sex": None,
           "goal_type": None, "start_weight_kg": None}
    guion = ScriptedDocs([ext, _analitica().model_dump()]).instalar(monkeypatch)
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                  files={"file": ("analitica.pdf", _pdf(1), "application/pdf")})
    assert r.status_code == 200, r.text
    b = r.json()
    assert b["redirected_to"] == "adjunto" and b["name"].startswith("adjunto_")
    assert "parece «analitica»" in b["document_warning"] and "no se ha tocado" in b["document_warning"]
    assert b["read_ok"] is True and b["attachment"]["document_kind"] == "analitica"
    # una lectura de anamnesis (sin 2º pase, que sería gasto tirado) + una de adjunto
    assert len(guion.calls) == 2 and "REVISOR" not in guion.calls[1]["system"]
    # la anamnesis es la misma; la ficha conserva lo suyo y SUMA el bloque
    assert [d["name"] for d in anamnesis_documents(c.id)] == [nombre_anamnesis]
    assert sum(1 for d in list_documents(c.id) if d["kind"] == "adjunto") == 1
    db.expire_all()
    cli = db.get(Client, c.id)
    assert cli.medical_notes.startswith(notas_antes.splitlines()[0])
    assert "[Adjunto:" in cli.medical_notes and "Glucosa 101" in cli.medical_notes
    assert cli.sex == "female" and cli.start_weight_kg == 68.4   # nada pisado


@db_required
def test_un_fichero_ilegible_no_destruye_la_anamnesis_anterior(http, db, monkeypatch):
    from app.services.storage import anamnesis_documents

    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    assert http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                     files={"file": ("a.pdf", _pdf(1), "application/pdf")}).status_code == 200
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                  files={"file": ("IMG.HEIC", b"\x00\x00\x00\x18ftypheic" + b"\x00" * 40, "image/heic")})
    assert r.status_code == 422 and "HEIC" in r.json()["detail"]
    assert len(anamnesis_documents(c.id)) == 1


@db_required
def test_el_adjunto_se_lee_entra_en_la_ficha_se_relee_y_se_borra(http, db, monkeypatch):
    from app.models import Client
    from app.services.storage import anamnesis_documents, list_documents

    guion = ScriptedDocs([_analitica().model_dump()]).instalar(monkeypatch)
    c = _cliente(db, medical_notes="- Clínica: hipotiroidismo", status="active")
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(), data={"kind": "adjunto"},
                  files={"file": ("analitica mayo.pdf", _pdf(2), "application/pdf")})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["read_ok"] is True and body["name"].startswith("adjunto_")
    assert body["attachment"]["alerts"] == ["Glucosa en ayunas 101 mg/dL"]
    assert body["attachment"]["out_of_range"] == ["Glucosa: 101 mg/dL (alto)"]
    assert anamnesis_documents(c.id) == []                       # la anamnesis no se toca
    db.expire_all()
    cli = db.get(Client, c.id)
    assert cli.medical_notes.startswith("- Clínica: hipotiroidismo")
    assert "[Adjunto:" in cli.medical_notes and "Glucosa 101 mg/dL (ALTO" in cli.medical_notes
    assert "Metformina — 850 mg (cena)" in cli.medication_notes
    # listado de adjuntos leídos + el documento con su formato
    adj = http.get(f"/api/clients/{c.id}/attachments", headers=_auth()).json()
    assert len(adj) == 1 and adj[0]["n_lab_values"] == 4 and adj[0]["document_kind"] == "analitica"
    assert [d["format"] for d in list_documents(c.id)] == ["pdf"]
    # releer: idempotente
    ScriptedDocs([_analitica().model_dump()]).instalar(monkeypatch)
    r2 = http.post(f"/api/clients/{c.id}/documents/{body['name']}/read", headers=_auth())
    assert r2.status_code == 200 and r2.json()["read_ok"] is True
    db.expire_all()
    assert db.get(Client, c.id).medical_notes.count("[Adjunto:") == 1
    # el contexto de generación lo incluye
    from app.services.attachments import attachment_context

    assert "Glucosa 101" in attachment_context(c.id)
    # borrar: fuera el fichero, el sidecar y su bloque; lo del coach queda
    r3 = http.delete(f"/api/clients/{c.id}/documents/{body['name']}", headers=_auth())
    assert r3.status_code == 204
    assert list_documents(c.id) == [] and attachment_context(c.id) is None
    db.expire_all()
    assert db.get(Client, c.id).medical_notes == "- Clínica: hipotiroidismo"
    assert db.get(Client, c.id).medication_notes is None
    assert guion.calls and guion.calls[0]["bloques"][0]["type"] == "document"


@db_required
def test_leer_la_anamnesis_no_borra_lo_que_aporto_un_adjunto(http, db, monkeypatch):
    """La lectura SUSTITUYE las columnas de notas: se llevaba por delante el
    bloque de la analítica (glucosa alta) y del fisio (evitar sentadilla),
    mientras la ficha seguía diciendo «adjunto leído». Esos textos los leen el
    filtro de ejercicios y la lista roja."""
    from app.models import Client

    c = _cliente(db, status="active")
    ScriptedDocs([_analitica().model_dump()]).instalar(monkeypatch)
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(), data={"kind": "adjunto"},
                  files={"file": ("analitica.pdf", _pdf(1), "application/pdf")})
    assert r.status_code == 200 and r.json()["read_ok"] is True
    # ahora llega el cuestionario y la IA reescribe las notas
    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    r2 = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                   files={"file": ("cuestionario.pdf", _pdf(1), "application/pdf")})
    assert r2.status_code == 200 and r2.json()["read_ok"] is True
    db.expire_all()
    cli = db.get(Client, c.id)
    assert "hipotiroidismo" in cli.medical_notes          # lo que leyó la anamnesis
    assert "[Adjunto:" in cli.medical_notes               # y el bloque del adjunto
    assert "Glucosa 101" in cli.medical_notes
    assert "Metformina" in cli.medication_notes


@db_required
def test_si_la_ia_no_puede_leer_no_se_borra_la_anamnesis_anterior(http, db, monkeypatch):
    """Con la API caída no sabemos qué se acaba de subir: el desvío a adjunto
    no puede actuar, así que la anamnesis anterior se CONSERVA (como adjunto)
    en vez de borrarse para siempre."""
    from app.services.ai.client import AIClient
    from app.services.storage import anamnesis_documents, list_documents

    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    r0 = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                   files={"file": ("cuestionario.pdf", _pdf(1), "application/pdf")})
    assert r0.status_code == 200
    viejo = anamnesis_documents(c.id)[0]["name"]

    def _boom(self, **kw):
        raise RuntimeError("API sin crédito")

    monkeypatch.setattr(AIClient, "_raw_call_with_blocks", _boom)
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(),
                  files={"file": ("analitica.pdf", _pdf(1), "application/pdf")})
    assert r.status_code == 200 and r.json()["read_ok"] is False
    nombres = [d["name"] for d in list_documents(c.id)]
    assert f"adjunto_{viejo}" in nombres, nombres     # el cuestionario sigue ahí
    assert any("se ha conservado entre los adjuntos" in a for a in r.json()["avisos"])
    assert len(anamnesis_documents(c.id)) == 1        # y sigue habiendo UNA anamnesis


@db_required
def test_si_la_ia_falla_el_adjunto_queda_guardado_y_se_dice(http, db, monkeypatch):
    from app.services.ai.client import AIClient

    def _boom(self, **kw):
        raise RuntimeError("API sin crédito")

    monkeypatch.setattr(AIClient, "_raw_call_with_blocks", _boom)
    c = _cliente(db)
    r = http.post(f"/api/clients/{c.id}/documents", headers=_auth(), data={"kind": "adjunto"},
                  files={"file": ("informe.txt", b"Informe fisio: rodilla", "text/plain")})
    assert r.status_code == 200
    assert r.json()["read_ok"] is False and r.json()["read_error"]
    assert r.json()["attachment"] is None
    assert http.get(f"/api/clients/{c.id}/attachments", headers=_auth()).json() == []


@db_required
def test_el_cliente_sube_su_analitica_desde_el_portal_y_el_coach_se_entera(http, db, monkeypatch):
    from app.models import Client
    from app.services import push as push_svc

    avisos = []
    monkeypatch.setattr(push_svc, "send_to_coach", lambda db_, payload: avisos.append(payload))
    ScriptedDocs([_analitica().model_dump()]).instalar(monkeypatch)
    c = _cliente(db, status="active")
    fotos = [("files", (f"a{i}.png", _png(400, 500), "image/png")) for i in range(2)]
    r = http.post(f"/api/p/{c.portal_token}/adjuntos", files=fotos)
    assert r.status_code == 200, r.text
    assert r.json()["ok"] is True and r.json()["read_ok"] is True
    assert "2 fotos" in r.json()["document"]
    # al cliente no le devolvemos valores clínicos
    assert "alerts" not in r.json() and "attachment" not in r.json()
    db.expire_all()
    assert "Glucosa 101" in db.get(Client, c.id).medical_notes
    assert avisos and "ha subido un documento" in avisos[0]["title"]
    assert "Glucosa" in avisos[0]["body"]


@db_required
def test_el_cliente_que_manda_un_informe_como_anamnesis_no_lo_pierde_y_el_coach_lo_sabe(http, db, monkeypatch):
    from app.services import push as push_svc
    from app.services.storage import anamnesis_documents, list_documents

    avisos = []
    monkeypatch.setattr(push_svc, "send_to_coach", lambda db_, payload: avisos.append(payload))
    ext = {**_EXTRACCION_WHATSAPP, "document_kind": "informe_medico"}
    ScriptedDocs([ext, _analitica().model_dump()]).instalar(monkeypatch)
    c = _cliente(db)
    r = http.post(f"/api/p/{c.portal_token}/anamnesis-pdf",
                  files={"file": ("informe.pdf", _pdf(1), "application/pdf")})
    assert r.status_code == 200, r.text
    assert r.json()["ok"] is True and r.json()["redirected_to"] == "adjunto"
    assert anamnesis_documents(c.id) == []                     # sigue faltando el cuestionario
    assert [d["kind"] for d in list_documents(c.id)] == ["adjunto"]
    assert avisos and "Falta su cuestionario" in avisos[0]["body"]


@db_required
def test_portal_adjuntos_rechaza_inactivos_y_ficheros_raros(http, db):
    c = _cliente(db, status="inactive")
    r = http.post(f"/api/p/{c.portal_token}/adjuntos", files={"file": ("a.pdf", _pdf(1), "application/pdf")})
    assert r.status_code == 409
    c2 = _cliente(db, status="active")
    r2 = http.post(f"/api/p/{c2.portal_token}/adjuntos",
                   files={"file": ("x.bin", b"\x00\x01" * 30, "application/octet-stream")})
    assert r2.status_code == 422 and "No reconozco" in r2.json()["detail"]


@db_required
def test_el_cliente_manda_su_anamnesis_en_fotos_desde_el_portal(http, db, monkeypatch):
    from app.models import Client
    from app.services.storage import anamnesis_documents

    ScriptedDocs([_EXTRACCION_WHATSAPP, _VERIFICACION_OK]).instalar(monkeypatch)
    c = _cliente(db)
    fotos = [("files", (f"h{i}.png", _png(400, 500), "image/png")) for i in range(2)]
    r = http.post(f"/api/p/{c.portal_token}/anamnesis-pdf", files=fotos)
    assert r.status_code == 200, r.text
    docs = anamnesis_documents(c.id)
    assert len(docs) == 1 and docs[0]["format"] == "pdf"
    db.expire_all()
    assert db.get(Client, c.id).sex == "female"


# ============================================ plan desde documento ajeno ====

def _plan_doc_json(nombre_ej: str, nombre_ej2: str) -> dict:
    return {
        "document_kind": "plan_completo",
        "nutrition": {
            "declared_kcal": "2300 kcal",
            "meals": [
                {"name": "Desayuno", "time": "8:00", "foods": [
                    {"food": "Avena", "amount": "60 g"}, {"food": "Leche", "amount": "250 ml"},
                    {"food": "Plátano", "amount": "1 ud"}]},
                {"name": "Comida", "foods": [{"food": "Arroz", "grams": "120"},
                                              {"food": "Pechuga de pollo", "amount": "150 g"}]},
                {"name": "Cena", "foods": [{"food": "Salmón", "amount": "150 g"}]},
            ],
            "supplements": [{"name": "Creatina", "dose": "5 g"}],
            "rules": "- Bebe 2 L de agua",
        },
        "training": {
            "split_name": "Torso-Pierna",
            "sessions": [
                {"day": "Lunes", "name": "Torso", "exercises": [
                    {"name": nombre_ej, "sets": "4", "reps": "6-8", "rir": 2, "rest_sec": "2 min"},
                    {"name": "Ejercicio inventado XYZ", "sets": 3, "reps": "10"}]},
                {"day": "Día 2", "name": "Pierna", "exercises": [
                    {"name": nombre_ej2, "sets": 4, "reps": "8-10"}]},
            ],
            "progression": ["Sube 2,5 kg cuando completes"],
            "cardio": {"daily_steps": "9000", "sessions": [{"type": "LISS", "minutes": 30, "times_per_week": 2}]},
        },
        "inventory": ["- Tabla de comidas", "- Tabla de entreno"], "unmapped": [], "warnings": [],
    }


def _cliente_completo(db):
    return _cliente(db, status="active", sex="male", birth_date=date(1990, 3, 12), height_cm=178,
                    start_weight_kg=82, goal_type="fat_loss", level="intermediate", training_days=4,
                    daily_activity_level="light", session_max_min=60, training_place="gym",
                    diet_mode="flexible_7", meals_per_day=4, food_allergies=["lactosa"])


@db_required
def test_plan_ajeno_en_excel_se_transcribe_y_las_cifras_las_pone_el_backend(http, db, monkeypatch):
    from app.models import Exercise, Plan

    ej = db.query(Exercise).filter(Exercise.canonical_name.ilike("%press%banca%")).first()
    ej2 = db.query(Exercise).filter(Exercise.canonical_name.ilike("%sentadilla%")).first()
    assert ej and ej2, "la biblioteca base debe estar sembrada"
    guion = ScriptedDocs([_plan_doc_json(ej.canonical_name, ej2.canonical_name)]).instalar(monkeypatch)
    c = _cliente_completo(db)
    xlsx = _xlsx([["Comida", "Alimento", "Cantidad"], ["Desayuno", "Avena", "60 g"]])
    r = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("dieta otro coach.xlsx", xlsx, "application/octet-stream")})
    assert r.status_code == 200, r.text
    p = r.json()
    nut, tr = p["nutrition_json"], p["training_json"]
    # kcal del CONTRATO, no del documento (y se avisa)
    assert nut["target_kcal"] == p["resumen"]["kcal_contrato"] != 2300
    assert any("declara 2300 kcal" in a for a in p["avisos"])
    # tomas del documento con sus horas (las que faltan, por nombre)
    assert [(m["name"], m["time"]) for m in nut["meals"]] == [("Desayuno", "08:00"), ("Comida", "14:00"), ("Cena", "21:00")]
    # Alimentos contra la base y macros RECALCULADOS desde ella. El plátano
    # venía sin gramos: los pone el SOLVER del backend (la IA no calcula) y se
    # dice; nunca se copian los macros del objetivo de la toma como si fueran
    # los del plato.
    ops = {s["slot"]: s["options"] for s in nut["meal_bank"]["slots"]}
    desayuno = ops[1][0]
    assert desayuno["ingredients"][0]["food_id"] and desayuno["ingredients"][2]["grams"] > 0
    assert "gramos_del_sistema" in desayuno["tags"] and "gramos_del_sistema" not in ops[2][0]["tags"]
    assert desayuno["macros"] != nut["meals"][0]["target"]
    assert any("venían sin gramos" in a for a in p["avisos"])
    # ejercicios: los de la biblioteca entran; el inventado se avisa y no entra
    assert p["resumen"]["ejercicios_reconocidos"] == 2 and p["resumen"]["sesiones"] == 2
    assert any("Ejercicio inventado XYZ" in a for a in p["avisos"])
    ex = tr["sessions"][0]["exercises"][0]
    assert ex["exercise_id"] == ej.id and ex["rest_sec"] == 120 and ex["rep_range"] == "6-8"
    assert tr["sessions"][1]["day"] == "Jueves"            # «Día 2» → reparto real de 2 días
    assert tr["cardio"]["daily_steps"] == 9000
    # la IA recibió la hoja como texto y NO se le pidió calcular
    assert guion.calls[0]["bloques"][0]["type"] == "text" and "No calcules" in guion.calls[0]["user"]
    # nada persistido todavía
    assert db.query(Plan).filter_by(client_id=c.id).count() == 0
    # CONFIRMAR → borrador por el camino de la biblioteca, con el alérgeno cazado
    r2 = http.post(f"/api/clients/{c.id}/plans/import-document/confirm", headers=_auth(),
                   json={"nutrition_json": nut, "training_json": tr, "origen": p["document"]})
    assert r2.status_code == 200, r2.text
    b = r2.json()
    assert b["status"] == "draft" and b["guardrail_flags"][0].startswith("copiado de el documento")
    assert any("ALÉRGENO" in w and "lactosa" in w for w in b["warnings"])
    plan = db.get(Plan, b["id"])
    assert plan.generated_by == "document" and len(plan.training_json["sessions"]) == 2


@db_required
def test_el_plan_importado_no_se_publica_solo_al_guardar_en_el_editor(http, db, monkeypatch):
    """Un borrador «copiado, revísalo» se monta en varias tandas. El importado
    no estaba en la lista de los que NO se activan al guardar: el primer PATCH
    del editor lo PUBLICABA al cliente y `activate_plan` borraba de paso los
    avisos de «copia:» — entre ellos el del ALÉRGENO."""
    from app.models import Exercise, Plan

    ej = db.query(Exercise).filter(Exercise.canonical_name.ilike("%press%banca%")).first()
    ej2 = db.query(Exercise).filter(Exercise.canonical_name.ilike("%sentadilla%")).first()
    ScriptedDocs([_plan_doc_json(ej.canonical_name, ej2.canonical_name)]).instalar(monkeypatch)
    c = _cliente_completo(db)
    prev = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                     files={"file": ("plan.txt", b"Desayuno: avena 60 g", "text/plain")}).json()
    b = http.post(f"/api/clients/{c.id}/plans/import-document/confirm", headers=_auth(),
                  json={"nutrition_json": prev["nutrition_json"],
                        "training_json": prev["training_json"], "origen": prev["document"],
                        "violaciones": prev.get("violaciones") or []}).json()
    plan_id = b["id"]
    assert any("ALÉRGENO" in w for w in b["warnings"])
    # el coach toca UNA cosa en el editor y guarda
    nut = dict(b["nutrition"])
    nut["meals"][0]["time"] = "09:00"
    r = http.patch(f"/api/plans/{plan_id}", headers=_auth(),
                   json={"nutrition_json": nut, "base_rev": (nut.get("rev") or 0)})
    assert r.status_code == 200, r.text
    db.expire_all()
    plan = db.get(Plan, plan_id)
    assert plan.status == "draft", "guardar en el editor NO puede publicar el importado"
    assert plan.published_at is None
    assert any("ALÉRGENO" in str(f) for f in (plan.guardrail_flags or []))


@db_required
def test_documento_mixto_no_tira_la_mitad_de_las_comidas(http, db, monkeypatch):
    """Con comidas con día Y comidas sueltas, el modo lo decide la mayoría y la
    otra lista no se descarta en silencio."""
    doc = {"document_kind": "plan_dieta", "nutrition": {"meals": [
        {"name": "Desayuno", "foods": [{"food": "Avena", "amount": "60 g"}]},
        {"name": "Comida", "foods": [{"food": "Arroz", "amount": "120 g"}]},
        {"name": "Cena", "day": "Lunes", "foods": [{"food": "Salmón", "amount": "150 g"}]},
    ]}, "training": None, "inventory": [], "unmapped": [], "warnings": []}
    ScriptedDocs([doc]).instalar(monkeypatch)
    c = _cliente_completo(db)
    r = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("mixto.txt", b"...", "text/plain")})
    assert r.status_code == 200, r.text
    p = r.json()
    bank = p["nutrition_json"]["meal_bank"]
    assert bank["mode"] == "flexible_7"                     # mayoría sin día
    titulos = [o["title"] for s_ in bank["slots"] for o in s_["options"]]
    assert {"Desayuno", "Comida", "Cena"} <= set(titulos)   # las TRES entran
    assert p["resumen"]["comidas"] == 3


@db_required
def test_un_plato_con_un_alimento_desconocido_no_entra_y_se_dice(http, db, monkeypatch):
    """Sin ficha en la base no se pueden calcular sus macros: NO se inventan
    con el objetivo de la toma (arrastraban el reescalado de las opciones
    buenas). El plato se queda fuera y el aviso lo nombra."""
    doc = {"document_kind": "plan_dieta", "nutrition": {"meals": [
        {"name": "Desayuno", "foods": [{"food": "Batido misterioso XYZ", "amount": "1 vaso"}]},
        {"name": "Comida", "foods": [{"food": "Arroz", "amount": "120 g"}]},
    ]}, "training": None, "inventory": [], "unmapped": [], "warnings": []}
    ScriptedDocs([doc]).instalar(monkeypatch)
    c = _cliente_completo(db)
    p = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("d.txt", b"...", "text/plain")}).json()
    titulos = [o["title"] for s_ in p["nutrition_json"]["meal_bank"]["slots"] for o in s_["options"]]
    assert titulos == ["Comida"]
    assert any("no está en la base" in a and "Desayuno" in a for a in p["avisos"])
    assert any("se han importado 1" in a for a in p["avisos"])


@db_required
def test_importar_plan_exige_la_anamnesis_del_cliente(http, db, monkeypatch):
    ScriptedDocs([_plan_doc_json("Press banca con barra", "Sentadilla trasera con barra")]).instalar(monkeypatch)
    c = _cliente(db, status="active")           # sin sexo, peso, objetivo…
    r = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("plan.txt", b"Desayuno: avena 60 g", "text/plain")})
    assert r.status_code == 422
    assert r.json()["detail"]["missing"]


@db_required
def test_importar_plan_con_fichero_ilegible_o_sin_nada_reconocible(http, db, monkeypatch):
    c = _cliente_completo(db)
    r = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("x.bin", b"\x00\x01" * 30, "application/octet-stream")})
    assert r.status_code == 422 and "No reconozco" in r.json()["detail"]
    ScriptedDocs([{"document_kind": "otro", "nutrition": None, "training": None,
                   "inventory": ["- Una carta"], "unmapped": [], "warnings": []}]).instalar(monkeypatch)
    r2 = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                   files={"file": ("carta.txt", "Estimado cliente…".encode(), "text/plain")})
    assert r2.status_code == 422 and "No se ha reconocido" in r2.json()["detail"]


@db_required
def test_menu_cerrado_por_dias_se_importa_en_modo_strict(http, db, monkeypatch):
    doc = {"document_kind": "plan_dieta", "nutrition": {"meals": [
        {"name": "Desayuno", "day": "Lunes", "foods": [{"food": "Avena", "amount": "60 g"}]},
        {"name": "Cena", "day": "lunes", "foods": [{"food": "Salmón", "amount": "150 g"}]},
        {"name": "Desayuno", "day": "Martes", "foods": [{"food": "Pechuga de pavo", "amount": "100 g"}]},
    ]}, "training": None, "inventory": [], "unmapped": [], "warnings": []}
    ScriptedDocs([doc]).instalar(monkeypatch)
    c = _cliente_completo(db)
    r = http.post(f"/api/clients/{c.id}/plans/import-document", headers=_auth(),
                  files={"file": ("menu.txt", "Lunes…".encode(), "text/plain")})
    assert r.status_code == 200, r.text
    bank = r.json()["nutrition_json"]["meal_bank"]
    # Un menú CERRADO tiene que cubrir los 7 días: el portal se quedaba sin
    # comidas el resto de la semana. Los que faltan se completan repitiendo los
    # que hay, y se dice cuáles son copia.
    assert bank["mode"] == "strict" and [d["day"] for d in bank["days"]] == list(_DIAS_SEMANA)
    assert len(bank["days"][0]["meals"]) == 2          # lunes: desayuno + cena
    assert any("solo cubría 2 de los 7 días" in a for a in r.json()["avisos"])
    copia = next(d for d in bank["days"] if d["day"] == "miercoles")
    assert copia["meals"][0]["dish"]["title"] == bank["days"][0]["meals"][0]["dish"]["title"]
