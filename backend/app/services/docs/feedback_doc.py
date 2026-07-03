"""Documento de feedback quincenal/mensual con gráficas (H.4).

Estructura: resumen del período en datos (peso+tendencia, adherencia,
perímetros, volumen) → progresión de fuerza (e1RM) → composición física (fotos
lado a lado + análisis IA) → análisis en lenguaje natural → "qué ha cambiado y
por qué" (máx 5 bullets) → respuesta a dudas + objetivos + cierre.

Las gráficas (services/docs/charts) usan datos ya calculados por
services/metrics. Las imágenes se incrustan desde BytesIO.
"""

from __future__ import annotations

import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.services.docs import charts
from app.services.docs.word_base import (
    DocBrand,
    add_bullets,
    add_cards_row,
    add_cover,
    add_section_heading,
    clean_table,
    init_document,
)


def _add_chart(doc: Document, png: bytes, width_in: float = 6.0) -> None:
    doc.add_picture(io.BytesIO(png), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_feedback_doc(
    *,
    brand: DocBrand,
    client_name: str,
    period_index: int,
    metrics: dict,
    weight_points: list[tuple[str, float]],
    goal_kg: float | None,
    e1rm_exercises: list[dict],
    perimeters: dict[str, list[tuple[str, float]]] | None,
    volume_by_group: dict[str, float] | None,
    photo_pairs: list[tuple[str, str]] | None,
    ai_photo_analysis: str | None,
    natural_analysis: str,
    changes_bullets: list[str],
    answers: str | None,
    next_objectives: list[str],
    closing_message: str,
    plan_adjustments: list[dict] | None = None,
) -> bytes:
    doc = init_document(brand)
    accent = brand.color_primary

    add_cover(doc, brand, client_name,
              subtitle=f"Informe de progreso · Período {period_index}",
              goal="Tu evolución en datos")

    # 1) Resumen del período en datos
    add_section_heading(doc, brand, "Tu período en datos")
    adh = metrics.get("adherence", {})
    weight = metrics.get("weight", {})
    add_cards_row(doc, brand, [
        ("Cambio de peso", _fmt_delta(weight.get("delta_kg"), "kg")),
        ("Adherencia dieta", f"{round(adh.get('diet_adherence_ratio', 0) * 100)}%"),
        ("Días registrados", f"{adh.get('days_logged', 0)}/{adh.get('period_days', 0)}"),
    ])
    doc.add_paragraph()

    if weight_points:
        doc.add_heading("Evolución de peso", level=2)
        _add_chart(doc, charts.weight_trend_chart(weight_points, goal_kg, accent))

    doc.add_heading("Adherencia", level=2)
    diet_pct = adh.get("diet_adherence_ratio", 0) * 100
    train_pct = min(100, adh.get("log_ratio", 0) * 100)
    _add_chart(doc, charts.adherence_chart(diet_pct, train_pct, accent), width_in=5.5)

    if perimeters:
        doc.add_heading("Perímetros", level=2)
        _add_chart(doc, charts.perimeters_chart(perimeters, accent))

    if volume_by_group:
        doc.add_heading("Volumen por grupo muscular", level=2)
        _add_chart(doc, charts.volume_by_group_chart(volume_by_group, accent))

    # 2) Progresión de fuerza
    if e1rm_exercises:
        doc.add_page_break()
        add_section_heading(doc, brand, "Progresión de fuerza")
        doc.add_paragraph(
            "Fuerza estimada (1RM por Epley) de tus ejercicios principales."
        )
        _add_chart(doc, charts.e1rm_chart(e1rm_exercises, accent))

    # 3) Composición física
    if (photo_pairs or ai_photo_analysis):
        doc.add_page_break()
        add_section_heading(doc, brand, "Composición física")
        if photo_pairs:
            for before, after in photo_pairs:
                _add_photo_pair(doc, before, after)
        if ai_photo_analysis:
            doc.add_paragraph(ai_photo_analysis)

    # 4) Análisis en lenguaje natural
    doc.add_page_break()
    add_section_heading(doc, brand, "Cómo ha ido")
    doc.add_paragraph(natural_analysis)

    # 5) Qué ha cambiado y por qué (máx 5 bullets)
    if changes_bullets:
        doc.add_heading("Qué ha cambiado en tu plan y por qué", level=2)
        add_bullets(doc, changes_bullets[:5])

    # 5b) CUADRÍCULA DE CAMBIOS — tabla de ajustes aplicados al plan
    if plan_adjustments:
        doc.add_heading("Cuadrícula de cambios aplicados", level=2)
        rows = [
            [str(a.get("area", "")), str(a.get("change", "")), str(a.get("reason", ""))]
            for a in plan_adjustments
        ]
        clean_table(
            doc, ["Área", "Cambio", "Por qué"], rows, brand,
            header_color=brand.color_primary, header_text_color="FFFFFF",
            col_widths=[1800, 3600, 3626],
        )

    # 6) Dudas + objetivos + cierre
    if answers:
        doc.add_heading("Tus dudas", level=2)
        doc.add_paragraph(answers)

    if next_objectives:
        doc.add_heading("Objetivos para las próximas 2 semanas", level=2)
        add_bullets(doc, next_objectives)

    p = doc.add_paragraph()
    p.add_run(closing_message).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_photo_pair(doc: Document, before_path: str, after_path: str) -> None:
    """Dos fotos lado a lado (antes/después) emparejadas por ángulo."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    headers = table.rows[0].cells
    for i, label in enumerate(("Período anterior", "Período actual")):
        p = headers[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = run.font.color.rgb  # mantiene color por defecto
    cells = table.rows[1].cells
    for i, path in enumerate((before_path, after_path)):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path and os.path.exists(path):
            try:
                run = p.add_run()
                run.add_picture(path, width=Inches(2.6))
            except Exception:
                p.add_run("(imagen no disponible)")
        else:
            p.add_run("—")


def _fmt_delta(value: float | None, unit: str) -> str:
    if value is None:
        return "—"
    sign = "+" if value > 0 else ""
    return f"{sign}{value} {unit}"
