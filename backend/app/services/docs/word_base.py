"""Helpers de generación Word con python-docx, tema claro con marca (H.3).

Centraliza el estilo (tipografía, colores de marca, espaciados) y las
primitivas de maquetación (portada, cabeceras de sección, cards de resumen,
tablas limpias con ancho explícito). Las reglas de oro de tablas (ancho en DXA,
sin viñetas unicode, padding de celda, sombreado CLEAR) siguen las del skill de
docx, aplicadas al equivalente de python-docx.

Tanto el documento de plan como el de feedback construyen sobre estas piezas
para garantizar un aspecto coherente y profesional.
"""

from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Ancho de contenido en twips: A4 (11906) menos los márgenes reales de
# init_document (56 pt = 1120 twips por lado). Debe cuadrar con ellos para que
# barras (párrafo sombreado con sangría) y cajas (tablas) queden alineadas.
CONTENT_WIDTH_DXA = 9666


@dataclass
class DocBrand:
    name: str
    color_primary: str   # "#6EE7B7"
    color_secondary: str
    font_family: str
    tagline: str | None = None
    contact_email: str | None = None
    logo_path: str | None = None  # ruta absoluta a imagen, opcional


def _hex(color: str) -> RGBColor:
    return RGBColor.from_string(color.lstrip("#").upper())


def _shade_cell(cell, hex_color: str) -> None:
    """Sombreado de celda (equivale a ShadingType.CLEAR del skill).

    shd debe ir al inicio de tcPr (antes de tcMar/tcW) según el esquema OOXML.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    # Quita un shd previo si existiera (evita duplicados en zebra+header)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto",
        qn("w:fill"): hex_color.lstrip("#").upper(),
    })
    # shd va después de tcW/gridSpan si existen, antes de tcMar
    tcW = tcPr.findall(qn("w:tcW"))
    if tcW:
        tcW[-1].addnext(shd)
    else:
        tcPr.insert(0, shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    m = tcPr.makeelement(qn("w:tcMar"), {})
    # El esquema de tcMar exige el orden: top, left (start), bottom, right (end)
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = m.makeelement(qn(f"w:{side}"), {qn("w:w"): str(val), qn("w:type"): "dxa"})
        m.append(node)
    tcPr.append(m)


def float_image_right(paragraph, image_path: str, width) -> None:
    """Inserta una imagen FLOTANTE a la derecha con el texto fluyendo alrededor
    (como las fotos del ejemplo del coach, superpuestas entre las palabras)."""
    from docx.oxml import OxmlElement

    def _wp(tag):
        return qn("wp:" + tag)

    run = paragraph.add_run()
    run.add_picture(image_path, width=width)  # inline de partida
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(_wp("inline"))
    extent = inline.find(_wp("extent"))
    cx, cy = extent.get("cx"), extent.get("cy")
    docPr = inline.find(_wp("docPr"))
    cNv = inline.find(_wp("cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
                 "simplePos": "0", "relativeHeight": "2", "behindDoc": "0", "locked": "0",
                 "layoutInCell": "1", "allowOverlap": "1"}.items():
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); anchor.append(sp)
    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "column")
    al = OxmlElement("wp:align"); al.text = "right"; ph.append(al); anchor.append(ph)
    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "paragraph")
    off = OxmlElement("wp:posOffset"); off.text = "0"; pv.append(off); anchor.append(pv)
    ext = OxmlElement("wp:extent"); ext.set("cx", cx); ext.set("cy", cy); anchor.append(ext)
    ee = OxmlElement("wp:effectExtent")
    for k in ("l", "t", "r", "b"):
        ee.set(k, "0")
    anchor.append(ee)
    wrap = OxmlElement("wp:wrapSquare"); wrap.set("wrapText", "bothSides"); anchor.append(wrap)
    anchor.append(docPr); anchor.append(cNv); anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)


def _header_bg_image(paragraph, image_path: str, width) -> None:
    """Imagen de cabecera como FONDO flotante (detrás del texto, anclada al borde
    superior de la página), para que el contenido se superponga encima (como el
    ejemplo: título y cajas sobre la banda de comida)."""
    from docx.oxml import OxmlElement

    def _wp(tag):
        return qn("wp:" + tag)

    run = paragraph.add_run()
    run.add_picture(image_path, width=width)
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(_wp("inline"))
    extent = inline.find(_wp("extent"))
    cx, cy = extent.get("cx"), extent.get("cy")
    docPr = inline.find(_wp("docPr"))
    cNv = inline.find(_wp("cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0",
                 "relativeHeight": "0", "behindDoc": "1", "locked": "0", "layoutInCell": "1",
                 "allowOverlap": "1"}.items():
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); anchor.append(sp)
    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "page")
    al = OxmlElement("wp:align"); al.text = "center"; ph.append(al); anchor.append(ph)
    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "page")
    off = OxmlElement("wp:posOffset"); off.text = "0"; pv.append(off); anchor.append(pv)
    ext = OxmlElement("wp:extent"); ext.set("cx", cx); ext.set("cy", cy); anchor.append(ext)
    ee = OxmlElement("wp:effectExtent")
    for k in ("l", "t", "r", "b"):
        ee.set(k, "0")
    anchor.append(ee)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(docPr); anchor.append(cNv); anchor.append(graphic)
    drawing.remove(inline); drawing.append(anchor)


def _cant_split_rows(table) -> None:
    """Evita que una fila se parta entre páginas (texto/celdas cortados).

    OJO: úsalo solo en filas BAJAS. Una fila más alta que la página no puede
    partirse y su contenido sobrante se RECORTA (se pierde texto). Para filas
    potencialmente altas (cajas de equivalencias, listas largas, texto libre)
    hay que dejar que la fila se parta entre páginas.
    """
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))


def _mark_header_repeat(table) -> None:
    """Marca la fila 0 como cabecera repetible (w:tblHeader): si la tabla se
    parte entre páginas, la cabecera de color se vuelve a dibujar arriba de la
    página siguiente en lugar de dejar filas sin encabezado."""
    if not table.rows:
        return
    trPr = table.rows[0]._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        el = trPr.makeelement(qn("w:tblHeader"), {})
        el.set(qn("w:val"), "true")
        trPr.append(el)


def _keep_with_next(paragraph) -> None:
    """La barra/encabezado se queda con el contenido que le sigue (no se orfana)."""
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(pPr.makeelement(qn("w:keepNext"), {}))


def _keep_rows_together(table) -> None:
    """Mantiene TODAS las filas juntas (la tabla no se corta entre páginas)."""
    rows = table.rows
    for row in list(rows)[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                _keep_with_next(p)


def open_box(doc: Document, fill: str = "F5F0E8", cant_split: bool = False):
    """Crea una caja (cell) a todo el ancho con relleno y devuelve la celda para
    rellenarla con párrafos. Sin bordes.

    cant_split=False (por defecto): si el contenido no cabe en la página, la caja
    SE PARTE entre páginas conservando el sombreado a ambos lados — nunca se
    recorta texto. Solo pon cant_split=True para cajas cortas que quieras
    mantener enteras (p. ej. una nota de una línea)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Pt(CONTENT_WIDTH_DXA / 20)
    _shade_cell(cell, fill)
    _set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    _box_border(table)  # marco gris fino, como la referencia
    if cant_split:
        _cant_split_rows(table)
    return cell


def init_document(brand: DocBrand) -> Document:
    """Documento con estilos base de marca (tipografía y headings)."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = brand.font_family
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _hex("#1A1A24")

    for level, size in (("Heading 1", 20), ("Heading 2", 14), ("Heading 3", 11.5)):
        st = doc.styles[level]
        st.font.name = brand.font_family
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = _hex("#1A1A24")

    # A4 explícito (python-docx crea Letter por defecto) + márgenes de 2 cm.
    # CONTENT_WIDTH_DXA depende de estas medidas: 11906 − 2×1120 = 9666.
    from docx.shared import Mm
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = section.bottom_margin = Pt(56)
        section.left_margin = section.right_margin = Pt(56)

    # El zoom por defecto de python-docx (val="bestFit") sin percent falla la
    # validación OOXML estricta; fijamos percent=100.
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")
    return doc


def add_cover(doc: Document, brand: DocBrand, client_name: str, subtitle: str,
              goal: str) -> None:
    """Portada: marca, nombre del cliente, mes/objetivo."""
    import os

    from docx.shared import Inches

    if brand.logo_path and os.path.exists(brand.logo_path):
        try:
            doc.add_picture(brand.logo_path, width=Inches(1.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(brand.name)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = _hex(brand.color_primary)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(client_name)
    run.font.size = Pt(30)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = _hex("#6B6B76")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(goal)
    run.font.size = Pt(12)
    run.font.color.rgb = _hex(brand.color_secondary)

    doc.add_page_break()


def add_section_heading(doc: Document, brand: DocBrand, text: str) -> None:
    """Encabezado de sección con regla inferior de color de marca."""
    h = doc.add_heading(text, level=1)
    # Regla inferior (border en el párrafo, no tabla — regla del skill)
    pPr = h._p.get_or_add_pPr()
    borders = pPr.makeelement(qn("w:pBdr"), {})
    bottom = borders.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "12",
        qn("w:space"): "4", qn("w:color"): brand.color_primary.lstrip("#").upper(),
    })
    borders.append(bottom)
    pPr.append(borders)


def add_cards_row(doc: Document, brand: DocBrand, cards: list[tuple[str, str]]) -> None:
    """Fila de 'cards' visuales (label, value) para el resumen ejecutivo."""
    n = len(cards)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_w = CONTENT_WIDTH_DXA // n
    for i, (label, value) in enumerate(cards):
        cell = table.rows[0].cells[i]
        cell.width = Pt(col_w / 20)
        _shade_cell(cell, "F4F4F7")
        _set_cell_margins(cell, top=120, bottom=120)
        cell.paragraphs[0].text = ""
        pv = cell.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(value)
        rv.font.size = Pt(17)
        rv.font.bold = True
        rv.font.color.rgb = _hex(brand.color_primary)
        pl = cell.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(label)
        rl.font.size = Pt(8.5)
        rl.font.color.rgb = _hex("#6B6B76")
    _no_table_borders(table)


def clean_table(doc: Document, headers: list[str], rows: list[list[str]],
                brand: DocBrand, col_widths: list[int] | None = None,
                header_color: str | None = None, header_colors: list[str] | None = None,
                header_text_color: str = "0A0A0F", font_pt: float = 9.5,
                cant_split_rows: bool = True, keep_together: bool = True):
    """Tabla limpia con cabecera de color, ancho explícito y padding (skill).

    header_color: color único de la cabecera (por defecto el de marca).
    header_colors: color por columna (p. ej. los 4 grupos de alimentos).
    font_pt: tamaño de fuente de las celdas (baja a 8 en tablas muy anchas).
    cant_split_rows: True para filas bajas (no se parten a media fila). Ponlo a
        False cuando alguna fila pueda ser MÁS ALTA que la página (celdas con
        listas largas o texto libre): así la fila se parte entre páginas en vez
        de recortarse y perder contenido.
    keep_together: True intenta mantener la tabla entera en una página. Ponlo a
        False en tablas potencialmente largas (semanal, grupos de alimentos,
        cambios) para que paginen limpiamente repitiendo la cabecera."""
    spacer(doc, SPACE_INNER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths is None:
        col_widths = [CONTENT_WIDTH_DXA // len(headers)] * len(headers)
    base_hdr = (header_color or brand.color_primary)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Pt(col_widths[i] / 20)
        _shade_cell(hdr[i], (header_colors[i] if header_colors else base_hdr).lstrip("#"))
        _set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # cabeceras centradas (referencia)
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = _hex(header_text_color)

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Pt(col_widths[i] / 20)
            _set_cell_margins(cells[i])
            # SIEMPRE relleno opaco (blanco/crema) para que el texto sea legible
            # aunque la fila quede sobre la banda de comida de la cabecera.
            _shade_cell(cells[i], "F5F0E8" if r_idx % 2 == 1 else "FFFFFF")
            # Valor LISTA = varias líneas en la celda, cada una su párrafo, con
            # etiqueta opcional en negrita: [("Cereales", "avena, arroz…"), …]
            # — el formato de subgrupos del plan de referencia.
            if isinstance(val, (list, tuple)):
                first = True
                for item in val:
                    p = cells[i].paragraphs[0] if first else cells[i].add_paragraph()
                    first = False
                    p.paragraph_format.space_after = Pt(3)
                    label, text = (item if isinstance(item, (list, tuple)) else ("", item))
                    if label:
                        rl = p.add_run(f"{label}: ")
                        rl.font.bold = True
                        rl.font.size = Pt(font_pt)
                    rt = p.add_run(str(text))
                    rt.font.size = Pt(font_pt)
            else:
                p = cells[i].paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(font_pt)
    _thin_borders(table)
    # Cabecera repetible: si la tabla se parte, la cabecera de color reaparece
    # en la página siguiente (nunca filas huérfanas sin encabezado).
    _mark_header_repeat(table)
    if cant_split_rows:
        _cant_split_rows(table)
    if keep_together:
        _keep_rows_together(table)  # tablas cortas: enteras en una página
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    """Lista con viñetas usando el estilo nativo (nunca viñetas unicode, skill)."""
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


# Aire vertical del diseño (medido del plan de referencia del coach): ~19 pt
# entre tarjeta y tarjeta; ~8 pt entre una barra y una tabla de contenido.
SPACE_SECTION = 18
SPACE_INNER = 8


def spacer(doc: Document, pt: float = SPACE_SECTION, keep_next: bool = True) -> None:
    """Aire vertical EXACTO entre bloques (párrafo vacío con altura de línea
    exacta): invisible, estable en Word y LibreOffice, y de paso evita que dos
    tablas contiguas se fusionen al editar el .docx. keep_next: el aire viaja
    con el bloque siguiente (no se queda huérfano al pie de página)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(pt)
    if keep_next:
        _keep_with_next(p)


def section_bar(doc: Document, text: str, color: str, text_color: str = "FFFFFF",
                size: float = 11) -> None:
    """Barra de sección a ancho de contenido con fondo de color y texto centrado.
    SIEMPRE con su aire por delante: cada título abre una tarjeta nueva separada
    de la anterior (regla del diseño de referencia).

    Es un PÁRRAFO sombreado, no una tabla: keepNext en un párrafo sí ancla la
    barra al bloque siguiente en Word Y LibreOffice, mientras que una barra-tabla
    se queda huérfana al pie de página (LibreOffice ignora keepNext entre
    tablas). El grosor vertical lo dan bordes del mismo color que el fondo."""
    spacer(doc, SPACE_SECTION)
    fill = color.lstrip("#")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # el sombreado de párrafo va de margen a margen: sangrar para alinear la
    # barra con las cajas/tablas (que miden CONTENT_WIDTH_DXA centradas)
    sec = doc.sections[0]
    usable_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    indent_emu = (usable_emu - int(Pt(CONTENT_WIDTH_DXA / 20))) // 2
    if indent_emu > 0:
        pf.left_indent = indent_emu
        pf.right_indent = indent_emu
    r = p.add_run(text.upper())
    r.font.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = _hex(text_color)
    _keep_with_next(p)  # la barra no se queda sola al pie de página
    pPr = p._p.get_or_add_pPr()
    # bordes superior/inferior del color del fondo = "relleno" vertical de la
    # barra; sin bordes laterales para que el ancho coincida EXACTO con las cajas
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    for edge in ("top", "bottom"):
        e = pbdr.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "32",
            qn("w:space"): "0", qn("w:color"): fill,
        })
        pbdr.append(e)
    pPr.append(pbdr)
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill,
    })
    pPr.append(shd)


def info_box(doc: Document, items, fill: str = "F5F0E8", label_color: str = "8B1A2B",
             cant_split: bool = False, image_path: str | None = None) -> None:
    """Recuadro con fondo crema. items: str (línea) o (label, valor).

    cant_split=False (por defecto): si el recuadro no cabe en la página, se parte
    entre páginas conservando el fondo — nunca recorta texto. Ponlo a True solo
    en recuadros cortos que quieras mantener enteros."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Pt(CONTENT_WIDTH_DXA / 20)
    _shade_cell(cell, fill)
    _set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, (tuple, list)):
            label, value = item
            rl = p.add_run(f"{label}: ")
            rl.font.bold = True
            rl.font.color.rgb = _hex(label_color)
            rl.font.size = Pt(10)
            rv = p.add_run(value)
            rv.font.size = Pt(10)
        else:
            r = p.add_run(str(item))
            r.font.size = Pt(10)
    if image_path:
        import os

        from docx.shared import Inches

        if os.path.exists(image_path):
            pimg = cell.add_paragraph()
            pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pimg.paragraph_format.space_before = Pt(6)
            try:
                pimg.add_run().add_picture(image_path, width=Inches(2.4))
            except Exception:
                pass
    _box_border(table)  # marco gris fino, como la referencia
    if cant_split:
        _cant_split_rows(table)


def setup_reference_pages(doc: Document, logo_path: str | None,
                          right_title: str, right_sub: str | None = None,
                          footer_text: str | None = None) -> None:
    """Cabecera COMPACTA del plan de referencia, en TODAS las páginas: logo a la
    izquierda y "PLAN NUTRICIONAL | Cliente" (+ año debajo) a la derecha, sobre
    blanco. Sin banda de fondo: el contenido NUNCA se superpone a una imagen y
    todas las páginas empiezan a la misma altura."""
    import os

    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches

    # Pie reutilizado (texto + nº de página); sin banda.
    setup_branded_pages(doc, banner_path=None, footer_text=footer_text)

    section = doc.sections[0]
    section.different_first_page_header_footer = False  # misma cabecera SIEMPRE
    section.top_margin = Inches(1.15)
    section.header_distance = Inches(0.35)

    header = section.header
    ht = header.add_table(rows=1, cols=2, width=Pt(CONTENT_WIDTH_DXA / 20))
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    ht.autofit = False
    left, right = ht.rows[0].cells
    left.width = Pt(CONTENT_WIDTH_DXA / 40)
    right.width = Pt(CONTENT_WIDTH_DXA / 40)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    if logo_path and os.path.exists(logo_path):
        try:
            lp.add_run().add_picture(logo_path, height=Inches(0.42))
        except Exception:
            pass
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run(right_title)
    rr.font.size = Pt(10)
    rr.font.bold = True
    if right_sub:
        rp2 = right.add_paragraph()
        rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp2.paragraph_format.space_after = Pt(0)
        r2 = rp2.add_run(right_sub)
        r2.font.size = Pt(10)
    _no_table_borders(ht)
    # El párrafo vacío por defecto del header queda DESPUÉS de la tabla con
    # altura mínima (los headers OOXML deben terminar en párrafo).
    hp = header.paragraphs[0]
    hp._p.getparent().remove(hp._p)
    tail = header.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing = Pt(2)


def setup_branded_pages(doc: Document, banner_path: str | None = None,
                        footer_text: str | None = None) -> None:
    """Cabecera con banda de marca (en páginas de contenido, no en portada) + pie."""
    import os

    from docx.shared import Inches

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    if banner_path and os.path.exists(banner_path):
        # Banda de comida TRANSLÚCIDA a sangre completa como FONDO; el título y las
        # cajas se superponen encima y se leen (la banda está atenuada como el
        # ejemplo del coach). El contenido empieza arriba, sobre la banda.
        section.top_margin = Inches(1.45)
        section.header_distance = Inches(0.0)
        hp = section.header.paragraphs[0]
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        try:
            _header_bg_image(hp, banner_path, Inches(8.8))  # banda full-bleed, detrás
        except Exception:
            pass
    if footer_text:
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement

        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Como la referencia: nº de página a la IZQUIERDA y lema CENTRADO
        fp.paragraph_format.tab_stops.add_tab_stop(
            Pt(CONTENT_WIDTH_DXA / 40), WD_TAB_ALIGNMENT.CENTER
        )
        # campo PAGE (Word/LibreOffice lo numeran solos)
        pr = fp.add_run()
        pr.font.size = Pt(8)
        pr.font.color.rgb = _hex("#9A9AA6")
        beg = OxmlElement("w:fldChar"); beg.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = "PAGE"
        end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
        pr._r.append(beg); pr._r.append(ins); pr._r.append(end)
        fp.add_run("\t")
        r = fp.add_run(footer_text)
        r.font.size = Pt(8)
        r.font.color.rgb = _hex("#9A9AA6")


def branded_cover(doc: Document, cover_path: str | None) -> None:
    """Portada con la imagen de marca centrada."""
    import os

    from docx.shared import Inches

    for _ in range(5):
        doc.add_paragraph()
    if cover_path and os.path.exists(cover_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(cover_path, width=Inches(4.8))
        except Exception:
            pass
    doc.add_page_break()


def _thin_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "0", qn("w:color"): "E0E0E6",
        })
        borders.append(e)
    _insert_tbl_borders(tblPr, borders)


def _box_border(table, color: str = "999999") -> None:
    """Borde EXTERIOR fino (gris de la referencia) sin líneas interiores:
    el marco de las cajas de contenido (info_box / open_box)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        e = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "0", qn("w:color"): color,
        })
        borders.append(e)
    for edge in ("insideH", "insideV"):
        e = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none"})
        borders.append(e)
    _insert_tbl_borders(tblPr, borders)


def _no_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none"})
        borders.append(e)
    _insert_tbl_borders(tblPr, borders)


def _insert_tbl_borders(tblPr, borders) -> None:
    """Inserta tblBorders en la posición correcta del esquema OOXML.

    El orden en tblPr es estricto: ...tblW, jc, tblCellSpacing, tblInd,
    tblBorders, shd, ... Insertamos tblBorders justo después de tblInd/jc/tblW
    (lo que exista) y antes de cualquier shd.
    """
    after = ("w:tblInd", "w:tblCellSpacing", "w:jc", "w:tblW", "w:tblStyle")
    anchor = None
    for tag in after:
        found = tblPr.findall(qn(tag))
        if found:
            anchor = found[-1]
            break
    if anchor is not None:
        anchor.addnext(borders)
    else:
        tblPr.insert(0, borders)
