"""Documento Word del plan — diseño de marca DQ (réplica del ejemplo del coach).

Un único documento con la estética del plan oficial: portada con logo, banda de
comida en la cabecera, barras de sección de color, tablas con cabecera de color,
cajas crema. Incluye NUTRICIÓN (objetivos, resumen energético, estructura diaria,
alimentos por grupos, plato saludable, comidas, dieta semanal, ideas, recomenda-
ciones, suplementación) y, a continuación, ENTRENAMIENTO en el mismo estilo.

El contenido cambia según el cliente (datos ya calculados); el diseño es fijo.
Secciones genéricas (alimentos por grupos, plato, ideas, recomendaciones) son
plantilla, filtrando alimentos por alergias/aversiones.
"""

from __future__ import annotations

import io
import os
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.services.docs.word_base import (
    DocBrand,
    clean_table,
    float_image_right,
    info_box,
    init_document,
    open_box,
    section_bar,
    setup_reference_pages,
    _hex,
    _keep_lines,
    _keep_with_next,
)
from app.services.metrics import _rhu

ASSETS = Path(__file__).resolve().parent.parent.parent / "assets" / "plan"

# Paleta EXACTA extraída del PDF de ejemplo del coach
WINE = "8B1A2B"
BLUE = "4A7BA8"
GOLD = "C9A961"   # barra de "Estructura diaria"
CREAM = "F5F0E8"  # relleno de cajas y zebra de tablas
# Colores de las 4 columnas de "Alimentos por grupos" (verbatim del ejemplo)
FG_GREEN = "2E7D32"
FG_YELLOW = "F1C232"
FG_WINE = "8B1A2B"
FG_ORANGE = "E69138"

DAYS = ["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]

# --- Contenido de plantilla (genérico, VERBATIM del ejemplo del coach) ---
# Estructurado como [(etiqueta, [alimentos…])] para poder filtrar alergias/
# aversiones SIN romper etiquetas ni alimentos contiguos.
FOOD_GROUPS = {
    "VEGETALES": [
        ("", ["Acelga", "ajo", "alcachofas", "apio", "berenjena", "brócoli", "calabacín",
              "pepino", "pimiento", "puerro", "rábano", "remolacha", "zanahoria", "coliflor",
              "endivia", "escarola", "espárragos", "espinacas", "judías verdes", "calabaza",
              "nabo", "cebolla", "col lombarda", "coles de Bruselas"]),
    ],
    "CARBOHIDRATOS": [
        ("Féculas y tubérculos", ["patata", "boniato", "yuca"]),
        ("Pseudocereales", ["amaranto", "quinoa", "trigo sarraceno"]),
        ("Cereales", ["cebada", "maíz", "arroz integral", "kamut", "centeno", "sorgo",
                      "teff", "mijo", "avena", "bulgur", "espelta"]),
    ],
    "PROTEÍNAS": [
        ("Proteína animal", ["carne", "pescado", "huevo"]),
        ("Legumbres y derivados", ["cacahuete", "azukis", "edamame", "garbanzos", "habas",
                                   "lentejas", "guisantes", "soja", "tempeh", "tofu",
                                   "seitán", "heura"]),
        ("Lácteos", ["leche", "cuajada", "kéfir", "yogurt"]),
    ],
    "LÍPIDOS": [
        ("", ["Aguacate", "aceite de oliva", "aceitunas"]),
        ("Frutos secos", ["almendras", "anacardos", "nueces", "avellanas", "castañas",
                          "pistachos", "cacahuete"]),
        ("Semillas", ["chía", "calabaza", "lino", "girasol", "sésamo"]),
    ],
}
FOOD_GROUP_FOOTNOTE: dict[str, str] = {}  # (sin trivia: no cambia ninguna decisión)
FOOD_GROUP_COLORS = [FG_GREEN, FG_YELLOW, FG_WINE, FG_ORANGE]

PLATO_TEXT = [
    "• ½ plato: verdura y fruta (cuanta más variedad, mejor).",
    "• ¼ plato: integrales, féculas y tubérculos.",
    "• ¼ plato: proteína — limita carne roja y procesada.",
    "• Grasa: AOVE, aguacate o frutos secos. Bebida: agua.",
]

IDEAS_RAPIDAS = [
    # Una idea por línea: el filtro de alérgenos descarta la LÍNEA entera, así
    # que agruparlas escondería ideas seguras. La viñeta la pone el renderizador.
    "Pan integral con queso cottage y aguacate.",
    "Pan integral con pavo, jamón o huevo.",
    "Pan integral con crema de cacahuete y plátano.",
    "Pan integral con hummus y tomate.",
    "Yogur o queso batido 0% con avena y fruta.",
    "Tortitas de arroz con crema de cacahuete y plátano.",
]

SALSAS_TEXT = [
    "Tomate triturado natural (sin azúcar añadido), mostaza Dijon, vinagre balsámico/de "
    "manzana/de Módena, salsa de soja baja en sodio, salsa tamari, salsa Sriracha (con "
    "moderación), salsa de yogur natural con limón y especias, salsa romesco casera, pesto "
    "casero (con moderación por las grasas), guacamole casero, hummus, salsa tahini, mayonesa "
    "light o de aguacate (con moderación), mojo verde/rojo, chimichurri, tzatziki.",
]

YOGURES_TEXT = [
    ("Mejor opción", "yogur natural sin azúcar, yogur griego natural, yogur skyr (alto en "
     "proteína), yogur proteico tipo Hacendado/Pascual sin azúcar, kéfir natural."),
    ("Evitar", "yogures de sabores, edulcorados con azúcar añadido, con frutas en almíbar o "
     "con cereales tipo «de postre»."),
]

QUESOS_TEXT = [
    ("Diarios", "queso fresco batido 0%, queso cottage, requesón, queso de Burgos light, "
     "queso fresco bajo en grasa, queso havarti light, queso de untar 0%."),
    ("Ocasionales (1-2 veces/semana)", "mozzarella de búfala, queso feta, queso de cabra "
     "fresco, parmesano rallado (en pequeñas cantidades para dar sabor)."),
    ("Evitar/limitar", "quesos curados muy grasos, quesos azules, quesos cremosos tipo "
     "brie/camembert en grandes cantidades."),
]

RECOMENDACIONES = [
    ("Agua", "2-3 L al día."),
    ("Días de descanso", "realizar cardio y tomar batido post entreno (opciones del post entreno)."),
    ("Cocciones recomendadas", "vapor, plancha, horno, freidora de aire. Aceite de oliva virgen extra siempre."),
    ("Saciedad extra", "proteína de soja aislada o caseína; espesantes como goma guar, arábiga o xantana."),
    ("Ansiedad", "gelatinas 0%, infusiones o aumentar ración de verdura."),
    ("Frutos secos", "sin sal, ni fritos ni tostados. Sus cremas 100% son válidas."),
]

SUPLEMENTACION_DEFAULT = [
    "Multivitamínico",
    "Omega 3",
    "Creatina monohidrato Creapure (incluida en intra y post entreno)",
    "Vitamina C 1000 mg después de entrenar",
    "Bisglicinato de magnesio 400 mg después de entrenar",
]


def _goal_label(goal: str | None) -> str:
    return {"fat_loss": "Pérdida de grasa", "muscle_gain": "Ganancia muscular",
            "recomp": "Recomposición", "maintenance": "Mantenimiento",
            "injury_recovery": "Recuperación de lesión"}.get(goal or "", "Plan personalizado")


def _objetivo_pairs(goal: str | None) -> list[tuple[str, str]]:
    """OBJETIVOS como el ejemplo: dos líneas con etiqueta en negrita vino
    ("Antropométrico: …" / "Nutricional: …"). Los cinco objetivos del sistema
    tienen texto propio: maintenance/injury_recovery imprimían el relleno
    genérico "Según objetivo." en el PDF del cliente (auditoría)."""
    anthro = {
        "fat_loss": "Déficit.",
        "muscle_gain": "Superávit.",
        "recomp": "Mantenimiento / recomposición.",
        "maintenance": "Mantenimiento del peso.",
        "injury_recovery": "Mantenimiento durante la recuperación.",
    }.get(goal or "", "Según objetivo.")
    nutri = {
        "fat_loss": "organizar y planificar la alimentación diaria, manteniendo proteína "
                    "para preservar masa muscular.",
        "muscle_gain": "organizar y planificar la alimentación diaria, aportando energía y "
                       "proteína suficientes para ganar masa muscular.",
        "recomp": "organizar y planificar la alimentación diaria, con proteína alta para "
                  "perder grasa y ganar o mantener músculo.",
        "maintenance": "consolidar hábitos y sostener el peso con una alimentación "
                       "equilibrada, con proteína suficiente para conservar masa muscular.",
        "injury_recovery": "apoyar la recuperación con energía en mantenimiento y proteína "
                           "alta para minimizar la pérdida de masa muscular.",
    }.get(goal or "", "organizar y planificar la alimentación diaria según tu objetivo.")
    return [("Antropométrico", anthro), ("Nutricional", nutri)]


def _n(x) -> str:
    """Entero en formato es-ES ("2.150 kcal", no "2150"). El importador del
    Word vuelve a leerlo sin problema: `_num` quita el punto de millar."""
    try:
        return f"{_rhu(float(x)):,}".replace(",", ".")
    except (TypeError, ValueError):
        return "—"


def _nota(doc: Document, texto: str) -> None:
    """Línea guía bajo una barra de sección: qué es esto y cómo se usa. Es lo
    que convierte una tabla en una instrucción — el cliente leía cifras sin
    saber qué hacer con ellas."""
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    run = par.add_run(texto)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = _hex("#6B6B76")
    _keep_with_next(par)


def _indice(doc: Document, secciones: list[str]) -> None:
    """Mapa del documento en la primera página: el cliente sabe de un vistazo
    qué tiene delante y deja de leerlo de cabo a rabo para encontrar una cosa."""
    if not secciones:
        return
    section_bar(doc, "En este documento", BLUE)
    info_box(doc, [" · ".join(secciones)], fill=CREAM, cant_split=True)


def _macro_lines(macros: dict, kcal: float) -> list[str]:
    """Celda del reparto: gramos arriba y su peso en kcal debajo. El reparto
    deja de ser tres cifras sueltas y se lee como lo que es."""
    ch = float(macros.get("carbs_g") or 0)
    pr = float(macros.get("protein_g") or 0)
    gr = float(macros.get("fat_g") or 0)
    linea = f"CH {_rhu(ch)} g · P {_rhu(pr)} g · G {_rhu(gr)} g"
    total = 4 * ch + 4 * pr + 9 * gr
    if total <= 0:
        return [linea]
    p_ch = _rhu(4 * ch / total * 100)
    p_pr = _rhu(4 * pr / total * 100)
    p_gr = max(0, 100 - p_ch - p_pr)
    return [linea, f"{p_ch}% · {p_pr}% · {p_gr}% de tus calorías"]


def _dias_semana(n: int) -> str:
    return "1 día/semana" if n == 1 else f"{n} días/semana"


def _indice_nutricion(nutrition: dict, include_training: bool,
                      blocked: list[str] | None = None,
                      diet_pattern: str | None = None) -> list[str]:
    """Qué secciones lleva DE VERDAD este documento (no un índice de plantilla
    que promete cosas que no están)."""
    partes = ["Objetivos"]
    # Mismas condiciones EXACTAS que gobiernan el pintado: un índice que
    # promete algo que luego no está es peor que no tener índice.
    bl = blocked or []
    razon = (nutrition.get("rationale") or "").strip()
    if razon and not _food_blocked(razon, bl, diet_pattern):
        partes.append("Por qué este enfoque")
    partes.append("Tus cifras del día")
    if nutrition.get("meals"):
        partes.append("Estructura diaria")
    partes += ["Alimentos por grupos", "Tus comidas", "Dieta semanal"]
    reglas = [r for r in (nutrition.get("flexibility_rules") or []) if str(r).strip()
              and not _food_blocked(str(r), bl, diet_pattern)]
    refeed = (nutrition.get("refeed_or_break") or "").strip()
    if reglas or (refeed and not _food_blocked(refeed, bl, diet_pattern)):
        partes.append("Margen de maniobra")
    if nutrition.get("supplements"):
        partes.append("Suplementación")
    if include_training:
        partes.append("Entrenamiento")
    return partes


def _origen_calorias(doc: Document, nutrition: dict) -> None:
    """De dónde salen las calorías: gasto estimado → ajuste → objetivo. Sin
    esto la cifra parece arbitraria; con esto el cliente entiende que es SUYA.
    Solo imprime números ya calculados por el backend: aquí no se calcula nada
    (la única operación es la resta que ya hace `_ajuste_text`)."""
    tdee = nutrition.get("tdee_kcal") or 0
    target = nutrition.get("target_kcal") or 0
    if not tdee or not target:
        return
    delta = _rhu(float(target) - float(tdee))
    ajuste = ("no aplicamos ajuste" if delta == 0
              else f"le sumamos {_n(delta)} kcal" if delta > 0
              else f"le restamos {_n(-delta)} kcal")
    info_box(doc, [
        ("De dónde sale tu cifra",
         f"tu gasto diario estimado es de ≈ {_n(tdee)} kcal; para tu objetivo "
         f"{ajuste}, y así llegamos a las {_n(target)} kcal de arriba."),
    ], fill=CREAM, label_color=WINE, cant_split=True)


def _title(doc: Document, text: str, sub: str | None = None,
           meta: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = _hex(WINE)
    if sub:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = ps.add_run(sub)
        rs.font.size = Pt(16)
        rs.font.bold = True
        rs.font.color.rgb = _hex("#1A1A1A")
    if meta:
        # De qué mes es este documento y para qué objetivo: el cliente acumula
        # PDF y necesita saber cuál mira sin abrir la cabecera.
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(2)
        rm = pm.add_run(meta.upper())
        rm.font.size = Pt(8.5)
        rm.font.bold = True
        rm.font.color.rgb = _hex(BLUE)


def _norm(s: str) -> str:
    s = (s or "").strip().lower()
    return "".join(c for c in unicodedata.normalize("NFD", s)
                   if unicodedata.category(c) != "Mn")


def _food_blocked(food: str, blocked: list[str], diet_pattern: str | None = None) -> bool:
    """¿Este alimento choca con una alergia/aversión o con el patrón dietético?

    Delega en el MISMO motor que veta el plan (guardrails.food_allergen, con sus
    sinónimos: "gluten"→pan/pasta/trigo…, "lactosa"→leche/yogur/queso…) en vez de
    comparar palabra a palabra: la anamnesis recoge las alergias por CATEGORÍA
    ("gluten", "frutos secos") y el filtro literal no casaba con "cebada" ni con
    "almendras", así que el celíaco veía pan en su documento (auditoría)."""
    from app.services import guardrails as gr

    if blocked and gr.food_allergen(food, blocked) is not None:
        return True
    if diet_pattern:
        forbidden = gr._DIET_PATTERN_FORBIDDEN.get(
            gr._norm_food(diet_pattern).replace(" ", "_"))
        if forbidden and gr._match_term(forbidden, [_norm(food)]) is not None:
            return True
    return False


def _blocked_line(item, blocked: list[str], diet_pattern: str | None) -> bool:
    """¿Una línea de una tarjeta contiene algo bloqueado? Acepta tanto texto
    suelto ("Pan integral con…") como las tuplas (etiqueta, cuerpo) que usan
    las tarjetas de yogures/quesos/recomendaciones."""
    texto = " ".join(str(x) for x in item) if isinstance(item, (tuple, list)) else str(item)
    return _food_blocked(texto, blocked, diet_pattern)


def _food_group_lines(column: str, blocked: list[str],
                      diet_pattern: str | None = None) -> list[tuple[str, str]]:
    """Líneas de una columna de 'Alimentos por grupos': [(etiqueta, alimentos)],
    cada subgrupo en SU línea con la etiqueta en negrita (como la referencia),
    quitando SOLO los alimentos bloqueados y conservando etiquetas y alimentos
    contiguos (arregla el bug del filtro)."""
    lines: list[tuple[str, str]] = []
    for label, foods in FOOD_GROUPS[column]:
        kept = [f for f in foods if not _food_blocked(f, blocked, diet_pattern)]
        if not kept:
            continue
        body = ", ".join(kept)
        if not body.endswith("."):
            body += "."
        lines.append((label, body))
    foot = FOOD_GROUP_FOOTNOTE.get(column)
    if foot:
        lines.append(("", foot))
    return lines or [("", "—")]


def _ajuste_text(nutrition: dict, goal: str | None) -> str:
    """Celda 'Ajuste aplicado': el ajuste real sobre el TDEE estimado."""
    tdee = nutrition.get("tdee_kcal") or 0
    target = nutrition.get("target_kcal") or 0
    if not tdee:
        return _goal_label(goal)
    delta = _rhu(target - tdee)
    pct = _rhu(abs(delta) / tdee * 100)
    # La etiqueta la manda el SIGNO del delta real, no el objetivo: si el objetivo
    # es "ganancia" pero las kcal quedaron por debajo del TDEE (tras editar o por
    # el suelo calórico), decir "Superávit +-150" sería falso y contradictorio.
    if delta > 0:
        return f"Superávit de {_n(delta)} kcal ({pct}%)"
    if delta == 0:
        return "Mantenimiento · sin ajuste"
    return f"Déficit de {_n(-delta)} kcal ({pct}%)"




def _ingredients_str(opt: dict) -> str:
    """Ingredientes con gramos y MEDIDA CASERA cuando la hay ("Avena 60 g
    (6 cucharadas)"): la medida casera es lo que permite al cliente servirse sin
    báscula y la IA la genera siempre; antes se descartaba (auditoría)."""
    out = []
    for ing in opt.get("ingredients", []):
        g = ing.get("grams")
        food = ing.get("food", "")
        casera = (ing.get("household") or "").strip()
        txt = f"{food} {round(g)} g" if g else food
        if casera:
            txt += f" ({casera})"
        out.append(txt)
    return ", ".join(out)


def _prep_str(opt: dict) -> str:
    """Preparación de la opción ("Mezcla la avena con el yogur…", 10 min): sin
    esto el cliente tiene la lista de la compra pero no la receta."""
    prep = (opt.get("prep") or "").strip()
    if not prep:
        return ""
    minutos = opt.get("prep_minutes")
    return f"{prep} ({minutos} min)" if minutos else prep


def generate_plan_doc(
    *, brand: DocBrand, client_name: str, month_index: int, goal_type: str | None,
    diet_mode: str | None, nutrition: dict, training: dict, education: dict,
    exercise_names: dict | None = None,
    food_allergies: list[str] | None = None, food_dislikes: list[str] | None = None,
    include_training: bool = False, include_nutrition: bool = True,
    diet_pattern: str | None = None,
) -> bytes:
    # El documento lleva lo que el cliente tenga contratado: dieta, entreno o las
    # dos cosas. include_nutrition=False es el plan `train` (documento SOLO de
    # entrenamiento, sin un "PLAN NUTRICIONAL" lleno de ceros).
    exercise_names = exercise_names or {}
    # Términos CRUDOS (como los escribió el coach/la anamnesis): el motor de
    # guardrails ya normaliza y expande sinónimos.
    blocked = [x for x in (food_allergies or []) + (food_dislikes or []) if x]

    # Ninguna toma sin contenido en el PDF: los planes antiguos (guardados antes
    # del relleno automático) reciben aquí sus 3 opciones por defecto escaladas
    # a los macros de la toma — el cliente nunca lee una "toma libre". Sobre una
    # copia: el dict del caller (posible fila de BD en sesión) no se muta.
    import copy as _copy

    from app.services.meal_fallback import ensure_bank_slots

    nutrition = _copy.deepcopy(nutrition)
    if include_nutrition:
        # El patrón dietético también aquí: sin él, las opciones de relleno del
        # PDF colaban huevo y pavo en el plan de un vegano (auditoría).
        ensure_bank_slots(nutrition, allergies=food_allergies or [],
                          dislikes=food_dislikes or [], diet_pattern=diet_pattern)

    doc = init_document(brand)
    # El ejemplo usa Calibri (en el contenedor se sustituye por Carlito, idéntico).
    for _sname in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            doc.styles[_sname].font.name = "Calibri"
        except Exception:
            pass
    # Cabecera de la REFERENCIA (logo + "PLAN NUTRICIONAL | Cliente" y año) en
    # todas las páginas; sin banda de fondo ni portada — el documento empieza
    # directamente con el contenido, como la copia del coach.
    from datetime import date as _date

    _cabecera = ("PLAN DE DIETA Y ENTRENAMIENTO" if (include_nutrition and include_training)
                 else "PLAN NUTRICIONAL" if include_nutrition
                 else "PLAN DE ENTRENAMIENTO")
    # Cabecera con el MES del plan y la fecha de generación: el cliente acumula
    # varios PDF y sin esto no sabía cuál era el vigente. Pie desde la MARCA
    # (antes iba el nombre del coach hardcodeado: cambiar la marca dejaba el
    # documento desincronizado con el resto de comunicaciones).
    _pie = brand.name + (f" · {brand.tagline}" if brand.tagline else "")
    setup_reference_pages(
        doc, logo_path=str(ASSETS / "dq_logo.png"),
        right_title=f"{_cabecera} | {client_name}",
        right_sub=f"Mes {month_index} · {_date.today().strftime('%d/%m/%Y')}",
        footer_text=_pie,
    )

    if include_nutrition:
        # ======================= NUTRICIÓN =======================
        _title(doc, "PLAN NUTRICIONAL", client_name,
               meta=f"Mes {month_index} · {_goal_label(goal_type)}")
        macros = nutrition.get("macros", {})

        _indice(doc, _indice_nutricion(nutrition, include_training,
                                       blocked, diet_pattern))

        section_bar(doc, "Objetivos", WINE)
        info_box(doc, _objetivo_pairs(goal_type), fill=CREAM, label_color=WINE)

        # POR QUÉ este enfoque: lo escribe el coach (o la IA bajo su revisión),
        # se guarda en el plan y NUNCA llegaba al documento del cliente — que
        # recibía las cifras sin el criterio que las justifica.
        razon = (nutrition.get("rationale") or "").strip()
        # Si el argumentario nombra algo que el cliente NO puede comer, no se
        # imprime: el resto del documento ya filtra sus alimentos y dejar esta
        # caja fuera del filtro reabría el agujero del celíaco viendo pan.
        if razon and not _food_blocked(razon, blocked, diet_pattern):
            section_bar(doc, "Por qué este enfoque", GOLD)
            info_box(doc, [razon], fill=CREAM, cant_split=True)

        section_bar(doc, "Resumen energético diario", BLUE)
        _nota(doc, "Estas son tus cifras del día. Los gramos de cada toma ya salen "
                   "de aquí: si cumples las tomas, el día cuadra solo.")
        clean_table(
            doc, ["Calorías", "Reparto de macros", "Ajuste aplicado"],
            [[f"≈ {_n(nutrition.get('target_kcal', 0))} kcal",
              _macro_lines(macros, nutrition.get("target_kcal", 0)),
              _ajuste_text(nutrition, goal_type)]],
            brand, header_color=WINE, header_text_color="FFFFFF",
            col_widths=[2400, 4226, 2400],
        )
        _origen_calorias(doc, nutrition)

        meals = nutrition.get("meals", [])

        # Cambios aplicados en la última adaptación (revisión quincenal): el cliente
        # ve QUÉ cambió, DÓNDE y POR QUÉ directamente en su PDF.
        aa = nutrition.get("applied_adjustments") or {}
        aa_items = aa.get("items") or []
        if aa_items:
            section_bar(doc, f"Cambios de tu plan · revisión #{aa.get('period_index', '')}", GOLD)
            rows = [[
                (it.get("area") or "").capitalize(),
                it.get("detail") or it.get("change") or "",
                it.get("reason") or "",
            ] for it in aa_items]
            # "Qué cambia"/"Por qué" son texto libre (IA/coach): pueden ser largos,
            # así que las filas se parten y la tabla pagina con cabecera repetida.
            clean_table(doc, ["Área", "Qué cambia", "Por qué"], rows, brand,
                        header_color=WINE, header_text_color="FFFFFF",
                        col_widths=[1400, 3800, 3826],
                        cant_split_rows=False, keep_together=False)

        if meals:
            section_bar(doc, "Estructura diaria", GOLD)
            _nota(doc, "Tu reparto del día. Las horas son orientativas: respeta el "
                       "número de tomas y su contenido, no el reloj.")
            rows = [[m.get("time", ""), m.get("name", f"Comida {m.get('slot')}"),
                     (str(m.get("strategy") or "").strip()
                      or _estrategia(m.get("name", "")))] for m in meals]
            clean_table(doc, ["Hora", "Toma", "Estrategia"], rows, brand,
                        header_color=WINE, header_text_color="FFFFFF",
                        col_widths=[1500, 3000, 4526], keep_together=False)

        # Alimentos por grupos (plantilla, filtrada con precisión por alergias).
        # Es UNA sola fila con listas largas: puede ser más alta que la página, así
        # que la fila debe poder partirse (cant_split_rows=False) y la tabla paginar
        # repitiendo la cabecera (keep_together=False) para no recortar alimentos.
        section_bar(doc, "Alimentos por grupos", WINE)
        _nota(doc, "Tu despensa: de aquí sale todo lo que comes. Ya está filtrada "
                   "con tus alergias, aversiones y tu patrón de alimentación.")
        names = list(FOOD_GROUPS.keys())
        clean_table(
            doc, names, [[_food_group_lines(n, blocked, diet_pattern) for n in names]],
            brand, header_colors=FOOD_GROUP_COLORS, header_text_color="FFFFFF",
            cant_split_rows=False, keep_together=False,
        )

        # El plato saludable (plantilla + foto)
        section_bar(doc, "El plato saludable", BLUE)
        _nota(doc, "La regla para el día que comes fuera y no puedes pesar nada.")
        # La foto del plato va DENTRO de la caja y la caja entera es indivisible
        # (cant_split): si no cabe, la tarjeta completa salta a la página siguiente
        # con su barra — la foto nunca queda sola en un fragmento de caja.
        info_box(doc, PLATO_TEXT, fill=CREAM, label_color=WINE,
                 cant_split=True, image_path=str(ASSETS / "plate.png"))

        # Comidas detalladas (flexible) — como el ejemplo: comida/cena con sistema de
        # equivalencias por grupos; el resto, 3 opciones numeradas en prosa (sin kcal).
        # Comidas detalladas: cada comida = barra + CAJA CREMA con el contenido dentro
        # (como el ejemplo). Comida/cena en equivalencias; resto, 3 opciones numeradas.
        bank = nutrition.get("meal_bank") or {}
        # El formato lo decide el banco PERSISTIDO (bank["mode"]); diet_mode del
        # cliente es solo fallback: si el coach cambia diet_mode sin regenerar, el
        # PDF sigue mostrando el menú que existe (no una sección vacía/equivocada).
        diet_mode = bank.get("mode") or diet_mode
        if diet_mode == "strict" and bank.get("days"):
            # MENÚ CERRADO: el cliente necesita el detalle de CADA día — antes
            # solo se imprimía la rejilla de títulos ("Pollo con arroz") y no
            # había forma de saber cuánto pesar ni cómo cocinarlo (auditoría).
            nombres = {m.get("slot"): (m.get("name"), m.get("time")) for m in meals}
            for d in bank["days"]:
                dia = str(d.get("day", "")).capitalize()
                section_bar(doc, dia, WINE, size=10)
                cell = open_box(doc, CREAM, cant_split=True)
                first = True
                for entry in d.get("meals", []):
                    dish = entry.get("dish") or {}
                    nombre, hora = nombres.get(entry.get("slot"), (None, None))
                    etiqueta = nombre or f"Toma {entry.get('slot','')}"
                    if hora:
                        etiqueta += f" · {hora}"
                    p = cell.paragraphs[0] if first else cell.add_paragraph()
                    first = False
                    p.paragraph_format.space_after = Pt(4)
                    _keep_lines(p)
                    rl = p.add_run(f"{etiqueta}. ")
                    rl.font.bold = True
                    rl.font.color.rgb = _hex(WINE)
                    p.add_run(f"{dish.get('title','')} — {_ingredients_str(dish)}.")
                    prep = _prep_str(dish)
                    if prep:
                        pp = cell.add_paragraph()
                        pp.paragraph_format.space_after = Pt(4)
                        _keep_lines(pp)
                        rp = pp.add_run(prep)
                        rp.font.italic = True
                        rp.font.size = Pt(9)
            # Comida libre semanal pactada en la anamnesis: sin esta caja la
            # pauta se generaba y el cliente no la veía en ningún sitio.
            libre = (bank.get("free_meal_guidelines") or "").strip()
            if libre:
                section_bar(doc, "Tu comida libre semanal", GOLD)
                info_box(doc, [libre], fill=CREAM, cant_split=True)
        elif diet_mode != "strict" and meals:
            blocks = {s.get("slot"): s for s in bank.get("slots", [])}
            for m in meals:
                section_bar(doc, f"{m.get('name','Comida')} · {m.get('time','')}", WINE, size=10)
                sb = blocks.get(m.get("slot"), {})
                # Regla del diseño de referencia: NINGÚN corte visible. Las cajas de
                # opciones/toma libre (contenido acotado, ≤3 opciones) viajan ENTERAS
                # a la página siguiente si no caben. Las equivalencias (sin cota) sí
                # pueden fluir, pero cada grupo lleva keepLines: el corte cae ENTRE
                # grupos, nunca a mitad de una frase.
                is_equiv = bool(sb.get("fmt") == "equivalences" and sb.get("equivalences"))
                cell = open_box(doc, CREAM, cant_split=not is_equiv)
                if sb.get("fmt") == "equivalences" and sb.get("equivalences"):
                    # foto redonda flotante en la cena (como el ejemplo del coach)
                    img = str(ASSETS / "food_round.png") if "cena" in _norm(m.get("name", "")) else None
                    _render_equivalences(cell, sb["equivalences"], image_path=img)
                else:
                    first = True
                    for n, opt in enumerate(sb.get("options", [])[:3], start=1):
                        p = cell.paragraphs[0] if first else cell.add_paragraph()
                        first = False
                        p.paragraph_format.space_after = Pt(4)
                        _keep_lines(p)  # una opción nunca se parte entre páginas
                        rl = p.add_run(f"Opción {n}. ")
                        rl.font.bold = True
                        rl.font.color.rgb = _hex(WINE)
                        p.add_run(f"{opt.get('title','')} — {_ingredients_str(opt)}.")
                        prep = _prep_str(opt)
                        if prep:
                            pp = cell.add_paragraph()
                            pp.paragraph_format.space_after = Pt(4)
                            _keep_lines(pp)
                            rp = pp.add_run(prep)
                            rp.font.italic = True
                            rp.font.size = Pt(9)
                    if first:
                        # Toma añadida a mano (sin recetario aún): guía digna en vez
                        # de una caja vacía — sus macros están en Estructura diaria.
                        t = m.get("target") or {}
                        detail = ""
                        if t.get("kcal"):
                            detail = (f" (~{round(t['kcal'])} kcal · P {round(t.get('protein_g') or 0)} g · "
                                      f"CH {round(t.get('carbs_g') or 0)} g · G {round(t.get('fat_g') or 0)} g)")
                        cell.paragraphs[0].add_run(
                            "Toma libre: elige alimentos de los grupos de arriba que cuadren "
                            f"con los macros objetivo de esta comida{detail}."
                        )

        # Ejemplo de dieta semanal
        _weekly_section(doc, brand, diet_mode, nutrition, bank)

        # MARGEN DE MANIOBRA: las reglas de flexibilidad y el refeed/descanso se
        # pautan en el plan y se quedaban dentro del sistema. Son justo lo que
        # evita que el cliente abandone el primer día que se sale del guion.
        reglas = [r for r in (nutrition.get("flexibility_rules") or []) if str(r).strip()
                  and not _food_blocked(str(r), blocked, diet_pattern)]
        refeed = (nutrition.get("refeed_or_break") or "").strip()
        if refeed and _food_blocked(refeed, blocked, diet_pattern):
            refeed = ""
        if reglas or refeed:
            section_bar(doc, "Tu margen de maniobra", GOLD)
            _nota(doc, "Un plan que no se puede seguir no sirve. Esto es lo que "
                       "puedes mover sin salirte del objetivo.")
            lineas: list = [f"• {r}" for r in reglas]
            if refeed:
                lineas.append(("Recarga o descanso", refeed))
            info_box(doc, lineas, fill=CREAM, label_color=WINE, cant_split=True)

        # Tarjetas informativas de cierre: contenido FIJO y acotado (menos de media
        # página cada una) → cada tarjeta viaja ENTERA a la página siguiente si no
        # cabe. Regla del diseño de referencia: un título abre una tarjeta nueva y
        # una tarjeta jamás aparece partida con líneas sueltas en otra página.

        # Las secciones de plantilla también se FILTRAN por alergias/aversiones y
        # patrón dietético: un celíaco no puede recibir "ideas rápidas" con pan ni
        # un vegano una tarjeta de quesos (antes salían intactas — auditoría).
        # Si una tarjeta se queda sin contenido, no se pinta ni su barra de título.
        def _tarjeta(titulo: str, color, lineas: list, *, vineta: bool = False) -> None:
            utiles = [x for x in lineas if not _blocked_line(x, blocked, diet_pattern)]
            if not utiles:
                return
            section_bar(doc, titulo, color)
            info_box(doc, [f"• {x}" for x in utiles] if vineta else utiles,
                     fill=CREAM, cant_split=True)

        _tarjeta("Ideas rápidas de desayunos, snacks y meriendas", WINE,
                 list(IDEAS_RAPIDAS), vineta=True)
        _tarjeta("Salsas recomendables", BLUE, list(SALSAS_TEXT))
        _tarjeta("Yogures recomendables", BLUE, list(YOGURES_TEXT))
        _tarjeta("Quesos recomendables", BLUE, list(QUESOS_TEXT))
        _tarjeta("Recomendaciones generales", WINE, list(RECOMENDACIONES))

        # Suplementación: SOLO lo que el plan prescribe. Sin suplementos pautados
        # no se inventa un protocolo por defecto (antes se colaban 5 productos que
        # el coach no había aprobado — auditoría).
        supps = nutrition.get("supplements", [])
        if supps:
            section_bar(doc, "Suplementación recomendada", BLUE)
            info_box(doc, [f"{s.get('name','')} — {s.get('dose','')} ({s.get('timing','')})"
                           for s in supps], fill=CREAM, cant_split=True)


    if not include_training or not training:
        _education_section(doc, education, include_training=False)
        _contact_section(doc, brand)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ======================= ENTRENAMIENTO =======================
    if include_nutrition:
        doc.add_page_break()
    _title(doc, "PLAN DE ENTRENAMIENTO", client_name,
           meta=f"Mes {month_index} · {_goal_label(goal_type)}")

    # Cambios de la última adaptación también en el plan SOLO-ENTRENO (ahí el
    # sello vive en training_json; con dieta ya se imprimió en su sección).
    if not include_nutrition:
        aa_t = training.get("applied_adjustments") or {}
        aa_t_items = aa_t.get("items") or []
        if aa_t_items:
            section_bar(doc, f"Cambios de tu plan · revisión #{aa_t.get('period_index', '')}", GOLD)
            rows = [[
                (it.get("area") or "").capitalize(),
                it.get("detail") or it.get("change") or "",
                it.get("reason") or "",
            ] for it in aa_t_items]
            clean_table(doc, ["Área", "Qué cambia", "Por qué"], rows, brand,
                        header_color=WINE, header_text_color="FFFFFF",
                        col_widths=[1400, 3800, 3826],
                        cant_split_rows=False, keep_together=False)

    section_bar(doc, f"Estructura · {training.get('split_name','')}", BLUE)
    info_box(doc, [
        (_dias_semana(len(training.get("sessions", []))), training.get("split_rationale", "")),
    ])

    prog = training.get("weekly_progression", [])
    if prog:
        section_bar(doc, "Progresión semanal", WINE)
        _nota(doc, "El mes no se entrena igual las cuatro semanas: la carga sube y "
                   "el RIR baja. Esta tabla manda sobre la sensación del día.")
        rows = [[f"Sem {w.get('week')}", w.get("intent", ""), f"{w.get('load_pct','')}%",
                 f"RIR {w.get('rir_target','')}", w.get("volume_note", "")] for w in prog]
        clean_table(doc, ["Semana", "Enfoque", "Carga", "RIR", "Notas"], rows, brand,
                    header_color=WINE, header_text_color="FFFFFF",
                    col_widths=[1100, 1800, 1100, 1100, 3926], keep_together=False)

    for sess in training.get("sessions", []):
        section_bar(doc, f"{sess.get('day','')} · {sess.get('name','')}", WINE, size=10)
        # Calentamiento en caja opaca (legible aunque caiga sobre la banda)
        if sess.get("warmup"):
            info_box(doc, [("Calentamiento", sess["warmup"])])
        rows = []
        for ex in sess.get("exercises", []):
            name = exercise_names.get(ex.get("exercise_id"), f"Ejercicio #{ex.get('exercise_id','')}")
            cue = ex.get("technique_cue", "") or ""
            # La clave BIOMECÁNICA (el porqué del ejercicio) y el TEMPO se
            # prescriben en el plan y solo los veía el panel del coach: el
            # cliente no recibía ni el porqué ni la cadencia (auditoría 27-08).
            # Mismas etiquetas que re-lee word_import._parse_cue_cell.
            biome = (ex.get("biomech_cue") or "").strip()
            if biome:
                cue = f"{cue}\nClave biomecánica: {biome}" if cue else f"Clave biomecánica: {biome}"
            tempo = str(ex.get("tempo") or "").strip()
            if tempo:
                cue = f"{cue}\nTempo: {tempo}" if cue else f"Tempo: {tempo}"
            # Indicaciones personalizadas del coach: en la misma celda, en línea
            # aparte y con etiqueta, para que el cliente no se las salte.
            notes = (ex.get("coach_notes") or "").strip()
            if notes:
                cue = f"{cue}\nIndicación para ti: {notes}" if cue else f"Indicación para ti: {notes}"
            # CÓMO PROGRESAR: se prescribe por ejercicio y no se imprimía en
            # ningún sitio — sin ella el cliente no sabe cuándo subir el peso
            # y repite las mismas cargas un mes entero (auditoría).
            regla = (ex.get("progression_rule") or "").strip()
            if regla:
                cue = f"{cue}\nCómo progresar: {regla}" if cue else f"Cómo progresar: {regla}"
            rows.append([
                name, f"{ex.get('sets','')}×{ex.get('rep_range','')}", f"RIR {ex.get('rir','')}",
                f"{ex.get('rest_sec','')}s", cue,
            ])
        if rows:
            clean_table(doc, ["Ejercicio", "Series", "RIR", "Descanso", "Clave técnica"], rows,
                        brand, header_color=WINE, header_text_color="FFFFFF",
                        col_widths=[2600, 1300, 1100, 1100, 2926], keep_together=False)
        if sess.get("cooldown"):
            info_box(doc, [("Vuelta a la calma", sess["cooldown"])])

    cardio = training.get("cardio") or {}
    if cardio.get("daily_steps") or cardio.get("sessions"):
        section_bar(doc, "Cardio y NEAT", BLUE)
        _nota(doc, "Los pasos diarios pesan más que el cardio: son la mayor parte "
                   "de lo que gastas fuera del gimnasio.")
        items = [("Pasos diarios objetivo",
                  f"{_n(cardio['daily_steps'])} pasos" if cardio.get("daily_steps") else "—")]
        for cs in cardio.get("sessions", []):
            items.append((cs.get("type", "").upper(),
                          f"{cs.get('minutes','')} min × {cs.get('times_per_week','')}/sem"
                          + (f" — {cs.get('notes')}" if cs.get("notes") else "")))
        info_box(doc, items)

    if training.get("deload_instructions"):
        section_bar(doc, "Semana de descarga (deload)", BLUE)
        _nota(doc, "No es perder una semana: es lo que permite que la siguiente "
                   "vuelvas más fuerte.")
        info_box(doc, [training["deload_instructions"]])

    _education_section(doc, education, include_training=True)
    _contact_section(doc, brand)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _contact_section(doc: Document, brand: DocBrand) -> None:
    """Cierre profesional del documento: cómo resolver dudas y dónde seguir el
    día a día. Antes el PDF terminaba en seco tras la suplementación/FAQ."""
    lineas: list = []
    if brand.contact_email:
        lineas.append(("Escríbeme", brand.contact_email))
    lineas.append(("Tu portal", "registra tu día a día (peso, comidas y entrenos) "
                   "desde el enlace de tu portal — es lo que me permite ajustarte "
                   "el plan en cada revisión."))
    section_bar(doc, "Cualquier duda, aquí me tienes", BLUE)
    info_box(doc, lineas, fill=CREAM, label_color=WINE, cant_split=True)


def _education_section(doc: Document, education: dict | None, *, include_training: bool) -> None:
    """Contenido educativo del plan (píldoras, técnica por patrones, FAQ).
    Se generaba con IA y se guardaba en education_json pero NUNCA llegaba al
    documento del cliente (auditoría): el parámetro entraba y no se usaba.
    Cada tarjeta viaja entera a la página siguiente si no cabe (cant_split)."""
    edu = education or {}
    pills = [p for p in (edu.get("pills") or []) if (p.get("for_client") or "").strip()]
    biomech = [b for b in (edu.get("biomech_by_pattern") or []) if b.get("pattern")]
    faq = [f for f in (edu.get("faq") or []) if (f.get("q") or "").strip()]
    if not (pills or biomech or faq):
        return

    if pills:
        section_bar(doc, "Aprende con tu plan", GOLD)
        info_box(doc, [(p.get("topic", ""), p["for_client"].strip()) for p in pills],
                 fill=CREAM, label_color=WINE, cant_split=True)

    # La técnica por patrones solo tiene sentido si el documento lleva entreno.
    if biomech and include_training:
        section_bar(doc, "Técnica: claves por patrón de movimiento", BLUE)
        items = []
        for b in biomech:
            cues = " · ".join(c.strip() for c in (b.get("cues") or []) if c and c.strip())
            why = (b.get("why") or "").strip()
            body = cues + (f" — {why}" if why else "")
            items.append((b.get("pattern", ""), body))
        info_box(doc, items, fill=CREAM, label_color=WINE, cant_split=True)

    if faq:
        section_bar(doc, "Preguntas frecuentes", WINE)
        info_box(doc, [(f.get("q", "").strip(), (f.get("a") or "").strip()) for f in faq],
                 fill=CREAM, label_color=WINE, cant_split=True)


def _render_equivalences(container, eq: dict, image_path: str | None = None) -> None:
    """Renderiza una comida en formato de equivalencias DENTRO de una caja (cell):
    línea intro + un párrafo por grupo con sus alimentos intercambiables. Si se
    pasa image_path, una foto redonda flota a la derecha con el texto alrededor."""
    intro = (eq.get("intro") or "").strip()
    p = container.paragraphs[0]  # reutiliza el primer párrafo (vacío) de la caja
    p.paragraph_format.space_after = Pt(4)
    if image_path and os.path.exists(image_path):
        try:
            float_image_right(p, image_path, Inches(1.5))
        except Exception:
            pass
    # Frase guía sin duplicados: si el intro del plan ya empieza por "Elige una
    # opción..." no se antepone otra vez, y se normaliza el cierre a un solo ":".
    base = "Elige una opción de cada grupo."
    if intro:
        low = intro.lower()
        if low.startswith("elige una opción") or low.startswith("elige una opcion"):
            txt = intro.rstrip(".:") + ":"
        else:
            txt = f"{base} {intro.rstrip('.:')}:"
    else:
        txt = base
    r = p.add_run(txt)
    r.font.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = _hex("#5A5A5A")
    for g in eq.get("groups", []):
        p = container.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _keep_lines(p)  # un grupo de equivalencias no se parte entre páginas
        rl = p.add_run(f"{g.get('name','')}: ")
        rl.font.bold = True
        rl.font.color.rgb = _hex(WINE)
        note = (g.get("note") or "").strip()
        items = [f"{it.get('food','')} ({it.get('amount','')})"
                 for it in g.get("items", []) if it.get("food")]
        body = note
        if items:
            body = (note + " " if note else "") + " o ".join(items)
        if body and not body.endswith("."):
            body += "."
        p.add_run(body)


def _estrategia(name: str) -> str:
    n = _norm(name)
    if "pre" in n:
        return "Pre-entreno: CH refinados, proteína magra, grasas reducidas."
    if "post" in n:
        return "Post-entreno: recuperación (proteína + CH)."
    if "cena" in n:
        return "Recuperación: integrales, proteína completa, grasas saludables."
    if "desayuno" in n:
        return "Ligero y de fácil digestión."
    if "media" in n or "merienda" in n:
        return "Sustancioso entre comidas principales."
    return "Comida equilibrada."


def _weekly_section(doc: Document, brand: DocBrand, diet_mode: str | None,
                    nutrition: dict, bank: dict) -> None:
    meals = nutrition.get("meals", [])
    if not meals:
        return

    headers = ["Toma"] + DAYS
    rows: list[list[str]] = []
    if diet_mode == "strict":
        days = bank.get("days", [])
        by_day = {_norm(d.get("day")): d for d in days}
        order = ["lunes", "martes", "miercoles", "jueves", "viernes", "sabado", "domingo"]
        for m in meals:
            slot = m.get("slot")
            cells = []
            for dslug in order:
                d = by_day.get(dslug)
                title = ""
                if d:
                    for meal in d.get("meals", []):
                        if meal.get("slot") == slot:
                            title = meal.get("dish", {}).get("title", "")
                cells.append(title)
            rows.append([m.get("name", f"Comida {slot}")] + cells)
    else:
        blocks = {s.get("slot"): s for s in bank.get("slots", [])}
        for m in meals:
            sb = blocks.get(m.get("slot"), {})
            wk = [x for x in (sb.get("weekly_examples") or []) if x]
            opts = sb.get("options", [])
            cells = []
            for di in range(7):
                if wk:
                    cells.append(wk[di % len(wk)])
                elif opts:
                    cells.append(opts[di % len(opts)].get("title", ""))
                else:
                    cells.append("")
            rows.append([m.get("name", f"Comida {m.get('slot')}")] + cells)

    # Sin contenido no se imprime NADA (ni la barra): una rejilla de celdas en
    # blanco bajo "Ejemplo de dieta semanal" lee como un error del sistema en
    # un documento pagado. Mismo criterio que las tarjetas (_tarjeta).
    rows = [r for r in rows if any(c.strip() for c in r[1:])]
    if rows:
        section_bar(doc, "Ejemplo de dieta semanal", WINE)
        # 8 columnas estrechas: fuente 8pt para que los nombres de plato no
        # desborden, y paginación con cabecera repetida (keep_together=False)
        # por si hay muchas tomas.
        clean_table(doc, headers, rows, brand, header_color=WINE, header_text_color="FFFFFF",
                    col_widths=[1500] + [1075] * 7, font_pt=8, keep_together=False)
