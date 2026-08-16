"""Pool de rutinas/planificaciones (plantillas) — lógica compartida.

Tres caminos alimentan el pool y los tres acaban en el MISMO shape que
`plans.training_json` (ejercicios resueltos a `exercise_id` de la biblioteca):

  1. SEMBRADAS de fábrica (`seeds/templates_data.py`): 20 rutinas por carpeta,
     escritas con nombres EXACTOS de la biblioteca y resueltas aquí.
  2. SUBIDAS (PDF/Word externos): la IA extrae la rutina y MAPEA cada ejercicio
     al nombre más cercano de la biblioteca (lista cerrada en el prompt, como en
     la generación de planes: la IA nunca inventa ejercicios); el backend
     resuelve nombre→id de forma determinista.
  3. CREADAS/EDITADAS a mano en el editor (ya llegan con exercise_id).

"Usar con un cliente" copia la plantilla como Plan BORRADOR (el coach revisa y
publica: mismo patrón de seguridad que la generación con IA). El documento de
una plantilla se renderiza con el dossier de la marca (mismo motor que los
planes de clientes) — así cualquier plan externo subido queda re-maquetado al
diseño de Professional.
"""
from __future__ import annotations

import difflib
import unicodedata

from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.models import Exercise, PlanTemplate

# Los TRES grandes grupos del negocio: es como piensa el coach y como pide el
# cliente ("quiero ganar músculo" / "quiero definir" / "quiero mantenerme").
# El frontend los lee de GET /api/templates/categories (una sola verdad).
CATEGORIES: list[dict] = [
    {"key": "masa", "label": "Masa muscular",
     "blurb": "Ganar músculo y fuerza: volumen, progresión y prioridades."},
    {"key": "definicion", "label": "Definición",
     "blurb": "Perder grasa manteniendo el músculo: déficit bien entrenado."},
    {"key": "mantenimiento", "label": "Mantenimiento y salud",
     "blurb": "Mantener resultados, empezar de cero o entrenar con molestias."},
]
CATEGORY_KEYS = {c["key"] for c in CATEGORIES}

# Carpetas ANTIGUAS del pool → grupo actual (fichas y plantillas de antes).
LEGACY_CATEGORIES = {
    "ganancia_muscular": "masa", "fuerza": "masa",
    "perdida_grasa": "definicion",
    "salud_espalda": "mantenimiento", "principiantes": "mantenimiento",
}


def normalize_category(key: str | None) -> str:
    k = (key or "").strip().lower()
    k = LEGACY_CATEGORIES.get(k, k)
    return k if k in CATEGORY_KEYS else "mantenimiento"


# Objetivo interno del motor por grupo (pre-rellena la ficha al usar la
# plantilla y decide el reparto de macros).
CATEGORY_GOAL = {
    "masa": "muscle_gain", "definicion": "fat_loss", "mantenimiento": "maintenance",
}

_DEFAULT_PROGRESSION = [
    {"week": 1, "intent": "Adaptación", "load_pct": 100, "rir_target": "2-3",
     "volume_note": "Aprende los pesos de referencia."},
    {"week": 2, "intent": "Progresión", "load_pct": 102.5, "rir_target": "2",
     "volume_note": "Sube carga si cierras todas las series."},
    {"week": 3, "intent": "Carga", "load_pct": 105, "rir_target": "1-2",
     "volume_note": "Semana fuerte: prioriza técnica."},
    {"week": 4, "intent": "Descarga", "load_pct": 90, "rir_target": "3-4",
     "volume_note": "Mitad de series: recupera para el mes siguiente."},
]


class TemplateError(ValueError):
    """Plantilla inválida (ejercicios sin resolver, estructura incoherente…)."""


def _norm(s: str) -> str:
    s = unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode()
    return " ".join(s.lower().split())


def _library_index(db: Session) -> tuple[dict[str, int], list[str], dict[str, str]]:
    """Índices nombre→id (canónicos y alias, normalizados) + lista de canónicos."""
    by_name: dict[str, int] = {}
    canon_by_norm: dict[str, str] = {}
    canonicals: list[str] = []
    for ex in db.scalars(select(Exercise)):
        n = _norm(ex.canonical_name)
        by_name[n] = ex.id
        canon_by_norm[n] = ex.canonical_name
        canonicals.append(n)
        for a in ex.aliases or []:
            by_name.setdefault(_norm(a), ex.id)
    return by_name, canonicals, canon_by_norm


def resolve_training(db: Session, routine: dict) -> dict:
    """Convierte una rutina con ejercicios POR NOMBRE al shape de
    `plans.training_json` (exercise_id resueltos). Lanza TemplateError con la
    lista de nombres irresolubles — nunca inventa un ejercicio."""
    by_name, canonicals, _ = _library_index(db)
    unresolved: list[str] = []
    sessions_out = []
    for sess in routine.get("sessions", []):
        exs = []
        for ex in sess.get("exercises", []):
            if ex.get("exercise_id"):
                exs.append(ex)
                continue
            name = ex.get("name") or ""
            key = _norm(name)
            ex_id = by_name.get(key)
            if ex_id is None:
                # Tolerancia mínima a variaciones tipográficas (0.9): un nombre
                # realmente distinto NO se sustituye en silencio.
                close = difflib.get_close_matches(key, canonicals, n=1, cutoff=0.9)
                ex_id = by_name.get(close[0]) if close else None
            if ex_id is None:
                unresolved.append(name)
                continue
            exs.append({
                "exercise_id": ex_id,
                "sets": int(ex.get("sets") or 3),
                "rep_range": str(ex.get("rep_range") or "8-10"),
                "rir": str(ex.get("rir") or "2"),
                "rest_sec": int(ex.get("rest_sec") or 120),
                "technique_cue": ex.get("technique_cue") or "",
            })
        if not exs:
            raise TemplateError(
                f"La sesión «{sess.get('name') or sess.get('day')}» queda sin ejercicios reconocibles")
        sessions_out.append({
            "day": sess.get("day") or "Lunes",
            "name": sess.get("name") or "Sesión",
            "warmup": sess.get("warmup") or "5-10 minutos de cardio suave y movilidad general.",
            "exercises": exs,
            "cooldown": sess.get("cooldown") or "Estiramiento global 5 minutos.",
        })
    if unresolved:
        raise TemplateError(
            "Ejercicios fuera de la biblioteca: " + ", ".join(sorted(set(unresolved))))
    if not sessions_out:
        raise TemplateError("La rutina no tiene sesiones")
    return {
        "split_name": routine.get("split_name") or "Rutina",
        "split_rationale": routine.get("split_rationale") or "",
        "sessions": sessions_out,
        "weekly_progression": routine.get("weekly_progression") or _DEFAULT_PROGRESSION,
        "cardio": routine.get("cardio") or {"daily_steps": 8000, "sessions": []},
        "deload_instructions": routine.get("deload_instructions")
        or "Semana 4: misma rutina con la mitad de series y RIR 3-4.",
    }


# ------------------------------------------- dieta de la plantilla ----------
# Cliente de REFERENCIA de cada grupo: solo sirve para dar a la plantilla unos
# macros de partida coherentes. Al asignarla a una persona, `materialize_for_client`
# recalcula sus targets REALES (metrics) y reescala la dieta — la plantilla nunca
# impone las calorías de otro.
_REF_CLIENT = {
    "masa":          {"sex": "male",   "weight": 75.0, "height": 176.0, "age": 30, "days": 4},
    "definicion":    {"sex": "female", "weight": 70.0, "height": 166.0, "age": 35, "days": 3},
    "mantenimiento": {"sex": "male",   "weight": 78.0, "height": 175.0, "age": 45, "days": 3},
}
_MEAL_NAMES = {
    3: [("Desayuno", "08:00"), ("Comida", "14:00"), ("Cena", "21:00")],
    4: [("Desayuno", "08:00"), ("Comida", "14:00"), ("Merienda", "18:00"), ("Cena", "21:30")],
    5: [("Desayuno", "08:00"), ("Media mañana", "11:00"), ("Comida", "14:00"),
        ("Merienda", "18:00"), ("Cena", "21:30")],
}
# Reparto de las kcal del día por número de comidas (suma 1).
_MEAL_SPLIT = {3: (0.30, 0.40, 0.30), 4: (0.25, 0.35, 0.15, 0.25),
               5: (0.22, 0.10, 0.33, 0.13, 0.22)}


def build_diet(category: str, meals_per_day: int = 4, *, diet_pattern: str | None = None,
               allergies: list[str] | None = None, focus: str | None = None) -> dict:
    """Dieta de referencia de una plantilla: totales del cliente tipo del grupo,
    reparto por comidas y banco de 3 opciones por toma. Se reescala a la persona
    al asignarla."""
    from app.services import metrics as M
    from app.services.meal_fallback import ensure_bank_slots
    from app.services.nutrition_scale import reconcile_nutrition

    ref = _REF_CLIENT.get(category, _REF_CLIENT["mantenimiento"])
    goal = CATEGORY_GOAL.get(category, "maintenance")
    e = M.energy_targets(ref["sex"], ref["weight"], ref["height"], ref["age"], goal,
                         ref["days"], daily_activity="light")
    m = M.macro_targets(ref["sex"], ref["weight"], goal, e.target_kcal, ref["days"],
                        tdee=e.tdee)
    n = max(3, min(5, int(meals_per_day or 4)))
    split = _MEAL_SPLIT[n]
    meals = []
    for i, ((name, hora), frac) in enumerate(zip(_MEAL_NAMES[n], split), start=1):
        meals.append({
            "slot": i, "name": name, "time": hora,
            "target": {
                "kcal": round(m.kcal * frac),
                "protein_g": round(m.protein_g * frac),
                "carbs_g": round(m.carbs_g * frac),
                "fat_g": round(m.fat_g * frac),
            },
        })
    nut = {
        "target_kcal": m.kcal, "bmr_kcal": round(e.bmr), "tdee_kcal": round(e.tdee),
        "macros": {"protein_g": m.protein_g, "carbs_g": m.carbs_g, "fat_g": m.fat_g},
        "meals": meals,
        "meal_bank": {"mode": "flexible_7", "slots": []},
        "rationale": (focus or "Estructura de referencia del grupo") +
                     "; las cantidades se ajustan a tus datos al asignarla.",
    }
    ensure_bank_slots(nut, allergies=allergies, dislikes=None, diet_pattern=diet_pattern)
    return reconcile_nutrition(nut, ref["weight"])


def materialize_for_client(db: Session, tpl: PlanTemplate, client) -> tuple[dict | None, dict | None]:
    """Adapta la plantilla a UNA persona antes de copiarla como plan.

    · Entrenamiento: se copia tal cual (ya está resuelto a la biblioteca).
    · Dieta: se REESCALA a las kcal y macros reales del cliente (metrics manda,
      nunca la plantilla) y el banco se rehace con SUS alergias, aversiones y
      patrón dietético. Sin datos suficientes en la ficha, se copia la dieta de
      referencia y el coach la ajusta.
    """
    from app.services import packages as pkgs

    training = tpl.training_json if pkgs.has_training(client.package_tier) else None
    if not pkgs.has_nutrition(client.package_tier):
        return training, None

    base = tpl.nutrition_json
    if not base:
        return training, None

    import copy as _copy

    from app.services import metrics as M
    from app.services.meal_fallback import ensure_bank_slots
    from app.services.nutrition_scale import reconcile_nutrition, rescale_nutrition

    nut = _copy.deepcopy(base)
    peso = client.current_weight_kg or client.start_weight_kg
    edad = None
    if client.birth_date:
        from datetime import date as _date

        hoy = _date.today()
        edad = hoy.year - client.birth_date.year - (
            (hoy.month, hoy.day) < (client.birth_date.month, client.birth_date.day))
    if peso and client.height_cm and edad and client.sex:
        goal = client.goal_type or CATEGORY_GOAL.get(tpl.category, "maintenance")
        dias = client.training_days or 3
        e = M.energy_targets(client.sex, peso, client.height_cm, edad, goal, dias,
                             body_fat_pct=client.body_fat_pct,
                             daily_activity=client.daily_activity_level,
                             level=client.level)
        m = M.macro_targets(client.sex, peso, goal, e.target_kcal, dias, tdee=e.tdee)
        rescale_nutrition(nut, base, m.kcal, m.protein_g, m.carbs_g, m.fat_g)
        nut["bmr_kcal"], nut["tdee_kcal"] = round(e.bmr), round(e.tdee)
        nut = reconcile_nutrition(nut, peso)
    # El banco se rehace SIEMPRE con las restricciones de esta persona.
    nut["meal_bank"] = {"mode": "flexible_7", "slots": []}
    ensure_bank_slots(nut, allergies=client.food_allergies, dislikes=client.food_dislikes,
                      diet_pattern=getattr(client, "diet_pattern", None))
    return training, nut


def template_document(db: Session, tpl: PlanTemplate, fmt: str = "pdf",
                      service: str = "full") -> tuple[bytes, str, str]:
    """Renderiza la plantilla con el dossier de la MARCA (mismo motor que los
    planes de clientes): cualquier plan externo subido sale re-maquetado.

    `service` decide QUÉ se entrega, igual que el servicio contratado:
    "nutri" solo la dieta, "train" solo el entrenamiento, "full" ambos.
    """
    from app.services.docs.pdf_convert import docx_bytes_to_pdf
    from app.services.docs.plan_doc import generate_plan_doc
    from app.services.plan_delivery import DOCX_MEDIA, doc_brand

    training = tpl.training_json or {}
    ex_ids = {e.get("exercise_id") for s in training.get("sessions", [])
              for e in s.get("exercises", []) if e.get("exercise_id")}
    names: dict[int, str] = {}
    if ex_ids:
        for ex in db.scalars(select(Exercise).where(Exercise.id.in_(ex_ids))):
            names[ex.id] = ex.canonical_name

    servicio = service if service in ("full", "nutri", "train") else "full"
    con_dieta = bool(tpl.nutrition_json) and servicio in ("full", "nutri")
    con_entreno = bool(training) and servicio in ("full", "train")
    data = generate_plan_doc(
        brand=doc_brand(db),
        client_name=tpl.title,
        month_index=1,
        goal_type=tpl.goal_type,
        diet_mode=None,
        nutrition=tpl.nutrition_json or {},
        training=training,
        education={},
        exercise_names=names,
        include_nutrition=con_dieta,
        include_training=con_entreno,
        package_tier=servicio,
    )
    ascii_t = unicodedata.normalize("NFKD", tpl.title).encode("ascii", "ignore").decode()
    safe = "".join(c if c.isalnum() else "_" for c in ascii_t).strip("_").lower() or "rutina"
    pref = {"nutri": "dieta", "train": "entreno", "full": "plan"}[servicio]
    if fmt == "docx":
        return data, DOCX_MEDIA, f"{pref}_{safe}.docx"
    try:
        pdf = docx_bytes_to_pdf(data)
        return pdf, "application/pdf", f"{pref}_{safe}.pdf"
    except Exception:
        return data, DOCX_MEDIA, f"{pref}_{safe}.docx"


# ----------------------------------------------------- importación con IA ----
class _ImpExercise(BaseModel):
    name: str  # EXACTO de la lista de la biblioteca inyectada en el prompt
    sets: int = Field(default=3, ge=1, le=10)
    rep_range: str = "8-10"
    rir: str = "2"
    rest_sec: int = Field(default=120, ge=15, le=600)
    technique_cue: str = ""


class _ImpSession(BaseModel):
    day: str = "Lunes"
    name: str = "Sesión"
    warmup: str = ""
    exercises: list[_ImpExercise] = Field(min_length=1)
    cooldown: str = ""


class TemplateImport(BaseModel):
    """Rutina extraída de un documento externo (solo ENTRENAMIENTO: la dieta de
    un documento ajeno no se importa — la nutrición la calcula el sistema con la
    anamnesis de cada cliente)."""

    title: str
    case_note: str | None = None
    level: str | None = None            # beginner|intermediate|advanced
    days_per_week: int | None = Field(default=None, ge=1, le=7)
    training_place: str | None = None   # gym|home
    split_name: str | None = None
    split_rationale: str | None = None
    sessions: list[_ImpSession] = Field(min_length=1)
    deload_instructions: str | None = None


_IMPORT_SYSTEM = """Eres el asistente del preparador de un centro de fitness. Te llega un \
documento EXTERNO con una rutina/planificación de entrenamiento y debes extraerla a JSON \
para incorporarla a la biblioteca del centro, que la re-maquetará con su propio diseño.

REGLAS:
- Extrae SOLO el entrenamiento (sesiones, ejercicios, series, repeticiones, RIR/intensidad, \
descansos). Si el documento trae dieta, IGNÓRALA (la nutrición se genera aparte para cada \
cliente).
- Cada ejercicio debe mapearse al nombre MÁS CERCANO de la BIBLIOTECA que se te da (lista \
cerrada, cópialo EXACTO, tildes incluidas). Nunca inventes nombres fuera de la lista; si un \
ejercicio no tiene equivalente razonable, OMÍTELO.
- Si el documento no da algún dato (RIR, descanso…), usa valores prudentes por defecto \
(RIR "2", 120 s) en vez de inventar precisión.
- Textos en castellano, tono profesional, sin emojis.
- "title": nombre corto de la rutina. "case_note": para quién es (si el documento lo dice)."""


def _import_user_prompt(db: Session, extra: str = "") -> str:
    _, _, canon_by_norm = _library_index(db)
    nombres = sorted(canon_by_norm.values())
    return (
        "BIBLIOTECA DE EJERCICIOS (usa estos nombres EXACTOS):\n- "
        + "\n- ".join(nombres)
        + "\n\nExtrae la rutina del documento a JSON según el esquema." + extra
    )


def import_template_from_file(
    db: Session, *, filename: str, raw: bytes, category: str, ai=None,
) -> PlanTemplate:
    """PDF/Word externo → plantilla del pool (la IA extrae y mapea; el backend
    resuelve nombre→id de forma determinista). Devuelve la fila SIN commitear."""
    from app.config import settings
    from app.services.ai.client import AIClient

    ai = ai or AIClient()
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        imp = ai.read_pdf_json(
            model=settings.model_heavy, system=_IMPORT_SYSTEM,
            user=_import_user_prompt(db), pdf_bytes=raw, schema=TemplateImport,
        )
    elif lower.endswith(".docx"):
        import io

        from docx import Document

        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)[:60000]
        imp = ai.generate_json(
            model=settings.model_heavy, system=_IMPORT_SYSTEM,
            user=_import_user_prompt(db, f"\n\nDOCUMENTO (texto extraído):\n{text}"),
            schema=TemplateImport,
        )
    else:
        raise TemplateError("Formato no soportado: sube un PDF o un Word (.docx)")

    routine = imp.model_dump()
    training = resolve_training(db, routine)
    cat = normalize_category(category)
    return PlanTemplate(
        category=cat,
        title=imp.title.strip()[:160] or "Rutina importada",
        case_note=imp.case_note,
        goal_type=CATEGORY_GOAL.get(cat),
        level=imp.level if imp.level in ("beginner", "intermediate", "advanced") else None,
        days_per_week=imp.days_per_week or len(training["sessions"]),
        training_place=imp.training_place if imp.training_place in ("gym", "home") else "gym",
        training_json=training,
        # La dieta del documento externo NO se importa, pero la plantilla sí
        # recibe la dieta de referencia del grupo: así también sirve al pack.
        nutrition_json=build_diet(cat, 4),
        meals_per_day=4,
        source="upload",
    )


def seed_plan_templates(db: Session) -> int:
    """Siembra el pool de fábrica (insert-por-título dentro de cada carpeta, como
    la biblioteca de ejercicios: re-ejecutar no duplica ni pisa ediciones).

    Si el catálogo de fábrica cambia (una rutina se retira o se renombra), las
    filas `seed` que ya no están en él se RETIRAN — salvo que el coach las haya
    editado, en cuyo caso pasan a ser suyas (`manual`) y se conservan. Sin esto,
    actualizar el catálogo dejaría el pool con las viejas y las nuevas."""
    try:
        from app.seeds.templates_data import TEMPLATES
    except Exception:
        return 0
    catalogo = {(normalize_category(e["category"]), e["title"]) for e in TEMPLATES}
    for t in list(db.scalars(select(PlanTemplate).where(PlanTemplate.source == "seed"))):
        if (t.category, t.title) in catalogo:
            continue
        # ¿La tocó el coach? updated_at solo cambia con un PATCH real.
        editada = t.updated_at is not None and t.created_at is not None and (
            t.updated_at - t.created_at).total_seconds() > 1
        if editada:
            t.source = "manual"   # es suya: se queda como rutina propia
        else:
            db.delete(t)          # de fábrica y retirada del catálogo
    db.flush()

    # Plantillas ya sembradas: se COMPLETA lo que falte del eje de dieta (una
    # instancia que venía de antes del eje tenía las tomas a NULL y su listado
    # de "solo dieta" salía mudo). Nunca se pisa un valor existente.
    por_titulo = {(t.category, t.title): t for t in db.scalars(
        select(PlanTemplate).where(PlanTemplate.source == "seed"))}
    catalogo_por_titulo = {(normalize_category(e["category"]), e["title"]): e
                           for e in TEMPLATES}
    rellenadas = 0
    for clave, fila in por_titulo.items():
        entrada = catalogo_por_titulo.get(clave)
        if entrada is None or fila.meals_per_day is not None:
            continue
        fila.meals_per_day = entrada.get("meals_per_day") or 4
        fila.diet_pattern = fila.diet_pattern or entrada.get("diet_pattern")
        fila.diet_focus = fila.diet_focus or entrada.get("diet_focus")
        if not fila.nutrition_json:
            fila.nutrition_json = build_diet(
                fila.category, fila.meals_per_day, diet_pattern=fila.diet_pattern,
                allergies=entrada.get("diet_allergies"), focus=fila.diet_focus)
        rellenadas += 1
    if rellenadas:
        db.commit()

    existing = set(por_titulo)
    added = 0
    for entry in TEMPLATES:
        key = (normalize_category(entry["category"]), entry["title"])
        if key in existing:
            continue
        try:
            training = resolve_training(db, entry)
        except TemplateError:
            # Un nombre irresoluble NO tumba el arranque: se omite esa plantilla.
            continue
        cat = normalize_category(entry["category"])
        comidas = entry.get("meals_per_day") or 4
        patron = entry.get("diet_pattern")
        db.add(PlanTemplate(
            category=cat, title=entry["title"],
            case_note=entry.get("case"), goal_type=CATEGORY_GOAL.get(cat),
            level=entry.get("level"), days_per_week=entry.get("days_per_week"),
            training_place=entry.get("place") or "gym",
            training_json=training,
            # Cada plantilla trae TAMBIÉN su dieta: así sirve para los tres
            # servicios (dieta, entrenamiento o pack) sin duplicar el catálogo.
            # El CASO decide las tomas, el patrón y las exclusiones.
            nutrition_json=build_diet(cat, comidas, diet_pattern=patron,
                                      allergies=entry.get("diet_allergies"),
                                      focus=entry.get("diet_focus")),
            meals_per_day=comidas, diet_pattern=patron,
            diet_focus=entry.get("diet_focus"),
            source="seed",
        ))
        added += 1
    if added:
        db.commit()
    return added


# ================================ RECOMENDADOR ================================
# Elige del pool las plantillas que MEJOR encajan con esta persona y explica el
# porqué en una frase. Es DETERMINISTA (reglas, no IA): el coach entiende de
# dónde sale cada sugerencia y puede fiarse o ignorarla — siempre le quedan las
# otras dos vías: buscarla él o crearla.

_LEVEL_ORDER = {"beginner": 0, "intermediate": 1, "advanced": 2}
_GOAL_WORD = {"masa": "ganar músculo", "definicion": "definición",
              "mantenimiento": "mantenerse"}


def _client_age(client) -> int | None:
    if not getattr(client, "birth_date", None):
        return None
    from datetime import date as _date

    hoy = _date.today()
    return hoy.year - client.birth_date.year - (
        (hoy.month, hoy.day) < (client.birth_date.month, client.birth_date.day))


def _goal_category(goal: str | None) -> str:
    return {"muscle_gain": "masa", "fat_loss": "definicion",
            "recomp": "definicion", "maintenance": "mantenimiento",
            "injury_recovery": "mantenimiento"}.get(goal or "", "mantenimiento")


_PATRON_PALABRA = {
    "vegano": "que es vegana", "vegetariano": "que es vegetariana",
    "pescetariano": "sin carne", "sin_cerdo": "sin cerdo",
    "halal": "halal", "kosher": "kosher",
}


def portal_signals(db: Session, client) -> dict:
    """Lo que el propio cliente ha registrado en el portal: evolución de peso y
    constancia. Es la segunda fuente de recomendaciones (la primera es su
    anamnesis) y también es DETERMINISTA: reglas simples y explicables.

    Devuelve `{}` si aún no hay datos suficientes — sin inventar tendencias.
    """
    from app.models import DailyLog, Period, WorkoutLog

    periodos = list(db.scalars(
        select(Period).where(Period.client_id == client.id)
        .order_by(Period.period_index.desc()).limit(2)))
    if not periodos:
        return {}
    ids = [p.id for p in periodos]
    logs = list(db.scalars(select(DailyLog).where(DailyLog.period_id.in_(ids))
                           .order_by(DailyLog.log_date)))
    if len(logs) < 5:
        return {}

    from app.services import metrics as M

    señales: dict = {"dias_registrados": len(logs)}

    pesos = [(x.log_date, x.weight_kg) for x in logs if x.weight_kg]
    if len(pesos) >= 5:
        t = M.weight_trend(pesos)
        ritmo = t.weekly_rate_kg or 0.0
        señales["kg_semana"] = round(ritmo, 2)
        peso_ref = pesos[-1][1] or 0
        pct = (ritmo / peso_ref * 100) if peso_ref else 0.0
        if abs(pct) < 0.2:
            señales["estancado"] = True
        elif pct <= -0.2:
            señales["bajando"] = True
        else:
            señales["subiendo"] = True

    dieta = [x.diet_adherence for x in logs if x.diet_adherence]
    if len(dieta) >= 5:
        bien = sum(1 for d in dieta if d == "yes")
        señales["adherencia_dieta"] = round(bien / len(dieta), 2)
        if bien / len(dieta) < 0.6:
            señales["adherencia_baja"] = True

    # Días transcurridos del seguimiento (no la duración nominal del período,
    # que sin ciclo quincenal es de un año): lo que se mide es la constancia.
    hoy = logs[-1].log_date
    dias_totales = sum(
        (min(p.ends_on, hoy) - p.starts_on).days + 1
        for p in periodos if p.starts_on and p.ends_on and p.starts_on <= hoy
    )
    if dias_totales >= 10 and len(logs) / dias_totales < 0.5:
        señales["registra_poco"] = True

    from app.services import packages as pkgs

    if pkgs.has_training(getattr(client, "package_tier", None)):
        entrenos = db.scalar(select(func.count(func.distinct(WorkoutLog.daily_log_id)))
                             .where(WorkoutLog.daily_log_id.in_([x.id for x in logs])))
        previstos = (client.training_days or 3) * max(1, len(logs) // 7)
        if previstos and (entrenos or 0) < previstos * 0.5:
            señales["entrena_menos"] = True
    return señales


def recommend_templates(db: Session, client, limit: int = 5) -> list[dict]:
    """Las `limit` plantillas más afines a este cliente, cada una con su motivo.

    Puntúa lo que de verdad decide el encaje: objetivo (el grupo), días
    disponibles, dónde entrena, experiencia, edad y las molestias declaradas.
    Si el cliente aún no tiene anamnesis, devuelve lo más genérico de su grupo.
    """
    from app.services import packages as pkgs

    tier = pkgs.normalize(getattr(client, "package_tier", None))
    cat = _goal_category(getattr(client, "goal_type", None))
    dias = getattr(client, "training_days", None)
    lugar = getattr(client, "training_place", None)
    nivel = getattr(client, "level", None)
    edad = _client_age(client)
    molestias = _norm(" ".join(filter(None, [
        getattr(client, "injuries_notes", None) or "",
        getattr(client, "medical_notes", None) or "",
    ])))
    patron = getattr(client, "diet_pattern", None)
    comidas = getattr(client, "meals_per_day", None)
    alergias = [_norm(a) for a in (getattr(client, "food_allergies", None) or []) if a]
    # Segunda fuente: lo que el cliente registra en el portal (peso y constancia)
    señales = portal_signals(db, client)

    # Palabras de molestia → término que debe aparecer en el caso de la plantilla
    SENALES = [
        ("rodilla", "rodilla"), ("menisco", "rodilla"), ("artrosis", "artrosis"),
        ("hombro", "hombro"), ("manguito", "hombro"),
        ("lumbar", "lumbar"), ("lumbalgia", "lumbar"), ("espalda", "espalda"),
        ("hernia", "hernia"), ("cervical", "cervical"), ("cuello", "cervical"),
        ("tension", "tension"), ("hipertens", "tension"),
        ("diabet", "diabetes"), ("embaraz", "embarazo"), ("posparto", "posparto"),
        ("menopaus", "menopausia"), ("osteoporosis", "osteoporosis"),
        ("escoliosis", "escoliosis"), ("codo", "codo"), ("fascitis", "fascitis"),
    ]
    buscadas = {termino for palabra, termino in SENALES if palabra in molestias}

    out: list[dict] = []
    for t in db.scalars(select(PlanTemplate)):
        # Con servicio de SOLO DIETA el entreno no decide nada: se ordena por
        # objetivo y perfil; con entrenamiento, el encaje de la rutina manda.
        score = 0.0
        motivos: list[str] = []

        if t.category == cat:
            score += 5
            motivos.append("el mismo objetivo: " + _GOAL_WORD.get(cat, cat))
        else:
            score -= 1.5

        if pkgs.has_training(tier):
            if dias and t.days_per_week:
                dif = abs(int(dias) - int(t.days_per_week))
                score += {0: 3.0, 1: 1.5}.get(dif, -1.0 * dif)
                if dif == 0:
                    motivos.append(f"{dias} días, los que puede entrenar")
            if lugar and t.training_place:
                if lugar == t.training_place:
                    score += 2.0
                    motivos.append("en casa" if lugar == "home" else "en gimnasio")
                else:
                    score -= 2.5
            if nivel and t.level:
                dif = abs(_LEVEL_ORDER.get(nivel, 0) - _LEVEL_ORDER.get(t.level, 0))
                score += {0: 2.0, 1: 0.5}.get(dif, -1.5)
                if dif == 0:
                    motivos.append({"beginner": "nivel principiante",
                                    "intermediate": "nivel intermedio",
                                    "advanced": "nivel avanzado"}[nivel])

        caso = _norm(t.case_note or "") + " " + _norm(t.title or "") + " " + _norm(t.diet_focus or "")
        for termino in buscadas:
            if termino in caso:
                score += 4.0   # una molestia declarada pesa mucho: es un filtro
                motivos.append(f"contempla su {termino}")
        if edad is not None and edad >= 55 and ("55" in caso or "60" in caso or "65" in caso):
            score += 2.0
            motivos.append("pensada para su edad")

        # --- eje DIETA: manda cuando el servicio la incluye -------------------
        if pkgs.has_nutrition(tier):
            if patron and t.diet_pattern == patron:
                score += 4.0
                motivos.append("una dieta " + _PATRON_PALABRA.get(patron, patron))
            elif patron and t.diet_pattern and t.diet_pattern != patron:
                score -= 2.0
            if comidas and t.meals_per_day:
                if int(comidas) == int(t.meals_per_day):
                    score += 2.0
                    motivos.append(f"{comidas} comidas al día, como pide")
                else:
                    score -= 0.5 * abs(int(comidas) - int(t.meals_per_day))
            for alergia in alergias:
                if alergia and alergia in caso:
                    score += 3.0
                    motivos.append(f"ya contempla lo de {alergia}")

        # --- lo que dice el PORTAL: evolución de peso y constancia ------------
        for clave, palabras, peso, motivo in _REGLAS_PORTAL:
            if not señales.get(clave):
                continue
            # Las señales de entrenamiento no dicen nada de quien solo tiene
            # dieta: su portal ni siquiera registra sesiones.
            if clave == "entrena_menos" and not pkgs.has_training(tier):
                continue
            if any(p in caso for p in palabras):
                score += peso
                motivos.append(motivo)
        if (señales.get("entrena_menos") and pkgs.has_training(tier)
                and t.days_per_week and dias):
            if t.days_per_week < int(dias):        # menos días = más probable que los cumpla
                score += 1.5
                motivos.append("menos días, que es lo que está cumpliendo")

        if score <= 0:
            continue
        out.append({
            "id": t.id, "title": t.title, "category": t.category,
            "case_note": t.case_note, "level": t.level,
            "days_per_week": t.days_per_week, "training_place": t.training_place,
            "meals_per_day": t.meals_per_day, "diet_pattern": t.diet_pattern,
            "diet_focus": t.diet_focus, "source": t.source,
            "summary": _summary(t, tier),
            "score": round(score, 2),
            "why": _why(motivos),
        })
    out.sort(key=lambda x: (-x["score"], x["title"]))
    return out[:limit]


# Señal del portal → palabras que debe traer el caso para responder a ella.
_REGLAS_PORTAL = [
    ("estancado", ("cardio", "pasos", "resisten", "estancad", "fallidas", "restrictiva"),
     3.0, "su peso lleva días plano y esta ataca justo eso"),
    ("adherencia_baja", ("picoteo", "come fuera", "familia", "sencill", "menu", "decisiones",
                         "ansiedad", "noche"),
     3.0, "la dieta se le escapa y esta es más fácil de cumplir"),
    ("registra_poco", ("minutos", "expres", "2 dias", "corto", "habito", "20 minutos"),
     2.5, "está registrando poco: una más ligera engancha mejor"),
    ("entrena_menos", ("minutos", "expres", "2 dias", "corto", "habito"),
     2.5, "está entrenando menos días de los previstos"),
]


def _summary(t: PlanTemplate, tier: str) -> str:
    """Resumen de una línea de LA PLANTILLA (qué es), para que el coach decida
    de un vistazo sin abrirla."""
    from app.services import packages as pkgs

    trozos = [_GOAL_WORD.get(t.category, t.category).capitalize()]
    if pkgs.has_training(tier) and t.days_per_week:
        trozos.append(f"{t.days_per_week} días")
        trozos.append("en casa" if t.training_place == "home" else "en gimnasio")
        if t.level:
            trozos.append({"beginner": "principiante", "intermediate": "intermedio",
                           "advanced": "avanzado"}.get(t.level, t.level))
    if pkgs.has_nutrition(tier) and t.meals_per_day:
        trozos.append(f"{t.meals_per_day} comidas")
        if t.diet_pattern:
            trozos.append(t.diet_pattern)
    return " · ".join(trozos)


def _why(motivos: list[str]) -> str:
    """Frase corta y honesta del porqué (sin repetir ni inventar)."""
    vistos: list[str] = []
    for m in motivos:
        if m not in vistos:
            vistos.append(m)
    if not vistos:
        return "Encaja de forma general con su ficha."
    if len(vistos) == 1:
        return f"Encaja por {vistos[0]}."
    return "Encaja por " + ", ".join(vistos[:-1]) + " y " + vistos[-1] + "."
